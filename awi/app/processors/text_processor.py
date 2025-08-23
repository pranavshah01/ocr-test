"""
Enhanced Text Processor for DOCX document processing.

This module provides comprehensive text processing capabilities for DOCX documents including:
- Pattern matching using regex patterns
- Text replacement with formatting preservation
- Layout impact analysis for text changes
- Support for both append and replace modes
- Font and formatting preservation during replacements

The processor analyzes document structure, finds patterns in text content, applies replacements
while maintaining document formatting, and provides detailed analysis of potential layout impacts
from text modifications.
"""

import re
import logging
import time
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.section import Section
from docx.enum.section import WD_SECTION_START

from ..core.processor_interface import BaseProcessor, ProcessingResult
from ..utils.pattern_matcher import PatternMatcher, create_pattern_matcher
from ..utils.text_utils.text_docx_utils import (
    TextReconstructor, FontManager, TextReplacer, create_text_replacer
)
from ..utils.shared_constants import (
    DEFAULT_MAPPING,
    DEFAULT_SEPARATOR,
    PROCESSING_MODES
)

logger = logging.getLogger(__name__)

@dataclass
class EnhancedMatch:
    """
    Enhanced match information with comprehensive details about text replacements.
    
    This class extends basic pattern matching with additional context including
    font details, and confidence scores. It provides a complete record of each text replacement
    operation for reporting and analysis purposes.
    
    Attributes:
        pattern: The regex pattern that was matched
        original_text: The original text found in the document
        replacement_text: The text that will replace the original
        position: Character position where the match was found
        location: Human-readable location description (e.g., "body_paragraph_5")
        font_info: Dictionary containing font properties (name, size, style, etc.)
        confidence: Confidence score for the match (0.0 to 1.0)
        actual_pattern: The actual regex pattern string
        content_type: Type of content (Paragraph, Table, Header, Footer, etc.)
        dimension: Element dimensions
        processor: Name of the processor that found this match
    """
    pattern: str
    original_text: str
    replacement_text: str
    position: int
    location: str
    font_info: Dict[str, Any] = None
    confidence: Optional[float] = None
    actual_pattern: str = ""
    content_type: str = "Paragraph"
    dimension: str = ""
    processor: str = "Text"
    
    def __post_init__(self):
        """
        Initialize default values after object creation.
        
        This method is automatically called after the dataclass is created to ensure
        that font_info is always a dictionary, even if not provided during initialization.
        """
        if self.font_info is None:
            self.font_info = {}
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            'pattern': self.pattern,
            'original_text': self.original_text,
            'replacement_text': self.replacement_text,
            'position': self.position,
            'location': self.location,
            'font_info': self.font_info,
            'confidence': self.confidence,
            'actual_pattern': self.actual_pattern,
            'content_type': self.content_type,
            'dimension': self.dimension,
            'processor': self.processor
        }







class TextProcessor(BaseProcessor):
    """
    Enhanced text processor for DOCX documents with comprehensive analysis capabilities.
    
    This class provides the main interface for processing DOCX documents. It handles
    pattern matching, text replacement, and maintains detailed information about all 
    text modifications for reporting purposes.
    
    The processor integrates with core infrastructure components including memory
    management, performance monitoring, and GPU acceleration for optimal performance.
    """
    
    def __init__(self, patterns: Dict[str, Any] = None, mappings: Dict[str, Any] = None, 
                 mode: str = PROCESSING_MODES['APPEND'], separator: str = DEFAULT_SEPARATOR, default_mapping: str = DEFAULT_MAPPING):
        """
        Initialize the text processor with patterns, mappings, and processing mode.
        
        Args:
            patterns: Dictionary containing regex patterns for text matching
            mappings: Dictionary containing text mappings for replacements
            mode: Processing mode - "append" or "replace"
            separator: Separator between original and appended text in append mode
            default_mapping: Default text to append when no mapping is found
        """
        # Initialize base processor
        config = {
            'patterns': patterns or {},
            'mappings': mappings or {},
            'mode': mode,
            'separator': separator,
            'default_mapping': default_mapping
        }
        super().__init__("text", config)
        
        self.patterns = patterns or {}
        self.mappings = mappings or {}
        self.mode = mode
        self.separator = separator
        self.default_mapping = default_mapping
        
        # Initialize components
        self.pattern_matcher = None
        self.text_replacer = None
        self.font_manager = None
        
        logger.info(f"Text processor initialized with mode: {mode}, separator: '{separator}', default_mapping: '{default_mapping}'")
    
    def initialize(self, **kwargs) -> bool:
        """
        Initialize the processor with configuration parameters.
        
        This method loads patterns and mappings, configures the processing mode,
        and sets up all internal components for document processing. It can accept
        patterns and mappings directly or load them from JSON files.
        
        Args:
            **kwargs: Configuration parameters including:
                - patterns: Dictionary of pattern names to regex patterns
                - mappings: Dictionary of original text to replacement text
                - mode: Processing mode ('append' or 'replace')
                - separator: Separator between original and appended text in append mode
                - default_mapping: Default text to append when no mapping is found
                - patterns_file: Path to patterns JSON file (if patterns not provided)
                - mappings_file: Path to mappings JSON file (if mappings not provided)
        
        Returns:
            True if initialization was successful, False otherwise
        """
        try:
            logger.info("Initializing Text Processor...")
            
            # Extract configuration parameters
            patterns = kwargs.get('patterns', self.patterns)
            mappings = kwargs.get('mappings', self.mappings)
            mode = kwargs.get('mode', self.mode)
            separator = kwargs.get('separator', self.separator)
            default_mapping = kwargs.get('default_mapping', self.default_mapping)
            
            # Load patterns and mappings if not provided
            if not patterns or not mappings:
                patterns, mappings = self._load_patterns_and_mappings()
            
            self.patterns = patterns
            self.mappings = mappings
            self.mode = mode
            self.separator = separator
            self.default_mapping = default_mapping
            
            # Initialize components using utility classes
            self.pattern_matcher = create_pattern_matcher(patterns, mappings, enhanced_mode=True)
            self.text_replacer = create_text_replacer(mode, separator, default_mapping)
            
            self.initialized = True
            logger.info(f"Text processor initialized with {len(patterns)} patterns, {len(mappings)} mappings, mode: {mode}")
            return True
            
        except Exception as e:
            logger.error(f"Error initializing Text Processor: {e}")
            self.initialized = False
            return False
    
    def _load_patterns_and_mappings(self) -> Tuple[Dict[str, str], Dict[str, str]]:
        """
        Load patterns and mappings from JSON configuration files.
        
        This method reads the patterns.json and mapping.json files from the current
        directory. Patterns contain regex expressions for text matching, while mappings
        define the replacement text for each pattern. If files are not found, empty
        dictionaries are returned and warnings are logged.
        
        Returns:
            Tuple containing (patterns dictionary, mappings dictionary)
        """
        import json
        
        # Load patterns
        patterns_path = Path("patterns.json")
        if patterns_path.exists():
            with open(patterns_path, 'r') as f:
                patterns = json.load(f)
            logger.info(f"Loaded {len(patterns)} patterns from {patterns_path}")
        else:
            patterns = {}
            logger.warning(f"Patterns file not found: {patterns_path}")
        
        # Load mappings
        mappings_path = Path("mapping.json")
        if mappings_path.exists():
            with open(mappings_path, 'r') as f:
                mappings = json.load(f)
            logger.info(f"Loaded {len(mappings)} mappings from {mappings_path}")
        else:
            mappings = {}
            logger.warning(f"Mappings file not found: {mappings_path}")
        
        return patterns, mappings
    
    def process_document_text(self, document: Document) -> Tuple[List[EnhancedMatch], List[Dict[str, Any]]]:
        """
        Process all text content in the document with comprehensive analysis.
        
        This method processes the entire document by analyzing paragraphs, tables, headers,
        and footers for pattern matches. It tracks all detections including duplicates and 
        skipped ones for comprehensive reporting.
        
        Args:
            document: The DOCX document object to process
            
        Returns:
            Tuple of (List of EnhancedMatch objects, List of all detections including skipped ones)
        """
        matches = []
        all_detections = []  # Track all detections including skipped ones
        
        logger.info("Starting enhanced text processing...")
        logger.debug(f"TEXT PROCESSOR DEBUG: Document has {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
        
        # Process main document body paragraphs
        body_matches, body_detections = self._process_paragraphs(document.paragraphs, "body")
        matches.extend(body_matches)
        all_detections.extend(body_detections)
        
        # Process tables
        table_matches, table_detections = self._process_tables(document.tables, "")
        matches.extend(table_matches)
        all_detections.extend(table_detections)
        
        # Process headers and footers
        header_footer_matches, header_footer_detections = self._process_headers_footers(document)
        matches.extend(header_footer_matches)
        all_detections.extend(header_footer_detections)
        
        logger.info(f"Enhanced text processing completed: {len(matches)} matches found, {len(all_detections)} total detections")
        
        return matches, all_detections
    
    def _process_paragraphs(self, paragraphs: List[Paragraph], location: str) -> Tuple[List[EnhancedMatch], List[Dict[str, Any]]]:
        """
        Process a list of paragraphs for pattern matches.
        
        Args:
            paragraphs: List of paragraphs to process
            location: Location description for logging
            
        Returns:
            Tuple of (List of EnhancedMatch objects, List of all detections including skipped ones)
        """
        matches = []
        all_detections = []
        
        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_matches, paragraph_detections = self._process_paragraph(
                    paragraph, f"{location}_paragraph_{i}"
                )
                matches.extend(paragraph_matches)
                all_detections.extend(paragraph_detections)
            except Exception as e:
                logger.error(f"Error processing paragraph {i} in {location}: {e}")
        
        return matches, all_detections
    
    def _process_paragraph(self, paragraph: Paragraph, location: str) -> Tuple[List[EnhancedMatch], List[Dict[str, Any]]]:
        """
        Process individual paragraph for pattern matches.
        
        Args:
            paragraph: Paragraph to process
            location: Location description
            
        Returns:
            Tuple of (List of EnhancedMatch objects, List of all detections including skipped ones)
        """
        matches = []
        all_detections = []
        
        if not paragraph.runs:
            return matches, all_detections
        
        # Reconstruct full text from runs
        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)
        
        if not full_text.strip():
            return matches, all_detections
        
        # Find pattern matches (only those with mappings for processing)
        pattern_matches = self.pattern_matcher.find_matches(full_text)
        
        # Find ALL pattern matches (including those without mappings for comprehensive reporting)
        all_pattern_matches = self.pattern_matcher.find_all_pattern_matches(full_text)
        
        # DEBUG: Log what we found
        logger.debug(f"TEXT PROCESSOR DEBUG: Paragraph '{location}' has {len(pattern_matches)} pattern_matches and {len(all_pattern_matches)} all_pattern_matches")
        if all_pattern_matches:
            logger.debug(f"TEXT PROCESSOR DEBUG: First pattern match: {all_pattern_matches[0]}")
        if full_text and len(full_text) > 100:
            logger.debug(f"TEXT PROCESSOR DEBUG: Paragraph text (first 100 chars): '{full_text[:100]}...'")
        elif full_text:
            logger.debug(f"TEXT PROCESSOR DEBUG: Paragraph text: '{full_text}'")
        
        # Create a set of matched texts for quick lookup
        matched_texts = {match[1] for match in pattern_matches}
        
        # Add all detections to all_detections list (including non-matched ones)
        for pattern_name, matched_text, start_pos, end_pos in all_pattern_matches:
            # Check if this has a mapping
            replacement_text = self.pattern_matcher.get_replacement(matched_text)
            
            # Get font information for this detection using utility class
            font_info = {}
            try:
                # Find the text in the runs to get font info
                text_span = TextReconstructor.find_text_in_runs(runs, matched_text, start_pos)
                if text_span and text_span[2]:  # affected_runs
                    affected_runs = text_span[2]
                    # Use the utility class to get the best font info
                    font_info = FontManager.get_best_font_info_from_runs(affected_runs, getattr(self, 'document', None))
                else:
                    # Fallback to default font info if no runs found
                    font_info = FontManager.get_default_font_info()
            except Exception as e:
                logger.debug(f"Could not extract font info for detection '{matched_text}': {e}")
                # Fallback to default font info
                font_info = FontManager.get_default_font_info()
            
            # Get the actual pattern from patterns
            actual_pattern = self.patterns.get(pattern_name, pattern_name)
            
            # Determine content type and dimensions
            content_type = "Paragraph"
            dimension = ""
            
            if "table" in location.lower():
                content_type = "Table"
                # For tables, we'll need to get cell dimensions
                dimension = "Cell size"  # Placeholder - would need table cell info
            elif "header" in location.lower():
                content_type = "Header"
            elif "footer" in location.lower():
                content_type = "Footer"
            
            all_detections.append({
                'pattern_name': pattern_name,
                'actual_pattern': actual_pattern,
                'matched_text': matched_text,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'replacement_text': replacement_text,
                'location': location,
                'content_type': content_type,
                'dimension': dimension,
                'processor': 'Text',
                'font_info': font_info
            })
        
        # Process each match (including those that should use default mapping)
        for pattern_name, matched_text, start_pos, end_pos in all_pattern_matches:
            try:
                replacement_text = self.pattern_matcher.get_replacement(matched_text)
                
                # Mode-specific behavior:
                # - In "replace" mode: only process if we have a valid mapping (no default)
                # - In "append" mode: use default mapping if no valid mapping found
                if not replacement_text:
                    if self.mode == "replace":
                        # In replace mode, skip if no mapping found
                        logger.debug(f"REPLACE MODE: Skipping '{matched_text}' - no mapping found")
                        continue
                    elif self.mode == PROCESSING_MODES['APPEND']:
                        # In append mode, use default mapping
                        replacement_text = self.default_mapping
                        logger.info(f"APPEND MODE: Using default mapping '{self.default_mapping}' for '{matched_text}'")
                    else:
                        # Unknown mode, skip
                        logger.warning(f"Unknown mode '{self.mode}', skipping '{matched_text}'")
                        continue
                
                # Check for duplicates
                # The deduplication logic is removed, so we just add all detections
                # Note: all_detections are now added earlier in the function for comprehensive reporting
                
                # Find affected runs
                text_span = TextReconstructor.find_text_in_runs(runs, matched_text, start_pos)
                if not text_span:
                    continue
                
                span_start, span_end, affected_runs = text_span
                
                # Get font information from the affected runs
                font_info = {}
                if affected_runs:
                    # Find the best run with explicit font properties
                    best_font_info = None
                    best_score = 0
                    
                    for run in affected_runs:
                        run_font_info = FontManager.get_font_info(run)
                        
                        # Score this run's font info (higher score = better)
                        score = 0
                        if run_font_info.get('font_family') and run_font_info.get('font_family') != 'Arial':
                            score += 2  # Explicit font family
                        if run_font_info.get('font_size') and run_font_info.get('font_size') != '12.0':
                            score += 2  # Explicit font size
                        if run_font_info.get('font_family') == 'Arial' and run_font_info.get('font_size') == '12.0':
                            score = 0  # Default values, lowest priority
                        
                        if score > best_score:
                            best_score = score
                            best_font_info = run_font_info
                    
                    # Use the best font info found, or the first run's info if all are defaults
                    if best_font_info and best_score > 0:
                        font_info = best_font_info
                    else:
                        font_info = FontManager.get_font_info(affected_runs[0])
                    
                    # If still no explicit font info found, try to get from document styles
                    if (font_info.get('font_family') == 'Arial' and 
                        font_info.get('font_size') == '12.0' and 
                        hasattr(self, 'document')):
                        default_font_info = FontManager.get_document_default_font_info(self.document)
                        # Only use default if we don't have explicit font info
                        if font_info.get('font_family') == 'Arial':
                            font_info['font_family'] = default_font_info.get('font_family', 'Arial')
                        if font_info.get('font_size') == '12.0':
                            font_info['font_size'] = default_font_info.get('font_size', '12.0')
                
                # Perform text replacement
                success = self.text_replacer.replace_text_in_runs(
                    runs, matched_text, replacement_text, span_start, span_end
                )
                
                # DEBUG: Log text replacement success/failure
                logger.debug(f"TEXT PROCESSOR DEBUG: Text replacement for '{matched_text}' -> '{replacement_text}' at {location}: {'SUCCESS' if success else 'FAILED'}")
                
                if success:
                    # Determine mapping type
                    mapping_type = "Default" if replacement_text == self.default_mapping else "Mapped"
                    # Update corresponding detection in-place for 1:1 XML↔detection
                    actual_pattern = self.patterns.get(pattern_name, pattern_name)
                    updated_detection = False
                    for detection in all_detections:
                        if (detection.get('matched_text') == matched_text and
                            detection.get('location') == location and
                            detection.get('pattern_name') == pattern_name and
                            detection.get('start_pos') == start_pos and
                            detection.get('end_pos') == end_pos):
                            detection['is_matched'] = True
                            detection['replacement_text'] = replacement_text
                            detection['actual_pattern'] = actual_pattern
                            # Ensure font_info exists
                            if 'font_info' not in detection or not isinstance(detection['font_info'], dict):
                                detection['font_info'] = {}
                            detection['font_info'].update(font_info or {})
                            detection['font_info']['normalized'] = True
                            # Optionally include preview/length
                            updated_detection = True
                            break
                    if not updated_detection:
                        logger.warning(f"Could not update detection for '{matched_text}' at {location}; separate match record skipped to preserve 1:1 policy")
                    # Logging
                    logger.info(f"Text replacement ({self.mode} mode, {mapping_type}): '{matched_text}' -> '{replacement_text}' at {location}")
                
            except Exception as e:
                logger.error(f"Error processing match '{matched_text}' at {location}: {e}")
        
        return matches, all_detections
    
    def _process_tables(self, tables: List[Table], location_prefix: str) -> Tuple[List[EnhancedMatch], List[Dict[str, Any]]]:
        """
        Process tables in the document.
        
        Args:
            tables: List of tables to process
            location_prefix: Prefix for location description
            
        Returns:
            Tuple of (List of EnhancedMatch objects, List of all detections including skipped ones)
        """
        matches = []
        all_detections = []
        
        for table_idx, table in enumerate(tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if location_prefix:
                            location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                        else:
                            location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                        cell_matches, cell_detections = self._process_paragraphs(cell.paragraphs, location)
                        matches.extend(cell_matches)
                        all_detections.extend(cell_detections)
            except Exception as e:
                logger.error(f"Error processing table {table_idx} in {location_prefix}: {e}")
        
        return matches, all_detections
    
    def _process_headers_footers(self, document: Document) -> Tuple[List[EnhancedMatch], List[Dict[str, Any]]]:
        """
        Process headers and footers in the document.
        
        Args:
            document: Document to process
            
        Returns:
            Tuple of (List of EnhancedMatch objects, List of all detections including skipped ones)
        """
        matches = []
        all_detections = []
        
        try:
            # Process headers
            for section_idx, section in enumerate(document.sections):
                # Primary header
                if section.header:
                    header_matches, header_detections = self._process_paragraphs(
                        section.header.paragraphs, 
                        f"header_section_{section_idx}"
                    )
                    matches.extend(header_matches)
                    all_detections.extend(header_detections)
                    
                    header_table_matches, header_table_detections = self._process_tables(
                        section.header.tables,
                        f"header_section_{section_idx}"
                    )
                    matches.extend(header_table_matches)
                    all_detections.extend(header_table_detections)
                
                # Primary footer
                if section.footer:
                    footer_matches, footer_detections = self._process_paragraphs(
                        section.footer.paragraphs,
                        f"footer_section_{section_idx}"
                    )
                    matches.extend(footer_matches)
                    all_detections.extend(footer_detections)
                    
                    footer_table_matches, footer_table_detections = self._process_tables(
                        section.footer.tables,
                        f"footer_section_{section_idx}"
                    )
                    matches.extend(footer_table_matches)
                    all_detections.extend(footer_table_detections)
                    
        except Exception as e:
            logger.error(f"Error processing headers/footers: {e}")
        
        return matches, all_detections
    
    def process(self, document_or_path, **kwargs) -> ProcessingResult:
        """
        Process a DOCX document with comprehensive text replacement and analysis.
        
        This is the main entry point for document processing. It handles both
        Document objects and file paths, performs text processing with pattern 
        matching and replacement, and returns comprehensive results including 
        match information and metadata.
        
        Args:
            document_or_path: Either a Document object or Path to the DOCX file
            **kwargs: Additional processing parameters
            
        Returns:
            ProcessingResult containing comprehensive processing results
            
        Raises:
            RuntimeError: If the processor is not initialized
        """
        if not self.initialized:
            raise RuntimeError("Text Processor not initialized")
        
        start_time = time.time()
        
        # Handle both Document objects and file paths
        if isinstance(document_or_path, Path):
            document_path = document_or_path
            logger.info(f"Processing DOCX document from path: {document_path.name}")
            try:
                document = Document(document_path)
            except Exception as e:
                logger.error(f"Failed to load document from path {document_path}: {e}")
                return ProcessingResult(
                    success=False,
                    processor_type="text",
                    matches_found=0,
                    processing_time=time.time() - start_time,
                    error_message=f"Failed to load document: {e}"
                )
        else:
            # Assume it's a Document object
            document = document_or_path
            document_path = None
            logger.info("Processing DOCX document object")
        
        try:
            # Store document for font information access
            self.document = document
            
            # Process the document
            matches, all_detections = self.process_document_text(document)
            
            # DEBUG: Log the exact number of matches and detections found
            logger.debug(f"TEXT PROCESSOR DEBUG: Found {len(matches)} matches and {len(all_detections)} detections")
            if matches:
                logger.debug(f"TEXT PROCESSOR DEBUG: First match pattern: {matches[0].pattern}, original_text: {matches[0].original_text}")
            if all_detections:
                logger.debug(f"TEXT PROCESSOR DEBUG: First detection: {all_detections[0]}")
            
            # Save the processed document to the processed directory (only if we have a path)
            output_path = None
            if document_path:
                from config import OUTPUT_DIR
                processed_dir = Path(OUTPUT_DIR)
                processed_dir.mkdir(parents=True, exist_ok=True)
                output_path = processed_dir / f"{document_path.stem}_processed{document_path.suffix}"
                document.save(output_path)
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            # Prepare metadata
            metadata = {
                'matches': [match.to_dict() for match in matches],
                'all_detections': all_detections,  # Add all detections to metadata
                'detailed_matches': [match.to_dict() for match in matches],  # Add detailed matches for reporting
                'file_size_mb': document_path.stat().st_size / (1024 * 1024) if document_path else 0
            }

            # Prefer reporting detections count so detections without replacement are visible in reports
            reported_matches = len(all_detections) if len(all_detections) > 0 else len(matches)

            logger.info(f"Text processing completed: {reported_matches} matches found (replacements: {len(matches)}, detections: {len(all_detections)})")

            return ProcessingResult(
                success=True,
                processor_type="text",
                matches_found=reported_matches,
                output_path=output_path,
                processing_time=processing_time,
                metadata=metadata
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Error processing document {document_path}: {e}")
            
            return ProcessingResult(
                success=False,
                processor_type="text",
                matches_found=0,
                processing_time=processing_time,
                error_message=str(e)
            )
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of file formats supported by this processor.
        
        Returns:
            List of supported file extensions
        """
        return ['.docx']
    
    def get_processing_info(self) -> Dict[str, Any]:
        """Get information about the text processor configuration."""
        return {
            'mode': self.mode,
            'patterns_count': len(self.patterns),
            'mappings_count': len(self.mappings),
            'compiled_patterns': len(self.pattern_matcher.compiled_patterns) if self.pattern_matcher else 0,
            'is_initialized': self.is_initialized
        }
    
    def cleanup(self):
        """
        Clean up processor resources and reset internal state.
        
        This method releases all resources used by the processor including pattern
        matchers, text replacers, and layout analyzers. It resets the initialization
        state to allow for proper cleanup and potential re-initialization.
        
        The cleanup process is logged for debugging purposes and any errors during
        cleanup are caught and logged without raising exceptions.
        """
        try:
            logger.info("Cleaning up Text Processor...")
            self.is_initialized = False
            self.pattern_matcher = None
            self.text_replacer = None
            logger.info("Text Processor cleanup completed")
        except Exception as e:
            logger.error(f"Error during Text Processor cleanup: {e}")

def create_text_processor(patterns: Dict[str, Any] = None, mappings: Dict[str, Any] = None, 
                         mode: str = PROCESSING_MODES['APPEND'], separator: str = DEFAULT_SEPARATOR, default_mapping: str = DEFAULT_MAPPING) -> TextProcessor:
    """
    Factory function to create a configured TextProcessor instance.
    
    This factory function provides a convenient way to create TextProcessor instances
    with proper configuration. It encapsulates the creation logic and ensures that
    all necessary dependencies are properly initialized.
    
    Args:
        patterns: Dictionary containing regex patterns for text matching
        mappings: Dictionary containing text mappings for replacements
        mode: Processing mode - "append" or "replace"
        separator: Separator between original and appended text in append mode
        default_mapping: Default text to append when no mapping is found
        
    Returns:
        A fully configured TextProcessor instance ready for initialization and use
    """
    return TextProcessor(patterns, mappings, mode, separator, default_mapping)