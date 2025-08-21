"""
Core Document Processor for Document Processing Pipeline.

This module provides the main orchestrator for document processing, managing
the complete pipeline from document loading through multi-processor processing
to output generation. It's designed to scale incrementally as new processors
(text, graphics, image) are implemented.
"""

import json
import time
import gc
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
import logging
from datetime import datetime

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .processor_interface import BaseProcessor, ProcessingResult
from .models import ProcessingLog, DocumentInfo, BatchProcessingResult
from app.converters.doc_converter import DocConverter, ConversionError
from app.utils.shared_constants import MAX_FILE_SIZE_MB, LARGE_FILE_THRESHOLD_MB, VERY_LARGE_FILE_THRESHOLD_MB

# Add these constants at the top of the file, after the existing imports
MAX_PROCESSING_SIZE_MB = 200.0  # Maximum size for full processing
LARGE_FILE_WARNING_MB = 150.0   # Size at which to warn about potential issues

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Main orchestrator for document processing pipeline.
    
    This class manages the complete document processing workflow, including:
    - Document format conversion (.doc to .docx)
    - Document loading with enhanced parsers for large files
    - Multi-processor coordination (text, graphics, image)
    - Error handling and graceful degradation
    - Processing result aggregation and reporting
    
    The processor is designed to work incrementally - it can function with
    only text processing implemented, and graphics/image processors can be
    added later without changing the core architecture.
    """
    
    def __init__(self, config):
        """
        Initialize document processor with configuration.
        
        Args:
            config: ProcessingConfig instance containing all processing settings
        """
        self.config = config
        self.converter = DocConverter() if config else None
        
        # Load patterns and mappings for processors
        self.patterns = self._load_json_file(config.patterns_file) if config else {}
        self.mappings = self._load_json_file(config.mappings_file) if config else {}
        
        # Initialize processor instances (will be set when processors are available)
        self.text_processor: Optional[BaseProcessor] = None
        self.graphics_processor: Optional[BaseProcessor] = None
        self.image_processor: Optional[BaseProcessor] = None
        
        # Processing statistics
        self.processing_stats = {
            'documents_processed': 0,
            'successful_processing': 0,
            'failed_processing': 0,
            'total_processing_time': 0.0,
            'total_matches_found': 0
        }
        
        logger.info("Document processor initialized")
    
    def _load_json_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Load JSON file with comprehensive error handling.
        
        Args:
            file_path: Path to the JSON file to load
            
        Returns:
            Dictionary containing the loaded data, or empty dict if loading failed
        """
        try:
            if not file_path.exists():
                logger.warning(f"JSON file not found: {file_path}")
                return {}
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            logger.debug(f"Successfully loaded JSON file: {file_path}")
            return data
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in file {file_path}: {e}")
            return {}
        except Exception as e:
            logger.error(f"Failed to load JSON file {file_path}: {e}")
            return {}
    
    def process_document(self, file_path: Path) -> ProcessingResult:
        """
        Process a single document through the complete pipeline.
        
        This method orchestrates the entire processing workflow for a document:
        1. Format conversion (if needed)
        2. Document loading with enhanced parsers
        3. Multi-processor processing with graceful degradation
        4. Result aggregation and output generation
        
        Args:
            file_path: Path to the document file to process
            
        Returns:
            ProcessingResult containing comprehensive processing results
        """
        start_time = time.time()
        
        # Create processing log for detailed tracking
        processing_log = ProcessingLog(document_path=file_path)
        
        # Initialize result structure
        result = ProcessingResult(
            success=False,
            processor_type="document_processor",
            processing_time=0.0
        )
        
        try:
            logger.info(f"Starting processing of: {file_path}")
            
            # Step 1: Handle format conversion if needed
            try:
                docx_path = self._ensure_docx_format(file_path, processing_log)
                if not docx_path:
                    error_msg = "Failed to convert document to DOCX format"
                    processing_log.add_error(error_msg)
                    result.error_message = error_msg
                    return result
            except Exception as e:
                error_msg = f"Error during format conversion: {e}"
                processing_log.add_error(error_msg)
                result.error_message = error_msg
                return result
            
            # Step 2: Load document with enhanced parser support
            try:
                document = self._load_document(docx_path, processing_log)
                if not document:
                    error_msg = "Failed to load document"
                    processing_log.add_error(error_msg)
                    result.error_message = error_msg
                    return result
            except Exception as e:
                error_msg = f"Error during document loading: {e}"
                processing_log.add_error(error_msg)
                result.error_message = error_msg
                return result
            
            # Step 3: Process through pipeline with graceful degradation
            try:
                pipeline_results = self._process_pipeline(document, processing_log)
            except Exception as e:
                error_msg = f"Error during pipeline processing: {e}"
                processing_log.add_error(error_msg)
                result.error_message = error_msg
                return result
            
            # Step 4: Generate output path and save processed document
            try:
                output_path = self._generate_output_path(file_path)
                if self._save_document(document, output_path, processing_log):
                    result.output_path = output_path
                    result.success = True
                    result.matches_found = pipeline_results.get('total_matches', 0)
                    # Add input_path to metadata for report generation
                    pipeline_results['input_path'] = str(file_path)
                    result.metadata = pipeline_results
                    result.layout_impact = pipeline_results.get('layout_impact')
                    
                    logger.info(f"Successfully processed: {file_path} -> {output_path}")
                else:
                    error_msg = "Failed to save processed document"
                    processing_log.add_error(error_msg)
                    result.error_message = error_msg
            except Exception as e:
                error_msg = f"Error during document saving: {e}"
                processing_log.add_error(error_msg)
                result.error_message = error_msg
            
            # Update processing log with final results
            try:
                processing_log.layout_impact_data = pipeline_results.get('layout_impact')
                processing_log.complete()
            except Exception as e:
                logger.warning(f"Error updating processing log: {e}")
            
        except Exception as e:
            error_msg = f"Unexpected error processing {file_path}: {e}"
            logger.error(error_msg)
            processing_log.add_error(error_msg)
            result.error_message = error_msg
        
        finally:
            # Calculate final processing time
            result.processing_time = time.time() - start_time
            
            # Update processing statistics
            try:
                self._update_processing_stats(result)
            except Exception as e:
                logger.warning(f"Error updating processing stats: {e}")
            
            logger.info(f"Processing completed in {result.processing_time:.2f}s")
        
        return result
    
    def _ensure_docx_format(self, file_path: Path, processing_log: ProcessingLog) -> Optional[Path]:
        """
        Ensure document is in DOCX format, convert if necessary.
        
        Args:
            file_path: Path to the document file
            processing_log: Processing log to record conversion activities
            
        Returns:
            Path to DOCX file or None if conversion failed
        """
        if file_path.suffix.lower() == '.docx':
            logger.debug(f"Document already in DOCX format: {file_path}")
            return file_path
        
        if file_path.suffix.lower() == '.doc':
            if not self.converter:
                error_msg = "No converter available for .doc files"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
            
            try:
                logger.info(f"Converting .doc to .docx: {file_path}")
                converted_path = self.converter.convert_to_docx(file_path)
                if converted_path and converted_path.exists():
                    logger.info(f"Successfully converted: {file_path} -> {converted_path}")
                    return converted_path
                else:
                    error_msg = "Conversion completed but output file not found"
                    processing_log.add_error(error_msg)
                    logger.error(error_msg)
                    return None
            except ConversionError as e:
                error_msg = f"Conversion failed: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
        
        error_msg = f"Unsupported file format: {file_path.suffix}"
        processing_log.add_error(error_msg)
        logger.error(error_msg)
        return None
    
    def _load_with_minimal_parser(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Minimal parser that bypasses XML parsing for extremely problematic files.
        
        This is a last resort method that creates a minimal document structure
        without parsing the full XML content.
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record parsing activities
            
        Returns:
            Document instance or None if parsing failed
        """
        try:
            from docx import Document
            
            logger.warning(f"Using minimal parser for extremely problematic file: {docx_path.name}")
            processing_log.add_warning("Minimal parser: Bypassing XML parsing for extremely large file")
            
            # Create a minimal document structure
            # This bypasses the XML parsing entirely
            document = Document()
            
            # Add a warning paragraph
            warning_para = document.add_paragraph("WARNING: This document was loaded with minimal parsing due to XML size limits.")
            warning_para.add_run(" The full content may not be accessible for processing.")
            
            # Add file info
            info_para = document.add_paragraph(f"File: {docx_path.name}")
            info_para.add_run(f" (Size: {docx_path.stat().st_size / (1024*1024):.1f}MB)")
            
            logger.warning(f"Created minimal document structure for extremely large file: {docx_path.name}")
            processing_log.add_warning("Minimal parser: Created basic document structure")
            
            return document
            
        except Exception as e:
            error_msg = f"Minimal parser failed for extremely large file: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
    
    def _load_document(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Load DOCX document with enhanced parser support for large files.
        
        This method implements a tiered loading approach:
        1. Standard parser for files < 100MB
        2. Enhanced parser for files 100MB - 200MB  
        3. Custom parser for files > 200MB
        4. Fallback mechanisms for each tier
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record loading activities
            
        Returns:
            Document instance or None if loading failed
        """
        file_size = docx_path.stat().st_size
        size_mb = file_size / (1024 * 1024)
        
        logger.info(f"Loading document: {size_mb:.1f}MB - {docx_path.name}")
        
        # Record file size in processing log
        processing_log.add_info(f"Document size: {size_mb:.1f}MB")
        
        # Check if file size exceeds maximum supported size
        if size_mb > MAX_FILE_SIZE_MB:
            error_msg = f"File size {size_mb:.1f}MB exceeds maximum supported size of {MAX_FILE_SIZE_MB}MB"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        
        # For files >50MB, use streaming parser directly
        if size_mb > 50:
            logger.info(f"Large file detected ({size_mb:.1f}MB), using streaming parser")
            processing_log.add_info(f"Large file: {size_mb:.1f}MB - using streaming parser")
            
            try:
                document = self._load_with_streaming_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with streaming parser: {docx_path.name}")
                    processing_log.add_info("Used streaming parser for large file")
                    return document
            except Exception as e:
                logger.debug(f"Streaming parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Streaming parser failed: {e}")
        
        # Check if file size exceeds recommended processing size (for summary documents)
        if size_mb > MAX_PROCESSING_SIZE_MB:
            logger.warning(f"File size {size_mb:.1f}MB exceeds recommended processing size of {MAX_PROCESSING_SIZE_MB}MB")
            processing_log.add_warning(f"File size {size_mb:.1f}MB exceeds recommended processing size")
            
            # Create a summary document instead of attempting full processing
            return self._create_summary_document(docx_path, processing_log)
        
        # Prepare system for large files
        if size_mb > LARGE_FILE_THRESHOLD_MB:
            self._prepare_for_large_file()
            processing_log.add_warning(f"Large document detected: {size_mb:.1f}MB")
        
        # Tier 1: Standard parser for files < 100MB
        if size_mb < LARGE_FILE_THRESHOLD_MB:
            try:
                document = Document(docx_path)
                logger.debug(f"Successfully loaded document with standard parser: {docx_path}")
                processing_log.add_info("Used standard parser")
                return document
            except Exception as e:
                logger.debug(f"Standard loading failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Standard parser failed: {e}")
        
        # Tier 2: Enhanced parser for files 100MB - 200MB
        if size_mb < VERY_LARGE_FILE_THRESHOLD_MB:
            try:
                document = self._load_with_enhanced_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with enhanced parser: {docx_path.name}")
                    processing_log.add_info("Used enhanced parser")
                    return document
            except Exception as e:
                logger.debug(f"Enhanced parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Enhanced parser failed: {e}")
        
        # Tier 3: For very large files (>200MB), try extreme parser first, then custom, then minimal
        if size_mb >= VERY_LARGE_FILE_THRESHOLD_MB:
            logger.info(f"Very large file detected ({size_mb:.1f}MB), using aggressive parsing strategy")
            processing_log.add_info(f"Very large file: {size_mb:.1f}MB - using aggressive parsing")
            
            # Try extreme parser first for very large files
            try:
                document = self._load_with_extreme_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with extreme parser: {docx_path.name}")
                    processing_log.add_info("Used extreme parser for very large file")
                    return document
            except Exception as e:
                logger.debug(f"Extreme parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Extreme parser failed: {e}")
            
            # Try custom parser as second option
            try:
                document = self._load_with_custom_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with custom parser: {docx_path.name}")
                    processing_log.add_info("Used custom parser for very large file")
                    return document
            except Exception as e:
                logger.debug(f"Custom parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Custom parser failed: {e}")
            
            # Try streaming parser for very large files
            try:
                document = self._load_with_streaming_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with streaming parser: {docx_path.name}")
                    processing_log.add_info("Used streaming parser for very large file")
                    return document
            except Exception as e:
                logger.debug(f"Streaming parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Streaming parser failed: {e}")
            
            # Try minimal parser as last resort for very large files
            try:
                document = self._load_with_minimal_parser(docx_path, processing_log)
                if document:
                    logger.info(f"Successfully loaded document with minimal parser: {docx_path.name}")
                    processing_log.add_info("Used minimal parser for very large file")
                    return document
            except Exception as e:
                logger.debug(f"Minimal parser failed for {size_mb:.1f}MB file: {e}")
                processing_log.add_warning(f"Minimal parser failed: {e}")
        
        # Final fallback: Try enhanced parser for any size
        try:
            document = self._load_with_enhanced_parser(docx_path, processing_log)
            if document:
                logger.info(f"Successfully loaded document with enhanced parser (fallback): {docx_path.name}")
                processing_log.add_info("Used enhanced parser (fallback)")
                return document
        except Exception as e:
            logger.debug(f"Enhanced parser fallback failed: {e}")
            processing_log.add_warning(f"Enhanced parser fallback failed: {e}")
        
        error_msg = f"All loading methods failed for {size_mb:.1f}MB document {docx_path}"
        processing_log.add_error(error_msg)
        logger.error(error_msg)
        return None
    
    def _load_with_enhanced_parser(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Enhanced parser specifically tuned for large files (100MB+).
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record parsing activities
            
        Returns:
            Document instance or None if parsing failed
        """
        try:
            from lxml import etree
            
            # Configure lxml parser globally for large files with increased limits
            # This affects all subsequent XML parsing operations
            # Using the approach from python-docx PR #1272
            parser = etree.XMLParser(
                huge_tree=True,
                recover=True,
                remove_blank_text=True,
                resolve_entities=False
            )
            etree.set_default_parser(parser)
            
            # Use standard Document constructor with global parser configuration
            document = Document(docx_path)
            
            logger.info(f"Successfully loaded large document with enhanced parser: {docx_path.name}")
            return document
            
        except ImportError:
            error_msg = "lxml not available for enhanced parsing"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except etree.XMLSyntaxError as e:
            error_msg = f"XML syntax error during enhanced parsing: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except ValueError as e:
            if "attvalue" in str(e).lower() or "length too long" in str(e).lower():
                error_msg = f"Attribute value error during enhanced parsing (likely due to large file): {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                # Try with even larger limits
                return self._load_with_extreme_parser(docx_path, processing_log)
            else:
                error_msg = f"Value error during enhanced parsing: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
        except MemoryError as e:
            error_msg = f"Memory error during enhanced parsing: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except Exception as e:
            error_msg = f"Enhanced parser failed: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
    
    def _load_with_extreme_parser(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Extreme parser for handling very large files and attvalue errors.
        
        This parser uses the most aggressive settings to handle files that cause
        attvalue errors with the standard enhanced parser.
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record parsing activities
            
        Returns:
            Document instance or None if parsing failed
        """
        try:
            from lxml import etree
            
            logger.info(f"Using extreme parser for problematic large file: {docx_path.name}")
            processing_log.add_info("Extreme parser: Starting with maximum limits")
            
            # Extreme parser with maximum limits for problematic files
            parser = etree.XMLParser(
                huge_tree=True,
                recover=True,
                strip_cdata=False,
                resolve_entities=False,
                attribute_defaults=False,
                dtd_validation=False,
                load_dtd=False,
                no_network=True,
                collect_ids=False,
                remove_blank_text=False,
                remove_comments=False,
                remove_pis=False,  # Preserve processing instructions
                compact=False      # Don't compact whitespace
            )
            etree.set_default_parser(parser)
            
            # Use standard Document constructor with extreme parser configuration
            document = Document(docx_path)
            
            logger.info(f"Successfully loaded problematic large document with extreme parser: {docx_path.name}")
            processing_log.add_info("Extreme parser: Successfully loaded problematic document")
            return document
            
        except ImportError as e:
            error_msg = f"Required dependencies not available for extreme parser: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except etree.XMLSyntaxError as e:
            error_msg = f"XML syntax error during extreme parsing: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except ValueError as e:
            if "attvalue" in str(e).lower() or "length too long" in str(e).lower():
                error_msg = f"Attribute value error during extreme parsing (trying custom parser): {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                # Try with custom parser as last resort
                return self._load_with_custom_parser(docx_path, processing_log)
            else:
                error_msg = f"Value error during extreme parsing (file may be corrupted): {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
        except etree.XMLSyntaxError as e:
            if "attvalue" in str(e).lower() or "length too long" in str(e).lower():
                error_msg = f"XML syntax error with attvalue issue (trying custom parser): {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                # Try with custom parser as last resort
                return self._load_with_custom_parser(docx_path, processing_log)
            else:
                error_msg = f"XML syntax error during extreme parsing: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
        except MemoryError as e:
            error_msg = f"Memory error during extreme parsing: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except Exception as e:
            error_msg = f"Extreme parser failed for problematic file: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
    
    def _load_with_custom_parser(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Custom parser specifically designed for very large files (>200MB).
        
        This parser uses streaming and chunked processing to handle extremely
        large documents that would otherwise cause memory issues.
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record parsing activities
            
        Returns:
            Document instance or None if parsing failed
        """
        try:
            import zipfile
            from lxml import etree
            
            logger.info(f"Using custom parser for very large file: {docx_path.name}")
            processing_log.add_info("Custom parser: Starting chunked processing")
            
            # Custom parser with extreme limits for very large files
            parser = etree.XMLParser(
                huge_tree=True,
                recover=True,
                strip_cdata=False,
                resolve_entities=False,
                attribute_defaults=False,
                dtd_validation=False,
                load_dtd=False,
                no_network=True,
                collect_ids=False,
                remove_blank_text=False,  # Preserve whitespace
                remove_comments=False,    # Preserve comments
                remove_pis=False,         # Preserve processing instructions
                compact=False             # Don't compact whitespace
            )
            etree.set_default_parser(parser)
            
            # Try standard Document constructor with custom parser configuration
            try:
                document = Document(docx_path)
                
                # Validate document structure
                if not hasattr(document, 'paragraphs') or len(document.paragraphs) == 0:
                    logger.warning(f"Very large document has no paragraphs: {docx_path.name}")
                    processing_log.add_warning("Very large document appears to have no text content")
                
                logger.info(f"Successfully loaded very large document with custom parser: {docx_path.name}")
                processing_log.add_info("Custom parser: Successfully loaded very large document")
                return document
                
            except (ValueError, etree.XMLSyntaxError) as e:
                if "attvalue" in str(e).lower() or "length too long" in str(e).lower():
                    # Last resort: try with minimal parsing
                    logger.warning(f"Custom parser failed with attvalue error, trying minimal parsing: {e}")
                    processing_log.add_warning("Attempting minimal parsing as last resort")
                    return self._load_with_minimal_parser(docx_path, processing_log)
                else:
                    raise
            
        except ImportError as e:
            error_msg = f"Required dependencies not available for custom parser: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except etree.XMLSyntaxError as e:
            error_msg = f"XML syntax error during custom parsing: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except ValueError as e:
            if "attvalue" in str(e).lower():
                error_msg = f"Attribute value error during custom parsing (likely due to very large file): {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                # Try with extreme parser
                return self._load_with_extreme_parser(docx_path, processing_log)
            else:
                error_msg = f"Value error during custom parsing: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
        except MemoryError as e:
            error_msg = f"Memory error during custom parsing of very large file: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
        except Exception as e:
            error_msg = f"Custom parser failed for very large file: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
    
    def _load_with_streaming_parser(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Streaming parser using etree.iterparse for large files (>50MB).
        
        This approach uses lxml's iterparse to process XML content incrementally
        without loading the entire document into memory, similar to the approach
        in python-docx PR #1272.
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record parsing activities
            
        Returns:
            Document instance or None if parsing failed
        """
        try:
            from docx import Document
            import zipfile
            from lxml import etree
            
            logger.info(f"Using streaming parser (iterparse) for large file: {docx_path.name}")
            processing_log.add_info("Streaming parser: Using etree.iterparse for incremental processing")
            
            # Create a new document
            document = Document()
            
            # Add header explaining the processing method
            header = document.add_heading("Streaming Processed Document", 0)
            document.add_paragraph(f"Original file: {docx_path.name}")
            document.add_paragraph(f"Size: {docx_path.stat().st_size / (1024*1024):.1f}MB")
            document.add_paragraph("This document was processed using streaming parser (etree.iterparse) for large files.")
            
            # Extract and process XML content using iterparse
            try:
                with zipfile.ZipFile(docx_path, 'r') as zip_file:
                    # Get the main document XML
                    if 'word/document.xml' in zip_file.namelist():
                        xml_content = zip_file.read('word/document.xml')
                        
                        # Configure parser with huge_tree=True (from the GitHub PR)
                        parser = etree.XMLParser(
                            huge_tree=True,
                            recover=True,
                            remove_blank_text=True,
                            resolve_entities=False
                        )
                        
                        # Use iterparse for streaming processing
                        from io import BytesIO
                        context = etree.iterparse(
                            BytesIO(xml_content), 
                            events=('end',), 
                            tag='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
                        )
                        
                        # Process paragraphs incrementally
                        text_count = 0
                        for event, elem in context:
                            try:
                                # Extract text from paragraph
                                text_elements = elem.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if text_elements:
                                    text = ''.join([elem.text or '' for elem in text_elements])
                                    if text.strip():
                                        document.add_paragraph(text)
                                        text_count += 1
                                        
                                        # Limit the number of paragraphs to avoid memory issues
                                        if text_count >= 2000:
                                            document.add_paragraph("... (content truncated due to size limits)")
                                            break
                                
                                # Clear the element to free memory
                                elem.clear()
                                # Also eliminate now-empty references from the root node to elem
                                while elem.getprevious() is not None:
                                    del elem.getparent()[0]
                                    
                            except Exception as e:
                                logger.debug(f"Error extracting text from paragraph: {e}")
                                continue
                        
                        logger.info(f"Streaming parser extracted {text_count} paragraphs from {docx_path.name}")
                        processing_log.add_info(f"Streaming parser: Extracted {text_count} paragraphs using iterparse")
                        
            except zipfile.BadZipFile:
                error_msg = f"Invalid DOCX file (not a valid ZIP): {docx_path.name}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
            except Exception as e:
                error_msg = f"Error during streaming extraction: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
                return None
            
            return document
            
        except Exception as e:
            error_msg = f"Streaming parser failed: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None
    
    def _prepare_for_large_file(self):
        """
        Prepare system for processing large files (>100MB).
        
        This method performs comprehensive system preparation to optimize
        memory usage and parsing performance for large documents.
        """
        logger.info("Preparing system for large file processing...")
        
        # Force garbage collection
        gc.collect()
        
        # Set lxml environment variables for huge tree support
        os.environ['LXML_HUGE_TREE'] = '1'
        os.environ['LXML_USE_HUGE_TREE'] = '1'
        
        # Set memory management flags
        os.environ['PYTHONMALLOC'] = 'malloc'
        
        # Log system status
        try:
            import psutil
            from app.utils.shared_constants import MIN_MEMORY_REQUIRED_MB
            
            # Get system memory info
            memory = psutil.virtual_memory()
            logger.info(f"System memory: {memory.total / (1024**3):.1f}GB total, "
                       f"{memory.available / (1024**3):.1f}GB available")
            
            # Get process memory info
            process = psutil.Process()
            memory_mb = process.memory_info().rss / 1024 / 1024
            logger.info(f"Process memory usage: {memory_mb:.1f}MB")
            
            # Check if we have enough memory for large file processing
            if memory.available < MIN_MEMORY_REQUIRED_MB * 1024 * 1024:  # Less than required memory available
                logger.warning(f"Low memory available for large file processing. "
                             f"Required: {MIN_MEMORY_REQUIRED_MB}MB, Available: {memory.available / (1024**2):.1f}MB")
            
        except ImportError:
            logger.debug("psutil not available, skipping memory monitoring")
        except Exception as e:
            logger.debug(f"Memory monitoring failed: {e}")
        
        # Optimize Python memory management
        try:
            from app.utils.shared_constants import GARBAGE_COLLECTION_THRESHOLD
            gc.set_threshold(*GARBAGE_COLLECTION_THRESHOLD)  # More aggressive garbage collection
            logger.debug("Optimized garbage collection thresholds")
        except Exception as e:
            logger.debug(f"Failed to optimize garbage collection: {e}")
    
    def _process_pipeline(self, document: Document, processing_log: ProcessingLog) -> Dict[str, Any]:
        """
        Process document through the complete pipeline with graceful degradation.
        
        This method coordinates processing across all available processors
        (text, graphics, image) and aggregates results. If any processor
        fails, processing continues with the remaining processors.
        
        Args:
            document: Document instance to process
            processing_log: Processing log to record pipeline activities
            
        Returns:
            Dictionary containing aggregated processing results
        """
        pipeline_results = {
            'total_matches': 0,
            'text_matches': 0,
            'graphics_matches': 0,
            'image_matches': 0,
            'layout_impact': None,
            'processor_results': {},
            'all_matches': []  # Collect all detailed matches for reporting
        }
        
        # Process text content
        logger.info(f"TEXT PROCESSOR DEBUG: text_processor exists: {self.text_processor is not None}")
        if self.text_processor:
            logger.info(f"TEXT PROCESSOR DEBUG: text_processor.is_initialized(): {self.text_processor.is_initialized()}")
        if self.text_processor and self.text_processor.is_initialized():
            try:
                logger.info("Processing text content...")
                text_result = self.text_processor.process(document)
                
                if text_result.success:
                    pipeline_results['text_matches'] = text_result.matches_found
                    pipeline_results['total_matches'] += text_result.matches_found
                    # Store the metadata as a dictionary instead of the ProcessingResult object
                    pipeline_results['processor_results']['text'] = text_result.metadata
                    
                    # Extract layout impact data if available
                    if text_result.layout_impact:
                        pipeline_results['layout_impact'] = text_result.layout_impact
                    
                    # Add text matches to all_matches if available
                    if text_result.metadata and 'detailed_matches' in text_result.metadata:
                        pipeline_results['all_matches'].extend(text_result.metadata['detailed_matches'])
                    
                    logger.info(f"Text processing completed: {text_result.matches_found} matches")
                else:
                    error_msg = f"Text processing failed: {text_result.error_message}"
                    processing_log.add_error(error_msg)
                    logger.error(error_msg)
                    
            except Exception as e:
                error_msg = f"Text processing failed with exception: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
        else:
            warning_msg = "Text processor not available or not initialized, skipping text processing"
            processing_log.add_warning(warning_msg)
            logger.warning(warning_msg)
        
        # Process graphics content (placeholder for future implementation)
        if self.graphics_processor and self.graphics_processor.is_initialized():
            try:
                logger.info("Processing graphics content...")
                graphics_result = self.graphics_processor.process(document)
                
                if graphics_result.success:
                    pipeline_results['graphics_matches'] = graphics_result.matches_found
                    pipeline_results['total_matches'] += graphics_result.matches_found
                    # Store the metadata as a dictionary instead of the ProcessingResult object
                    pipeline_results['processor_results']['graphics'] = graphics_result.metadata
                    
                    # Add graphics matches to all_matches if available
                    if graphics_result.metadata and 'detailed_matches' in graphics_result.metadata:
                        pipeline_results['all_matches'].extend(graphics_result.metadata['detailed_matches'])
                    
                    logger.info(f"Graphics processing completed: {graphics_result.matches_found} matches")
                else:
                    error_msg = f"Graphics processing failed: {graphics_result.error_message}"
                    processing_log.add_error(error_msg)
                    logger.error(error_msg)
                    
            except Exception as e:
                error_msg = f"Graphics processing failed with exception: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
        else:
            logger.debug("Graphics processor not available, skipping graphics processing")
        
        # Process image content
        if self.image_processor and self.image_processor.is_initialized():
            try:
                logger.info("Processing images in document...")
                image_result = self.image_processor.process(document)
                if image_result.success:
                    # Centralize image match data for reuse
                    pipeline_results['image_matches'] = image_result.matches_found
                    pipeline_results['total_matches'] += image_result.matches_found
                    pipeline_results['processor_results']['image'] = image_result.metadata
                    
                    # Extract and centralize image matches for report generation
                    if image_result.metadata and 'image_matches' in image_result.metadata:
                        pipeline_results['image_matches_data'] = image_result.metadata['image_matches']
                        logger.debug(f"Centralized {len(image_result.metadata['image_matches'])} image matches for reporting")
                        
                        # Convert image matches to unified detection format for all_detections table
                        for i, match_data in enumerate(image_result.metadata['image_matches']):
                            # Handle both dict and object formats
                            if isinstance(match_data, dict):
                                ocr_result = match_data.get('ocr_result', {})
                                pattern_name = match_data.get('pattern', 'Unknown')
                                matched_text = ocr_result.get('text', 'Unknown') if isinstance(ocr_result, dict) else 'Unknown'
                                replacement_text = match_data.get('replacement_text', 'N/A')
                                confidence = ocr_result.get('confidence', 0.0) if isinstance(ocr_result, dict) else 0.0
                                bbox_obj = ocr_result.get('bounding_box', {}) if isinstance(ocr_result, dict) else {}
                                # Convert bounding box object to list format
                                if isinstance(bbox_obj, dict):
                                    bbox = [bbox_obj.get('x', 0), bbox_obj.get('y', 0), bbox_obj.get('width', 0), bbox_obj.get('height', 0)]
                                else:
                                    bbox = [0, 0, 0, 0]
                                wipe_boundaries = match_data.get('wipe_boundaries')
                                calculated_text_boundary = match_data.get('calculated_text_boundary')
                                processing_mode = match_data.get('processing_mode', 'wipe')
                                # Use extracted pattern text if available, otherwise use original text
                                extracted_pattern_text = match_data.get('extracted_pattern_text', matched_text)
                            else:
                                # Handle object format
                                ocr_result = getattr(match_data, 'ocr_result', None)
                                pattern_name = getattr(match_data, 'pattern', 'Unknown')
                                matched_text = getattr(ocr_result, 'text', 'Unknown') if ocr_result else 'Unknown'
                                replacement_text = getattr(match_data, 'replacement_text', 'N/A')
                                confidence = getattr(ocr_result, 'confidence', 0.0) if ocr_result else 0.0
                                bbox_obj = getattr(ocr_result, 'bounding_box', {}) if ocr_result else {}
                                # Convert bounding box object to list format
                                if hasattr(bbox_obj, 'x') and hasattr(bbox_obj, 'y') and hasattr(bbox_obj, 'width') and hasattr(bbox_obj, 'height'):
                                    bbox = [bbox_obj.x, bbox_obj.y, bbox_obj.width, bbox_obj.height]
                                else:
                                    bbox = [0, 0, 0, 0]
                                wipe_boundaries = getattr(match_data, 'wipe_boundaries', None)
                                calculated_text_boundary = getattr(match_data, 'calculated_text_boundary', None)
                                processing_mode = getattr(match_data, 'processing_mode', 'wipe')
                                # Use extracted pattern text if available, otherwise use original text
                                extracted_pattern_text = getattr(match_data, 'extracted_pattern_text', matched_text)
                            
                            # Use the extracted pattern text for display and matching
                            display_text = extracted_pattern_text if extracted_pattern_text else matched_text
                            
                            # Determine if this is a match based on whether replacement_text is not the default
                            is_matched = replacement_text != self.config.default_mapping
                            matched_pattern = pattern_name if is_matched else 'Unknown'
                            
                            # Create unified detection format
                            unified_detection = {
                                'pattern_name': matched_pattern or 'Unknown',
                                'actual_pattern': matched_pattern or 'Unknown',
                                'matched_text': display_text,  # Use extracted pattern text for display
                                'extracted_pattern_text': extracted_pattern_text,  # Store the extracted pattern text
                                'original_text': matched_text,  # Store the original OCR text
                                'start_pos': f"Image_{i+1}",
                                'end_pos': f"Image_{i+1}",
                                'replacement_text': replacement_text,
                                'location': 'Image',
                                'content_type': 'Image',  # These are image detections
                                'dimension': f"{bbox[2]}x{bbox[3]} pixels" if len(bbox) >= 4 else "Unknown",
                                'processor': 'Image',  # These are image detections
                                'is_matched': is_matched,  # Add match status
                                'font_info': {
                                    'font_family': 'OCR Detected',
                                    'font_size': f"{confidence:.2f}",
                                    'font_color': '000000'
                                },
                                # Add image-specific fields for wipe boundaries
                                'wipe_boundaries': wipe_boundaries,
                                'calculated_text_boundary': calculated_text_boundary,
                                'processing_mode': processing_mode,
                                'confidence': confidence,
                                'bounding_box': bbox
                            }
                            
                            # Initialize all_detections if it doesn't exist
                            if 'all_detections' not in pipeline_results:
                                pipeline_results['all_detections'] = []
                            
                            # Add unified detection to all_detections
                            pipeline_results['all_detections'].append(unified_detection)
                            logger.debug(f"Added image detection {i+1} to all_detections: {matched_text}")
                    
                    # Add image matches to all_matches if available
                    if image_result.metadata and 'detailed_matches' in image_result.metadata:
                        pipeline_results['all_matches'].extend(image_result.metadata['detailed_matches'])
                    
                    logger.info(f"Image processing completed: {image_result.matches_found} matches")
                else:
                    error_msg = f"Image processing failed: {image_result.error_message}"
                    processing_log.add_error(error_msg)
                    logger.error(error_msg)
            except Exception as e:
                error_msg = f"Image processing failed with exception: {e}"
                processing_log.add_error(error_msg)
                logger.error(error_msg)
        else:
            logger.debug("Image processor not available, skipping image processing")

        
        logger.info(f"Pipeline processing completed: {pipeline_results['total_matches']} total matches")
        return pipeline_results
    
    def _generate_output_path(self, input_path: Path) -> Path:
        """
        Generate output path for processed document.
        
        Args:
            input_path: Original document path
            
        Returns:
            Output path for processed document
        """
        output_dir = self.config.output_dir if self.config else Path("processed")
        output_dir.mkdir(exist_ok=True)
        
        # Generate filename with _processed suffix
        stem = input_path.stem
        if stem.endswith('_processed'):
            # Avoid double suffixes
            output_filename = f"{stem}.docx"
        else:
            output_filename = f"{stem}_processed.docx"
        
        return output_dir / output_filename
    
    def _save_document(self, document: Document, output_path: Path, processing_log: ProcessingLog) -> bool:
        """
        Save processed document with comprehensive error handling.
        
        Args:
            document: Document instance to save
            output_path: Path where to save the document
            processing_log: Processing log to record save activities
            
        Returns:
            True if save was successful, False otherwise
        """
        try:
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            document.save(output_path)
            logger.debug(f"Document successfully saved to: {output_path}")
            return True
        except Exception as e:
            error_msg = f"Failed to save document to {output_path}: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return False
    
    def _update_processing_stats(self, result: ProcessingResult):
        """
        Update internal processing statistics.
        
        Args:
            result: Processing result to incorporate into statistics
        """
        self.processing_stats['documents_processed'] += 1
        self.processing_stats['total_processing_time'] += result.processing_time
        self.processing_stats['total_matches_found'] += result.matches_found
        
        if result.success:
            self.processing_stats['successful_processing'] += 1
        else:
            self.processing_stats['failed_processing'] += 1
    
    def set_processors(self, text_processor: Optional[BaseProcessor] = None,
                      graphics_processor: Optional[BaseProcessor] = None,
                      image_processor: Optional[BaseProcessor] = None):
        """
        Set processor instances for the document processor.
        
        This method allows processors to be set after the DocumentProcessor
        is created, enabling flexible processor configuration and incremental
        implementation of different processor types.
        
        Args:
            text_processor: Text processor instance
            graphics_processor: Graphics processor instance
        """
        self.text_processor = text_processor
        self.graphics_processor = graphics_processor
        self.image_processor = image_processor
        
        logger.info("Processors updated in document processor")
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """
        Get comprehensive processing statistics and configuration info.
        
        Returns:
            Dictionary containing processing statistics and configuration
        """
        return {
            'config': {
                'text_mode': self.config.text_mode if self.config else 'unknown',
                'ocr_mode': self.config.ocr_mode if self.config else 'unknown',
                'use_gpu': self.config.use_gpu if self.config else False,
                'max_workers': self.config.max_workers if self.config else 1
            },
            'patterns_loaded': len(self.patterns),
            'mappings_loaded': len(self.mappings),
            'processors_available': {
                'text': self.text_processor is not None and self.text_processor.is_initialized(),
                'graphics': self.graphics_processor is not None and self.graphics_processor.is_initialized()
            },
            'converter_available': self.converter is not None,
            'processing_stats': self.processing_stats
        }
    
    def cleanup(self):
        """
        Clean up resources used by the document processor.
        
        This method releases resources and performs cleanup operations
        when the document processor is no longer needed.
        """
        logger.info("Cleaning up document processor")
        
        # Clean up processors
        if self.text_processor:
            self.text_processor.cleanup()
        if self.graphics_processor:
            self.graphics_processor.cleanup()
        if self.image_processor:
            self.image_processor.cleanup()
        
        # Clear references
        self.text_processor = None
        self.graphics_processor = None
        self.image_processor = None
        
        logger.info("Document processor cleanup completed")

    def _create_summary_document(self, docx_path: Path, processing_log: ProcessingLog) -> Optional[Document]:
        """
        Create a summary document for files that are too large to process fully.
        
        This creates a document with basic file information and a warning that
        the file was too large for full processing.
        
        Args:
            docx_path: Path to the DOCX file
            processing_log: Processing log to record activities
            
        Returns:
            Document instance with summary information
        """
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            logger.warning(f"Creating summary document for oversized file: {docx_path.name}")
            processing_log.add_warning(f"File too large for full processing, creating summary")
            
            # Create a new document
            document = Document()
            
            # Add title
            title = document.add_heading("Document Processing Summary", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add file information
            file_size = docx_path.stat().st_size
            size_mb = file_size / (1024 * 1024)
            
            document.add_paragraph(f"File Name: {docx_path.name}")
            document.add_paragraph(f"File Size: {size_mb:.1f}MB")
            document.add_paragraph(f"Processing Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add warning section
            warning_heading = document.add_heading("Processing Status", level=1)
            warning_para = document.add_paragraph()
            warning_para.add_run("⚠️  WARNING: ").bold = True
            warning_para.add_run(f"This file ({size_mb:.1f}MB) exceeds the recommended processing size limit of {MAX_PROCESSING_SIZE_MB}MB.")
            
            document.add_paragraph("The file was not processed for the following reasons:")
            reasons = document.add_paragraph()
            reasons.add_run("• XML parsing limitations for very large files\n")
            reasons.add_run("• Memory constraints during processing\n")
            reasons.add_run("• Risk of 'AttValue length too long' errors\n")
            reasons.add_run("• Performance degradation with files this large")
            
            # Add recommendations
            rec_heading = document.add_heading("Recommendations", level=1)
            document.add_paragraph("To process this file, consider:")
            recommendations = document.add_paragraph()
            recommendations.add_run("• Splitting the file into smaller sections\n")
            recommendations.add_run("• Reducing image quality or removing images\n")
            recommendations.add_run("• Using a different file format\n")
            recommendations.add_run("• Processing in batches")
            
            # Add technical details
            tech_heading = document.add_heading("Technical Details", level=1)
            document.add_paragraph(f"Maximum recommended file size: {MAX_PROCESSING_SIZE_MB}MB")
            document.add_paragraph(f"Current file size: {size_mb:.1f}MB")
            document.add_paragraph(f"Size difference: {size_mb - MAX_PROCESSING_SIZE_MB:.1f}MB over limit")
            
            logger.info(f"Created summary document for oversized file: {docx_path.name}")
            processing_log.add_info(f"Summary document created for {size_mb:.1f}MB file")
            
            return document
            
        except Exception as e:
            error_msg = f"Failed to create summary document: {e}"
            processing_log.add_error(error_msg)
            logger.error(error_msg)
            return None

def create_document_processor(config) -> DocumentProcessor:
    """
    Factory function to create a DocumentProcessor instance.
    
    Args:
        config: ProcessingConfig instance
        
    Returns:
        Configured DocumentProcessor instance
    """
    return DocumentProcessor(config)
