#!/usr/bin/env python3
"""
Test to mimic the exact full pipeline sequence and identify where text processor changes are lost.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from app.processors.text_processor import TextProcessor
from app.processors.graphics_processor import GraphicsProcessor
from app.processors.image_processor import ImageProcessor
from app.config import OCRConfig
import json
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s | %(levelname)8s | %(name)s | %(message)s')

def get_header_content(document, step_name):
    """Extract and display header content."""
    print(f"\n--- {step_name} ---")
    found_content = False
    for section_idx, section in enumerate(document.sections):
        if section.header:
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text or "4022-620-19087-13" in cell_text:
                            print(f"  Header content: '{cell_text}' in table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
                            found_content = True
    if not found_content:
        print("  No relevant header content found")
    return found_content

def test_pipeline_sequence():
    """Test the exact pipeline sequence to identify where changes are lost."""
    
    config = OCRConfig()
    
    # Load patterns and mappings
    with open(config.patterns_path, 'r') as f:
        patterns = json.load(f)
    with open(config.mapping_path, 'r') as f:
        mappings = json.load(f)
    
    print("=== TESTING FULL PIPELINE SEQUENCE ===")
    
    # Step 1: Load fresh document
    document = Document("source_documents/test_file2.docx")
    get_header_content(document, "STEP 1: Initial document")
    
    # Step 2: Text processor (same as full pipeline)
    print("\n=== STEP 2: Text Processor ===")
    text_processor = TextProcessor(patterns=patterns, mappings=mappings, mode=config.text_mode)
    text_matches = text_processor.process_document_text(document)
    print(f"Text processor found {len(text_matches)} matches")
    get_header_content(document, "AFTER Text Processor")
    
    # Step 3: Graphics processor (same as full pipeline)
    print("\n=== STEP 3: Graphics Processor ===")
    graphics_processor = GraphicsProcessor(patterns=patterns, mappings=mappings, mode=config.text_mode)
    graphics_matches = graphics_processor.process_graphics(document)
    print(f"Graphics processor found {len(graphics_matches)} matches")
    get_header_content(document, "AFTER Graphics Processor")
    
    # Step 4: Image processor (same as full pipeline)
    print("\n=== STEP 4: Image Processor ===")
    try:
        image_processor = ImageProcessor(
            patterns=patterns,
            mappings=mappings,
            mode=config.ocr_mode,
            ocr_engine=config.ocr_engine,
            gpu=config.gpu
        )
        
        # Get media directory (same as full pipeline)
        media_dir = None  # We'll handle this if needed
        
        image_matches = image_processor.process_images(document, media_dir)
        print(f"Image processor found {len(image_matches)} matches")
        get_header_content(document, "AFTER Image Processor")
        
    except Exception as e:
        print(f"Image processor error: {e}")
        print("Skipping image processor...")
        get_header_content(document, "AFTER Image Processor (skipped)")
    
    # Step 5: Save document (same as full pipeline)
    print("\n=== STEP 5: Save Document ===")
    output_path = "test_pipeline_sequence_result.docx"
    document.save(output_path)
    print(f"Saved document to: {output_path}")
    
    # Step 6: Verify saved document
    print("\n=== STEP 6: Verify Saved Document ===")
    saved_doc = Document(output_path)
    found_replacement = get_header_content(saved_doc, "FINAL Saved Document")
    
    print(f"\n=== SUMMARY ===")
    print(f"Text matches: {len(text_matches)}")
    print(f"Graphics matches: {len(graphics_matches) if 'graphics_matches' in locals() else 'N/A'}")
    print(f"Image matches: {len(image_matches) if 'image_matches' in locals() else 'N/A'}")
    
    if found_replacement:
        print("✅ SUCCESS: Replacement text found in final document!")
    else:
        print("❌ FAILURE: Replacement text lost somewhere in the pipeline!")
        print("Check the step-by-step output above to see where it was lost.")

if __name__ == "__main__":
    test_pipeline_sequence()
