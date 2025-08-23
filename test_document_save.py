#!/usr/bin/env python3
"""
Test to verify that text processor changes are actually saved to the document.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from app.processors.text_processor import TextProcessor
from app.config import OCRConfig
import json
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s | %(levelname)8s | %(name)s | %(message)s')

def test_document_save():
    """Test if text processor changes are properly saved."""
    
    config = OCRConfig()
    
    # Load patterns and mappings
    with open(config.patterns_path, 'r') as f:
        patterns = json.load(f)
    with open(config.mapping_path, 'r') as f:
        mappings = json.load(f)
    
    print("=== STEP 1: Load fresh document ===")
    document = Document("source_documents/test_file2.docx")
    
    # Check initial header content
    print("Initial header content:")
    for section_idx, section in enumerate(document.sections):
        if section.header:
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text:
                            print(f"  Found: '{cell_text}' in table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
    
    print("\n=== STEP 2: Process with text processor ===")
    text_processor = TextProcessor(patterns=patterns, mappings=mappings, mode=config.text_mode)
    matches = text_processor.process_document_text(document)
    
    print(f"Text processor found {len(matches)} matches")
    for match in matches:
        # Handle different possible attribute names
        original = getattr(match, 'original', getattr(match, 'original_text', 'N/A'))
        replacement = getattr(match, 'replacement', getattr(match, 'replacement_text', 'N/A'))
        location = getattr(match, 'location', 'N/A')
        
        if 'header' in str(location):
            print(f"  Header match: '{original}' -> '{replacement}' at {location}")
    
    print("\n=== STEP 3: Check document content after processing ===")
    print("Header content after text processing:")
    for section_idx, section in enumerate(document.sections):
        if section.header:
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text or "4022-620-19087-13" in cell_text:
                            print(f"  Content: '{cell_text}' in table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
    
    print("\n=== STEP 4: Save document ===")
    output_path = "test_save_verification.docx"
    document.save(output_path)
    print(f"Saved document to: {output_path}")
    
    print("\n=== STEP 5: Reload saved document and verify ===")
    reloaded_doc = Document(output_path)
    print("Header content in reloaded document:")
    found_replacement = False
    for section_idx, section in enumerate(reloaded_doc.sections):
        if section.header:
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text or "4022-620-19087-13" in cell_text:
                            print(f"  Content: '{cell_text}' in table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
                            if "4022-620-19087-13" in cell_text:
                                found_replacement = True
    
    print(f"\n=== RESULT ===")
    if found_replacement:
        print("✅ SUCCESS: Replacement text found in saved document!")
        print("The text processor is working correctly.")
    else:
        print("❌ FAILURE: Replacement text NOT found in saved document!")
        print("There's an issue with the text processor or document saving.")
    
    return found_replacement

if __name__ == "__main__":
    success = test_document_save()
    if not success:
        print("\n🔧 The issue is with text processor changes not being saved properly.")
    else:
        print("\n🤔 Text processor works fine in isolation - issue must be in full pipeline interaction.")
