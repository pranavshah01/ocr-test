#!/usr/bin/env python3
"""
Debug script to trace what happens to header text during full pipeline processing.
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from app.processors.text_processor import TextProcessor
from app.config import OCRConfig
import logging

# Set up logging to see debug messages
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s | %(levelname)8s | %(name)s | %(message)s')

def debug_header_processing():
    """Debug header processing in isolation."""
    
    config = OCRConfig()
    print(f"Config text_mode: {config.text_mode}")
    
    # Load patterns and mappings
    import json
    with open(config.patterns_path, 'r') as f:
        patterns = json.load(f)
    with open(config.mapping_path, 'r') as f:
        mappings = json.load(f)
    
    # Initialize text processor with same config as full pipeline
    text_processor = TextProcessor(
        patterns=patterns,
        mappings=mappings,
        mode=config.text_mode
    )
    
    # Load the document
    docx_path = "source_documents/test_file2.docx"
    document = Document(docx_path)
    
    print("\n=== BEFORE PROCESSING ===")
    # Check header content before processing
    for section_idx, section in enumerate(document.sections):
        if section.header:
            print(f"Header {section_idx} content:")
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text:
                            print(f"  Found target in table_{table_idx}_row_{row_idx}_cell_{cell_idx}: '{cell_text}'")
    
    print("\n=== PROCESSING ===")
    # Process the document
    matches = text_processor.process_document_text(document)
    print(f"Text processor returned {len(matches)} matches")
    
    for match in matches:
        if "77-620-1908713-03" in match.original:
            print(f"Header match found: '{match.original}' -> '{match.replacement}' at {match.location}")
    
    print("\n=== AFTER PROCESSING ===")
    # Check header content after processing
    for section_idx, section in enumerate(document.sections):
        if section.header:
            print(f"Header {section_idx} content after processing:")
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        if "77-620-1908713-03" in cell_text or "4022-620-19087-13" in cell_text:
                            print(f"  Content in table_{table_idx}_row_{row_idx}_cell_{cell_idx}: '{cell_text}'")
    
    # Save the document to see the actual result
    output_path = "debug_processed.docx"
    document.save(output_path)
    print(f"\nSaved debug document to: {output_path}")
    print("Check this file to see if the replacement was actually applied!")

if __name__ == "__main__":
    debug_header_processing()
