
import re
import logging
import time
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..core.models import ProcessingResult, MatchDetail, ProcessorType, MatchFlag, FallbackFlag
from ..utils.text_utils.text_docx_utils import TextReconstructor, FontManager, PatternMatcher, create_pattern_matcher, load_patterns_and_mappings
from config import DEFAULT_MAPPING, DEFAULT_SEPARATOR, PROCESSING_MODES

logger = logging.getLogger(__name__)


class TextReplacer:

    def __init__(self, mode: str = PROCESSING_MODES['APPEND'], separator: str = DEFAULT_SEPARATOR, default_mapping: str = DEFAULT_MAPPING):
        self.mode = mode
        self.separator = separator
        self.default_mapping = default_mapping

    def replace_text_in_runs(self, runs: List, original_text: str, replacement_text: str,
                           start_pos: int, end_pos: int) -> bool:
        try:

            current_pos = 0
            affected_runs = []

            for run in runs:
                run_start = current_pos
                run_end = current_pos + len(run.text)

                if run_start < end_pos and run_end > start_pos:
                    affected_runs.append((run, run_start, run_end))

                current_pos = run_end

            if not affected_runs:
                return False


            if self.mode == PROCESSING_MODES['APPEND']:


                expected_append = f"{original_text}{self.separator}{replacement_text}"


                current_text = ""
                for run in runs:
                    current_text += run.text


                if expected_append in current_text:
                    logger.info(f"APPEND MODE: Skipping duplicate append for '{original_text}' -> '{replacement_text}' (already exists)")
                    return True

                final_text = expected_append
            else:
                final_text = replacement_text


            first_run = affected_runs[0][0]
            font_info = FontManager.get_font_info(first_run)


            for i, (run, run_start, run_end) in enumerate(affected_runs):
                if i == 0:

                    run_match_start = max(0, start_pos - run_start)
                    run_match_end = min(len(run.text), end_pos - run_start)


                    before_text = run.text[:run_match_start]
                    after_text = run.text[run_match_end:]
                    run.text = before_text + final_text + after_text


                    FontManager.apply_font_info(run, font_info)
                else:

                    run_match_start = max(0, start_pos - run_start)
                    run_match_end = min(len(run.text), end_pos - run_start)

                    before_text = run.text[:run_match_start]
                    after_text = run.text[run_match_end:]
                    run.text = before_text + after_text

            return True

        except Exception as e:
            logger.error(f"Failed to replace text: {e}")
            return False


class TextProcessor:

    def __init__(self, config):
        self.config = config
        self.initialized = False
        self.patterns = {}
        self.mappings = {}
        self.pattern_matcher = None
        self.text_replacer = None
        self.mode = getattr(config, 'mode', PROCESSING_MODES['APPEND'])
        self.separator = getattr(config, 'separator', DEFAULT_SEPARATOR)
        self.default_mapping = getattr(config, 'default_mapping', DEFAULT_MAPPING)


        self._load_patterns_and_mappings()

    def _load_patterns_and_mappings(self):
        self.patterns, self.mappings = load_patterns_and_mappings(self.config)

    def initialize(self) -> bool:
        try:

            self.pattern_matcher = create_pattern_matcher(self.patterns, self.mappings)


            self.text_replacer = TextReplacer(self.mode, self.separator, self.default_mapping)

            logger.info(f"Pattern matcher created with {len(self.patterns)} patterns, {len(self.mappings)} mappings")

            self.initialized = True
            logger.info(f"Text processor initialized with {len(self.patterns)} patterns, {len(self.mappings)} mappings")
            return True
        except Exception as e:
            logger.error(f"Error initializing Text Processor: {e}")
            self.initialized = False
            return False

    def is_initialized(self) -> bool:
        return self.initialized

    def process_document(self, document: Document, processing_result: ProcessingResult) -> ProcessingResult:
        if not self.initialized:
            logger.error("Text processor not initialized")
            processing_result.error_message = "Text processor not initialized"
            return processing_result

        start_time = time.time()
        logger.info("Starting text processing...")

        try:

            self.document = document


            all_detections = []


            body_detections = self._process_paragraphs(document.paragraphs, "body")
            all_detections.extend(body_detections)


            table_detections = self._process_tables(document.tables)
            all_detections.extend(table_detections)


            header_footer_detections = self._process_headers_footers(document)
            all_detections.extend(header_footer_detections)


            logger.info(f"Found {len(all_detections)} pattern matches:")
            for i, detection in enumerate(all_detections, 1):
                pattern_name = detection.get('pattern_name', 'Unknown')
                matched_text = detection.get('matched_text', 'Unknown')
                replacement_text = detection.get('replacement_text', 'Unknown')
                location = detection.get('location', 'Unknown')
                logger.info(f"  {i}. Pattern '{pattern_name}' matched '{matched_text}' -> '{replacement_text}' at {location}")


            self._update_processing_result(processing_result, all_detections)


            reconstruction_results = self.reconstruct_document(document, all_detections)


            self._update_reconstruction_status(processing_result, all_detections, reconstruction_results)

            processing_time = time.time() - start_time
            logger.info(f"Text processing completed in {processing_time:.2f}s: {len(all_detections)} detections")

            return processing_result

        except Exception as e:
            error_msg = f"Error during text processing: {e}"
            logger.error(error_msg)
            processing_result.error_message = error_msg
            return processing_result

    def reconstruct_document(self, document: Document, detections: List[Dict[str, Any]]) -> Dict[str, bool]:
        logger.info("Starting document reconstruction...")

        reconstruction_results = {}

        try:

            body_results = self._reconstruct_paragraphs(document.paragraphs, detections, "body")
            reconstruction_results.update(body_results)


            table_results = self._reconstruct_tables(document.tables, detections)
            reconstruction_results.update(table_results)


            header_footer_results = self._reconstruct_headers_footers(document, detections)
            reconstruction_results.update(header_footer_results)

            logger.info(f"Document reconstruction completed: {sum(reconstruction_results.values())}/{len(reconstruction_results)} matches reconstructed successfully")

        except Exception as e:
            logger.error(f"Error during document reconstruction: {e}")

        return reconstruction_results

    def _reconstruct_paragraphs(self, paragraphs: List[Paragraph], detections: List[Dict[str, Any]], location: str) -> Dict[str, bool]:
        reconstruction_results = {}

        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_location = f"{location}_paragraph_{i}"
                paragraph_results = self._reconstruct_paragraph(paragraph, detections, paragraph_location)
                reconstruction_results.update(paragraph_results)
            except Exception as e:
                logger.error(f"Error reconstructing paragraph {i} in {location}: {e}")

        return reconstruction_results

    def _reconstruct_paragraph(self, paragraph: Paragraph, detections: List[Dict[str, Any]], location: str) -> Dict[str, bool]:
        reconstruction_results = {}

        if not paragraph.runs:
            return reconstruction_results


        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)

        if not full_text.strip():
            return reconstruction_results


        paragraph_detections = [
            detection for detection in detections
            if detection.get('location') == location and detection.get('is_matched', False)
        ]


        paragraph_detections.sort(key=lambda x: x.get('start_pos', 0), reverse=True)


        for detection in paragraph_detections:
            try:
                matched_text = detection.get('matched_text', '')
                replacement_text = detection.get('replacement_text', '')
                start_pos = detection.get('start_pos', 0)
                end_pos = detection.get('end_pos', 0)

                if not matched_text or not replacement_text:
                    continue


                detection_key = f"{location}_{matched_text}_{start_pos}_{end_pos}"


                text_span = TextReconstructor.find_text_in_runs(runs, matched_text, start_pos)
                if not text_span:
                    reconstruction_results[detection_key] = False
                    continue

                span_start, span_end, affected_runs = text_span


                success = self.text_replacer.replace_text_in_runs(
                    runs, matched_text, replacement_text, span_start, span_end
                )

                reconstruction_results[detection_key] = success

                if success:
                    logger.info(f"Text replacement successful: '{matched_text}' -> '{replacement_text}' at {location}")
                else:
                    logger.warning(f"Text replacement failed: '{matched_text}' -> '{replacement_text}' at {location}")

            except Exception as e:
                logger.error(f"Error applying replacement for detection at {location}: {e}")
                detection_key = f"{location}_{matched_text}_{start_pos}_{end_pos}"
                reconstruction_results[detection_key] = False

        return reconstruction_results

    def _reconstruct_tables(self, tables: List[Table], detections: List[Dict[str, Any]], location_prefix: str = "") -> Dict[str, bool]:
        reconstruction_results = {}


        from ..utils.table_utils import get_table_cells_to_process

        for table_idx, table in enumerate(tables):
            try:

                cells_to_process = get_table_cells_to_process(table)

                for row_idx, cell_idx, cell in cells_to_process:
                    if location_prefix:
                        location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    else:
                        location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    cell_results = self._reconstruct_paragraphs(cell.paragraphs, detections, location)
                    reconstruction_results.update(cell_results)

            except Exception as e:
                logger.error(f"Error reconstructing table {table_idx}: {e}")

        return reconstruction_results

    def _reconstruct_headers_footers(self, document: Document, detections: List[Dict[str, Any]]) -> Dict[str, bool]:
        reconstruction_results = {}

        try:

            for section_idx, section in enumerate(document.sections):

                if section.header:
                    header_results = self._reconstruct_paragraphs(section.header.paragraphs, detections, f"header_section_{section_idx}")
                    reconstruction_results.update(header_results)
                    header_table_results = self._reconstruct_tables(section.header.tables, detections, f"header_section_{section_idx}")
                    reconstruction_results.update(header_table_results)


                if section.footer:
                    footer_results = self._reconstruct_paragraphs(section.footer.paragraphs, detections, f"footer_section_{section_idx}")
                    reconstruction_results.update(footer_results)
                    footer_table_results = self._reconstruct_tables(section.footer.tables, detections, f"footer_section_{section_idx}")
                    reconstruction_results.update(footer_table_results)

        except Exception as e:
            logger.error(f"Error reconstructing headers/footers: {e}")

        return reconstruction_results

    def _process_paragraphs(self, paragraphs: List[Paragraph], location: str) -> List[Dict[str, Any]]:
        all_detections = []

        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_detections = self._process_paragraph(
                    paragraph, f"{location}_paragraph_{i}"
                )
                all_detections.extend(paragraph_detections)
            except Exception as e:
                logger.error(f"Error processing paragraph {i} in {location}: {e}")

        return all_detections

    def _process_paragraph(self, paragraph: Paragraph, location: str) -> List[Dict[str, Any]]:
        if not paragraph.runs:
            return []


        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)

        if not full_text.strip():
            return []


        all_pattern_matches = self.pattern_matcher.find_all_pattern_matches(full_text)

        detection_results = []

        for pattern_name, matched_text, start_pos, end_pos in all_pattern_matches:

            replacement_text = self.pattern_matcher.get_replacement(matched_text)


            if not replacement_text:
                if self.mode == "replace":

                    logger.debug(f"REPLACE MODE: Skipping '{matched_text}' - no mapping found")
                    replacement_text = None
                elif self.mode == PROCESSING_MODES['APPEND']:

                    replacement_text = self.default_mapping
                    logger.info(f"APPEND MODE: Using default mapping '{self.default_mapping}' for '{matched_text}'")
                else:

                    logger.warning(f"Unknown mode '{self.mode}', skipping '{matched_text}'")
                    replacement_text = None


            is_matched = replacement_text is not None


            actual_pattern = self.patterns.get(pattern_name, pattern_name)


            font_info = FontManager.extract_font_info_for_detection(
                runs, matched_text, start_pos, getattr(self, 'document', None)
            )


            content_type = "Paragraph"
            dimension = ""

            if "table" in location.lower():
                content_type = "Table"

                dimension = "Cell size"
            elif "header" in location.lower():
                content_type = "Header"
            elif "footer" in location.lower():
                content_type = "Footer"


            detection_result = {
                'pattern_name': pattern_name,
                'actual_pattern': actual_pattern,
                'matched_text': matched_text,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'replacement_text': replacement_text,
                'location': location,
                'content_type': content_type,
                'dimension': dimension,
                'processor': 'Text',
                'font_info': font_info,
                'is_matched': is_matched,
                'confidence': 1.0
            }

            detection_results.append(detection_result)

        return detection_results

    def _process_tables(self, tables: List[Table], location_prefix: str = "") -> List[Dict[str, Any]]:
        all_detections = []


        from ..utils.table_utils import get_table_cells_to_process

        for table_idx, table in enumerate(tables):
            try:

                cells_to_process = get_table_cells_to_process(table)

                for row_idx, cell_idx, cell in cells_to_process:
                    if location_prefix:
                        location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    else:
                        location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"

                    cell_detections = self._process_paragraphs(cell.paragraphs, location)
                    all_detections.extend(cell_detections)

                    logger.debug(f"Processed cell at {location} with {len(cell_detections)} detections")

            except Exception as e:
                logger.error(f"Error processing table {table_idx}: {e}")

        return all_detections

    def _process_headers_footers(self, document: Document) -> List[Dict[str, Any]]:
        all_detections = []

        try:

            for section_idx, section in enumerate(document.sections):

                if section.header:
                    header_detections = self._process_paragraphs(section.header.paragraphs, f"header_section_{section_idx}")
                    all_detections.extend(header_detections)


                    header_table_detections = self._process_tables(section.header.tables, f"header_section_{section_idx}")
                    all_detections.extend(header_table_detections)


                if section.footer:
                    footer_detections = self._process_paragraphs(section.footer.paragraphs, f"footer_section_{section_idx}")
                    all_detections.extend(footer_detections)


                    footer_table_detections = self._process_tables(section.footer.tables, f"footer_section_{section_idx}")
                    all_detections.extend(footer_table_detections)
        except Exception as e:
            logger.error(f"Error processing headers/footers: {e}")

        return all_detections

    def _update_processing_result(self, processing_result: ProcessingResult, detections: List[Dict[str, Any]]):

        text_matches = 0
        graphics_matches = 0
        image_matches = 0
        graphics_no_match = 0
        image_no_match = 0


        match_details = []
        # Continue sr_no from existing match_details to keep numbering continuous across processors
        sr_no = len(processing_result.match_details) + 1

        for detection in detections:

            if detection.get('processor') == 'Text':
                if detection.get('is_matched', False):
                    text_matches += 1
            elif detection.get('processor') == 'Graphics':
                if detection.get('is_matched', False):
                    graphics_matches += 1
                else:
                    graphics_no_match += 1
            elif detection.get('processor') == 'Image':
                if detection.get('is_matched', False):
                    image_matches += 1
                else:
                    image_no_match += 1


            match_detail_dict = self._convert_detection_to_match_detail(
                detection, sr_no, ProcessorType.TEXT.value
            )


            match_detail = MatchDetail(
                sr_no=match_detail_dict['sr_no'],
                type=ProcessorType.TEXT,
                orig_id_name=match_detail_dict['orig_id_name'],
                src_text=match_detail_dict['src_text'],
                src_text_font=match_detail_dict['src_text_font'],
                src_text_color=match_detail_dict['src_text_color'],
                src_text_size=match_detail_dict['src_text_size'],
                src_dimension=match_detail_dict['src_dimension'],
                mapped_text=match_detail_dict['mapped_text'],
                mapped_text_font=match_detail_dict['mapped_text_font'],
                mapped_text_color=match_detail_dict['mapped_text_color'],
                mapped_text_size=match_detail_dict['mapped_text_size'],
                match_flag=MatchFlag.YES if match_detail_dict['match_flag'] == 'Y' else MatchFlag.NO,
                is_fallback=FallbackFlag.NO,
                reasoning=None
            )

            match_details.append(match_detail)
            sr_no += 1


        processing_result.total_text_matches = text_matches
        processing_result.total_graphics_matches = graphics_matches
        processing_result.total_image_matches = image_matches
        processing_result.total_graphics_no_match = graphics_no_match
        processing_result.total_image_no_match = image_no_match
        processing_result.matches_found = text_matches + graphics_matches + image_matches
        processing_result.total_matches = len(detections)
        processing_result.match_details = match_details


        processing_result.processor_type = "text_processor"


        processing_result.success = True


        if hasattr(self, 'document') and self.document:

            base_name = Path(processing_result.file_name).stem
            processing_result.processed_file_name = f"{base_name}{getattr(self.config, 'suffix', '_12NC_processed')}.docx"

    def _convert_detection_to_match_detail(
        self,
        detection: Dict[str, Any],
        sr_no: int,
        processor_type: str
    ) -> Dict[str, Any]:

        font_info = detection.get('font_info', {})
        font_family = font_info.get('font_family', 'Arial')
        font_size = font_info.get('font_size', '12.0')
        font_color = font_info.get('color', '000000')

        return {
            'sr_no': sr_no,
            'type': processor_type,
            'orig_id_name': detection.get('location', ''),
            'src_text': detection.get('matched_text', ''),
            'src_text_font': font_family,
            'src_text_color': font_color,
            'src_text_size': font_size,
            'src_dimension': detection.get('dimension', ''),
            'mapped_text': detection.get('replacement_text', ''),
            'mapped_text_font': font_family,
            'mapped_text_color': font_color,
            'mapped_text_size': font_size,
            'match_flag': 'Y' if detection.get('is_matched', False) else 'N',
            'is_fallback': 'N',
            'reasoning': None,
            'reconstructed': False
        }

    def _update_reconstruction_status(self, processing_result: ProcessingResult, detections: List[Dict[str, Any]], reconstruction_results: Dict[str, bool]):
        for match_detail in processing_result.match_details:

            for detection in detections:
                if (detection.get('matched_text') == match_detail.src_text and
                    detection.get('location') == match_detail.orig_id_name):


                    detection_key = f"{detection.get('location')}_{detection.get('matched_text')}_{detection.get('start_pos', 0)}_{detection.get('end_pos', 0)}"


                    match_detail.reconstructed = reconstruction_results.get(detection_key, False)
                    break

    def cleanup(self):
        logger.info("Cleaning up text processor")
        self.initialized = False


def create_text_processor(config) -> TextProcessor:
    return TextProcessor(config)