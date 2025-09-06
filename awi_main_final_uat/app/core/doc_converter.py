"""
Enhanced Document Converter for Document Processing Pipeline.

This module provides comprehensive document conversion capabilities including:
- LibreOffice conversion (cross-platform)
- Microsoft Word COM conversion (Windows)
- Conversion validation and error handling
- Fallback mechanisms for conversion failures
"""

import subprocess
import shutil
import tempfile
import platform
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class ConversionError(Exception):
    """Exception raised when document conversion fails."""
    pass

class DocConverter:
    """
    Enhanced document converter with multiple conversion methods.
    
    This class provides comprehensive document conversion capabilities
    with multiple fallback mechanisms and robust error handling.
    """
    
    def __init__(self):
        """Initialize document converter with available conversion tools."""
        self.available_tools = self._detect_conversion_tools()
        self.preferred_tool = self._select_preferred_tool()
        
        if not self.available_tools:
            logger.info("No .doc file conversion tools detected. .docx files will work normally.")
            logger.debug("To enable .doc file support:")
            logger.debug("  - Install LibreOffice for .doc file support")
            logger.debug("  - On macOS: Download from https://www.libreoffice.org")
            logger.debug("  - On Linux: install via your package manager, e.g. 'sudo apt install libreoffice' (Debian/Ubuntu), 'sudo dnf install libreoffice' (Fedora), or 'sudo yum install libreoffice' (RHEL/CentOS)")
            logger.debug("  - On Windows: Word COM interface can also be used if Microsoft Word is installed")
        else:
            logger.info(f"Available conversion tools: {self.available_tools}")
            logger.info(f"Preferred conversion tool: {self.preferred_tool}")
    
    def _detect_conversion_tools(self) -> List[str]:
        """Detect available document conversion tools."""
        tools = []
        
        # Check for LibreOffice
        if self._check_libreoffice():
            tools.append('libreoffice')
        
        # Check for Microsoft Word COM (Windows only)
        if platform.system() == 'Windows' and self._check_word_com():
            tools.append('word_com')
        
        
        return tools
    
    def _check_libreoffice(self) -> bool:
        """Check if LibreOffice is available."""
        try:
            # Try different LibreOffice command names
            libreoffice_commands = ['libreoffice', 'soffice', 'loffice', 'libreoffice6.4', 'libreoffice7.0']

            # macOS bundled soffice path can exist even if not on PATH
            mac_soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            if platform.system() == "Darwin" and Path(mac_soffice).exists():
                libreoffice_commands.insert(0, mac_soffice)
            
            for cmd in libreoffice_commands:
                result = subprocess.run([cmd, '--version'], 
                                      capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    logger.info(f"LibreOffice detected: {cmd}")
                    self.libreoffice_cmd = cmd
                    return True
            
            return False
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
            return False
    
    def _check_word_com(self) -> bool:
        """Check if Microsoft Word COM interface is available (Windows only)."""
        try:
            import win32com.client
            # Try to create Word application object
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Quit()
            logger.info("Microsoft Word COM interface detected")
            return True
        except ImportError:
            logger.debug("pywin32 not available for Word COM interface")
            return False
        except Exception as e:
            logger.debug(f"Word COM interface not available: {e}")
            return False
    
    def _select_preferred_tool(self) -> Optional[str]:
        """Select the preferred conversion tool based on platform preferences and availability."""
        if not self.available_tools:
            return None

        system = platform.system()
        
        # Platform-specific preferences
        if system == 'Windows':
            # Prefer Word COM on Windows, then LibreOffice
            if 'word_com' in self.available_tools:
                return 'word_com'
            if 'libreoffice' in self.available_tools:
                return 'libreoffice'
        else:
            # Prefer LibreOffice on macOS/Linux; Word COM not applicable
            if 'libreoffice' in self.available_tools:
                return 'libreoffice'
            if 'word_com' in self.available_tools:
                return 'word_com'
        
        # Default to first available tool as a last resort
        return self.available_tools[0]
    
    def convert_to_docx(self, file_path: Path) -> Optional[Path]:
        """
        Convert document to DOCX format using available conversion tools.
        
        Args:
            file_path: Path to the document file to convert
            
        Returns:
            Path to the converted DOCX file or None if conversion failed
        """
        if not self.available_tools:
            raise ConversionError("No conversion tools available")
        
        # Create output directory for converted files
        converted_dir = Path("converted")
        converted_dir.mkdir(exist_ok=True)
        
        # Generate output path
        output_path = converted_dir / f"{file_path.stem}_converted.docx"
        
        logger.info(f"Converting {file_path} to {output_path}")
        
        # Try conversion with preferred tool first
        if self.preferred_tool:
            try:
                if self.preferred_tool == 'libreoffice':
                    return self._convert_with_libreoffice(file_path, output_path)
                elif self.preferred_tool == 'word_com':
                    return self._convert_with_word_com(file_path, output_path)
            except Exception as e:
                logger.warning(f"Preferred tool {self.preferred_tool} failed: {e}")
        
        # Try all available tools as fallback
        for tool in self.available_tools:
            if tool == self.preferred_tool:
                continue  # Already tried
            
            try:
                if tool == 'libreoffice':
                    result = self._convert_with_libreoffice(file_path, output_path)
                elif tool == 'word_com':
                    result = self._convert_with_word_com(file_path, output_path)
                
                if result:
                    logger.info(f"Conversion succeeded with {tool}")
                    return result
                    
            except Exception as e:
                logger.warning(f"Tool {tool} failed: {e}")
                continue
        
        raise ConversionError(f"All conversion tools failed for {file_path}")
    
    def _convert_with_libreoffice(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert document using LibreOffice."""
        try:
            # Create temporary directory for conversion
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_output = Path(temp_dir) / f"{input_path.stem}.docx"
                
                # LibreOffice command for conversion
                cmd = [
                    self.libreoffice_cmd,
                    '--headless',  # Run in headless mode
                    '--convert-to', 'docx',  # Convert to DOCX
                    '--outdir', temp_dir,  # Output directory
                    str(input_path)  # Input file
                ]
                
                logger.debug(f"Running LibreOffice command: {' '.join(cmd)}")
                
                # Execute conversion
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                
                if result.returncode != 0:
                    logger.error(f"LibreOffice conversion failed: {result.stderr}")
                    raise ConversionError(f"LibreOffice conversion failed: {result.stderr}")
                
                # Check if output file was created
                if not temp_output.exists():
                    logger.error("LibreOffice conversion completed but output file not found")
                    raise ConversionError("Output file not created")
                
                # Move file to final location
                shutil.move(str(temp_output), str(output_path))
                
                # Validate converted file
                if self._validate_docx_file(output_path):
                    logger.info(f"LibreOffice conversion successful: {output_path}")
                    return output_path
                else:
                    raise ConversionError("Converted file validation failed")
                
        except subprocess.TimeoutExpired:
            raise ConversionError("LibreOffice conversion timed out")
        except Exception as e:
            raise ConversionError(f"LibreOffice conversion error: {e}") from e
    
    def _convert_with_word_com(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert document using Microsoft Word COM interface (Windows only)."""
        try:
            import win32com.client
            
            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            try:
                # Open the document
                doc = word_app.Documents.Open(str(input_path.absolute()))
                
                # Save as DOCX
                doc.SaveAs2(str(output_path.absolute()), FileFormat=16)  # 16 = wdFormatDocumentDefault
                
                # Close document and quit Word
                doc.Close()
                word_app.Quit()
                
                # Validate converted file
                if self._validate_docx_file(output_path):
                    logger.info(f"Word COM conversion successful: {output_path}")
                    return output_path
                else:
                    raise ConversionError("Converted file validation failed")
                
            except Exception as e:
                # Ensure Word is closed even if conversion fails
                try:
                    word_app.Quit()
                except Exception as quit_err:
                    logger.debug(f"Failed to quit Word after exception: {quit_err}")
                raise
                
        except ImportError:
            raise ConversionError("pywin32 not available for Word COM interface")
        except Exception as e:
            raise ConversionError(f"Word COM conversion error: {e}") from e
    
    
    def _validate_docx_file(self, file_path: Path) -> bool:
        """Validate that a file is a valid DOCX file."""
        try:
            # Check if file exists and has content
            if not file_path.exists() or file_path.stat().st_size == 0:
                return False
            
            # Try to open with python-docx to validate
            from docx import Document
            doc = Document(file_path)
            
            # Basic validation - check if document has content
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                logger.warning(f"Converted file appears to be empty: {file_path}")
                return False
            
            return True
            
        except Exception as e:
            logger.warning(f"DOCX validation failed for {file_path}: {e}")
            return False
    
    def get_conversion_status(self) -> Dict[str, Any]:
        """Get status of conversion tools and capabilities."""
        return {
            'available_tools': self.available_tools,
            'preferred_tool': self.preferred_tool,
            'platform': platform.system(),
            'conversion_supported': len(self.available_tools) > 0
        }
    
    def cleanup_converted_files(self):
        """Clean up converted files directory."""
        try:
            converted_dir = Path("converted")
            if converted_dir.exists():
                for file_path in converted_dir.glob("*_converted.docx"):
                    file_path.unlink()
                logger.info("Cleaned up converted files")
        except Exception as e:
            logger.warning(f"Failed to cleanup converted files: {e}")