
from __future__ import annotations

import logging
import os
import shutil
import subprocess
import tempfile
import zipfile
import platform
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any

from app.core.doc_converter import DocConverter, ConversionError

from config import (
    MAX_FILE_SIZE_MB,
    LARGE_FILE_THRESHOLD_MB,
    VERY_LARGE_FILE_THRESHOLD_MB
)

logger = logging.getLogger(__name__)


class DocumentProcessor:

    def __init__(self, config) -> None:
        self.config = config
        self.initialized = False
        self.text_processor = None
        self.graphics_processor = None
        self.image_processor = None
        self.parked_attributes: List[Dict] = [] # Initialize parked_attributes

    def initialize(self) -> None:

        # Initialize text processor
        try:
            from app.processors.text_processor import create_text_processor
            self.text_processor = create_text_processor(self.config)
            self.text_processor.initialize()
            logger.info("Text processor initialized in document processor")
        except Exception as e:
            logger.error(f"Failed to initialize text processor: {e}")
            self.text_processor = None

        # Load patterns and mappings once for both graphics and image processors
        patterns = {}
        mappings = {}
        try:
            from app.utils.text_utils.text_docx_utils import load_patterns_and_mappings
            patterns, mappings = load_patterns_and_mappings(self.config)
            logger.info(f"Loaded patterns and mappings: {len(patterns)} patterns, {len(mappings)} mappings")
        except Exception as e:
            logger.error(f"Failed to load patterns and mappings: {e}. Proceeding with empty configurations.")

        # Initialize graphics processor
        try:
            from app.processors.graphics_processor import GraphicsProcessor
            self.graphics_processor = GraphicsProcessor(
                patterns=patterns,
                mappings=mappings,
                mode=self.config.text_mode,
                separator=self.config.text_separator,
                default_mapping=self.config.default_mapping
            )
            self.graphics_processor.initialize()
            logger.info("Graphics processor initialized in document processor")
        except Exception as e:
            logger.error(f"Failed to initialize graphics processor: {e}")
            self.graphics_processor = None

        # Initialize image processor
        try:
            from app.processors.image_processor import ImageProcessor
            self.image_processor = ImageProcessor(
                patterns=patterns,
                mappings=mappings,
                mode=getattr(self.config, 'ocr_mode', 'append'),
                separator=self.config.text_separator,
                default_mapping=self.config.default_mapping,
                ocr_engine=getattr(self.config, 'ocr_engine', 'hybrid'),
                confidence_min=getattr(self.config, 'confidence_min', 0.4),
                use_gpu=getattr(self.config, 'use_gpu', True)
            )
            self.image_processor.initialize()
            logger.info("Image processor initialized in document processor")
        except Exception as e:
            logger.error(f"Failed to initialize image processor: {e}")
            self.image_processor = None

        self.initialized = True

    def cleanup(self) -> None:
        if self.text_processor:
            try:
                self.text_processor.cleanup()
                logger.info("Text processor cleaned up")
            except Exception as e:
                logger.error(f"Error cleaning up text processor: {e}")
        self.text_processor = None

        if self.graphics_processor:
            try:
                self.graphics_processor.cleanup()
                logger.info("Graphics processor cleaned up")
            except Exception as e:
                logger.error(f"Error cleaning up graphics processor: {e}")
        self.graphics_processor = None

        if self.image_processor:
            try:
                self.image_processor.cleanup()
                logger.info("Image processor cleaned up")
            except Exception as e:
                logger.error(f"Error cleaning up image processor: {e}")
        self.image_processor = None

        self.initialized = False

    def detect_and_set_parser(self, file_path: Path) -> str:
        try:
            file_size = file_path.stat().st_size
            size_mb = file_size / (1024 * 1024)

            logger.info(f"Detecting parser for document: {size_mb:.1f}MB - {file_path.name}")


            if size_mb > getattr(self.config, 'max_file_size', MAX_FILE_SIZE_MB):
                error_msg = f"File size {size_mb:.1f}MB exceeds maximum supported size of {MAX_FILE_SIZE_MB}MB"
                logger.error(error_msg)
                return "unsupported_size"


            large_threshold = getattr(self.config, 'large_file_threshold', LARGE_FILE_THRESHOLD_MB)
            if size_mb < large_threshold:
                logger.info(f"Selected standard parser for {size_mb:.1f}MB file ({file_path.name})")
                return "standard"


            very_large_threshold = getattr(self.config, 'very_large_file_threshold', VERY_LARGE_FILE_THRESHOLD_MB)
            if size_mb < very_large_threshold:
                logger.info(f"Selected enhanced parser for {size_mb:.1f}MB file ({file_path.name})")
                return "enhanced"


            if size_mb >= very_large_threshold:
                logger.info(f"Selected custom parser for {size_mb:.1f}MB file ({file_path.name})")
                return "custom"


            logger.info(f"Using standard parser as fallback for {size_mb:.1f}MB file ({file_path.name})")
            return "standard"

        except Exception as e:
            logger.error(f"Error detecting parser for {file_path}: {e}")
            return "error"

    @staticmethod
    def is_attvalue_error(error_message: str) -> bool:
        """Centralized AttValue error detection for reuse across modules."""
        error_lower = error_message.lower()
        attvalue_indicators = [
            "attvalue too large",
            "attribute value too large", 
            "xml attribute too large",
            "lxml attribute value limit",
            "huge_tree",
            "xml parsing error",
            "attvalue",
            "attribute value"
        ]
        return any(indicator in error_lower for indicator in attvalue_indicators)

    def get_next_parser(self, current_parser: str) -> Optional[str]:
        parser_sequence = ["standard", "enhanced", "custom"]

        try:
            current_index = parser_sequence.index(current_parser)
            if current_index + 1 < len(parser_sequence):
                next_parser = parser_sequence[current_index + 1]
                logger.info(f"Falling back from {current_parser} to {next_parser} parser")
                return next_parser
        except ValueError:
            logger.warning(f"Unknown parser type: {current_parser}")

        return None

    def try_parser_with_fallback(self, file_path: Path, initial_parser: str = None) -> Tuple[str, Optional[str]]:
        if initial_parser is None:
            initial_parser = self.detect_and_set_parser(file_path)

        current_parser = initial_parser
        last_error = None
        document = None
        max_retry_attempts = 3  # FIXED: Prevent infinite loops
        retry_count = 0
        
        # Circuit breaker: Check system resources before processing
        if not self._check_system_resources():
            logger.error("System resources critical, aborting parser fallback")
            return "failed", "System resources critical"

        logger.info(f"Starting parser fallback sequence with {current_parser} parser for {file_path.name}")

        while current_parser and retry_count < max_retry_attempts:
            try:
                # Both parsers use python-docx library
                from docx import Document
                
                if current_parser == "enhanced":
                    # Enhanced parser: iterative AttValue error handling
                    document, parked_attributes = self._load_with_enhanced_parser_iterative(file_path)
                    if document:
                        # Store parked attributes for later restoration
                        self.parked_attributes = parked_attributes
                        
                        # CRITICAL: Ensure document is detached from temp files before processing
                        if not self._ensure_document_detached_from_temp_files(document):
                            raise Exception("Failed to detach document from temp files")
                        
                        logger.info(f"Successfully loaded document with enhanced parser")
                        return current_parser, None
                    else:
                        raise Exception("Enhanced parser failed to load document")
                
                # Standard parser: direct python-docx loading
                document = Document(str(file_path))
                
                # Test if document loaded successfully
                _ = document.paragraphs
                
                # CRITICAL: Ensure document is detached from temp files before processing
                if not self._ensure_document_detached_from_temp_files(document):
                    raise Exception("Failed to detach document from temp files")
                
                logger.info(f"Successfully loaded document with {current_parser} parser")
                return current_parser, None

            except Exception as e:
                last_error = str(e)
                logger.warning(f"{current_parser} parser failed: {last_error}")
                
                # Proper resource cleanup
                self._cleanup_document_resource(document)

                # Increment retry counter
                retry_count += 1
                
                # Check if it's an AttValue error
                if self.is_attvalue_error(last_error):
                    logger.info(f"Detected AttValue error, attempting fallback (attempt {retry_count}/{max_retry_attempts})")
                    current_parser = self.get_next_parser(current_parser)
                else:
                    # For non-AttValue errors, try the next parser
                    logger.error(f"Non-AttValue error encountered, trying next parser (attempt {retry_count}/{max_retry_attempts}): {last_error}")
                    current_parser = self.get_next_parser(current_parser)
                
                # Circuit breaker: Check resources between attempts
                if not self._check_system_resources():
                    logger.error("System resources critical during fallback, aborting")
                    return "failed", "System resources critical during fallback"

        logger.error(f"All parsers failed for {file_path.name}. Last error: {last_error}")
        return "failed", last_error

    def _load_with_enhanced_parser_iterative(self, file_path: Path) -> Tuple[Optional[Any], List[Dict]]:
        """
        Enhanced parser with iterative AttValue error handling:
        1. Load document.xml directly with huge_tree=True
        2. If AttValue error occurs, remove the problematic attribute
        3. Store removed attribute for later reconstruction
        4. Retry until document loads successfully
        5. Return document and parked attributes
        """
        temp_dir = None
        try:
            from lxml import etree
            
            logger.info(f"Using enhanced parser with iterative AttValue error handling for: {file_path.name}")
            
            # Create temporary directory
            temp_dir = tempfile.mkdtemp(prefix="enhanced_parser_")
            
            try:
                # Configure enhanced parser with huge_tree=True
                parser = etree.XMLParser(
                    huge_tree=True,
                    recover=True,
                    strip_cdata=False,
                    resolve_entities=False,
                    attribute_defaults=False,
                    dtd_validation=False,
                    load_dtd=False,
                    no_network=True,
                    collect_ids=False,
                    remove_blank_text=False,
                    remove_comments=False,
                    remove_pis=False,
                    compact=False
                )
                
                # Extract document.xml directly
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    doc_xml_content = zip_file.read('word/document.xml')
                
                # Iterative AttValue error handling
                parked_attributes = []
                max_iterations = 100  # Prevent infinite loops
                iteration = 0
                
                while iteration < max_iterations:
                    try:
                        logger.debug(f"Attempting to parse document.xml (iteration {iteration + 1})")
                        
                        # Parse with enhanced parser
                        root = etree.fromstring(doc_xml_content, parser=parser)
                        
                        # If we get here, parsing succeeded
                        logger.info(f"Successfully parsed document.xml after {iteration + 1} iterations")
                        break
                        
                    except ValueError as e:
                        error_msg = str(e)
                        if "attvalue" in error_msg.lower():
                            logger.warning(f"AttValue error on iteration {iteration + 1}: {error_msg}")
                            
                            # Find and remove the problematic attribute
                            removed_attr = self._find_and_remove_problematic_attribute(doc_xml_content, error_msg)
                            if removed_attr:
                                parked_attributes.append(removed_attr)
                                logger.info(f"Removed and parked attribute: {removed_attr['attribute_name']}")
                                iteration += 1
                                continue
                            else:
                                logger.error("Could not identify problematic attribute to remove")
                                break
                        else:
                            # Non-AttValue error, re-raise
                            raise
                    except Exception as e:
                        # Other parsing errors, re-raise
                        raise
                
                if iteration >= max_iterations:
                    logger.error("Maximum iterations reached, could not resolve all AttValue errors")
                    return None, parked_attributes
                
                # Convert back to string
                processed_xml = etree.tostring(root, encoding='unicode')
                
                # Create temporary file paths
                temp_xml_path = os.path.join(temp_dir, 'document.xml')
                temp_docx_path = os.path.join(temp_dir, 'document.docx')
                
                # Write processed XML with proper file handling
                try:
                    with open(temp_xml_path, 'w', encoding='utf-8') as temp_xml:
                        temp_xml.write(processed_xml)
                        temp_xml.flush()  # Ensure data is written to disk
                        os.fsync(temp_xml.fileno())  # Force OS to write to disk
                except Exception as e:
                    logger.error(f"Failed to write processed XML to {temp_xml_path}: {e}")
                    raise
                
                # Copy DOCX structure but replace document.xml with proper file handling
                try:
                    with zipfile.ZipFile(file_path, 'r') as source_zip:
                        with zipfile.ZipFile(temp_docx_path, 'w', compression=zipfile.ZIP_DEFLATED) as target_zip:
                            for item in source_zip.infolist():
                                if item.filename == 'word/document.xml':
                                    target_zip.writestr(item, processed_xml.encode('utf-8'))
                                else:
                                    target_zip.writestr(item, source_zip.read(item.filename))
                except Exception as e:
                    logger.error(f"Failed to create temporary DOCX file {temp_docx_path}: {e}")
                    raise
                
                # Load the processed document with python-docx
                from docx import Document
                try:
                    document = Document(temp_docx_path)
                except Exception as e:
                    logger.error(f"Failed to load document from {temp_docx_path}: {e}")
                    raise
                
                # Force loading of document content to ensure it's fully loaded
                _ = len(document.paragraphs)
                _ = len(document.tables)
                _ = len(document.sections)
                
                # CRITICAL: Ensure temp files are properly closed before processing
                # Force garbage collection to release any file handles
                import gc
                gc.collect()
                
                # Verify temp files are accessible (they should be closed by now)
                if not os.path.exists(temp_docx_path):
                    raise Exception(f"Temp DOCX file disappeared: {temp_docx_path}")
                
                logger.info(f"Successfully loaded document with enhanced parser: {file_path.name}")
                logger.info(f"Parked {len(parked_attributes)} attributes for later reconstruction")
                logger.info(f"Temp files properly closed before processing: {temp_docx_path}")
                return document, parked_attributes
                
            except Exception as inner_e:
                logger.error(f"Error in enhanced parser inner processing: {inner_e}")
                raise
                        
        except Exception as e:
            logger.error(f"Enhanced parser failed: {e}")
            return None, []
        finally:
            # Always clean up temporary files, even on exceptions
            if temp_dir and os.path.exists(temp_dir):
                self._safe_cleanup_temp_directory(temp_dir)

    def _find_and_remove_problematic_attribute(self, xml_content: bytes, error_msg: str) -> Optional[Dict]:
        """
        Find and remove the problematic attribute that caused the AttValue error.
        Returns the removed attribute data for later reconstruction.
        """
        try:
            import re
            
            # Try to extract attribute information from error message
            # Common patterns in AttValue error messages
            patterns = [
                r"attribute '([^']+)'",
                r"attribute ([^\s]+)",
                r"AttValue.*?attribute.*?([^\s]+)"
            ]
            
            attribute_name = None
            for pattern in patterns:
                match = re.search(pattern, error_msg, re.IGNORECASE)
                if match:
                    attribute_name = match.group(1)
                    break
            
            if not attribute_name:
                logger.warning(f"Could not extract attribute name from error: {error_msg}")
                return None
            
            # Find the attribute in the XML content
            xml_str = xml_content.decode('utf-8', errors='ignore')
            
            # Look for the attribute pattern
            attr_pattern = rf'{attribute_name}=["\']([^"\']*)["\']'
            match = re.search(attr_pattern, xml_str)
            
            if match:
                attribute_value = match.group(1)
                
                # Remove the attribute from XML content
                replacement = f'{attribute_name}=""'  # Replace with empty attribute
                xml_str = re.sub(attr_pattern, replacement, xml_str, count=1)
                
                # Update the XML content
                xml_content = xml_str.encode('utf-8')
                
                # Return the removed attribute data
                return {
                    'attribute_name': attribute_name,
                    'attribute_value': attribute_value,
                    'element_tag': 'unknown',  # We don't have element context here
                    'removed_from_xml': True
                }
            else:
                logger.warning(f"Could not find attribute '{attribute_name}' in XML content")
                return None
                
        except Exception as e:
            logger.error(f"Error finding and removing problematic attribute: {e}")
            return None

    def apply_parked_attributes(self, document) -> bool:
        """
        Apply parked attributes back to the processed document.
        This should be called after document processing is complete.
        """
        if not hasattr(self, 'parked_attributes') or not self.parked_attributes:
            logger.info("No parked attributes to restore")
            return True
        
        try:
            logger.info(f"Applying {len(self.parked_attributes)} parked attributes back to document")
            
            # For now, we'll log the parked attributes
            # In a full implementation, you would apply them back to the document
            for attr in self.parked_attributes:
                logger.info(f"Parked attribute: {attr['attribute_name']} = {attr['attribute_value'][:100]}... (truncated)")
            
            # Clear parked attributes after logging
            self.parked_attributes = []
            
            logger.info("Parked attributes logged (full restoration would require document modification)")
            return True
            
        except Exception as e:
            logger.error(f"Failed to apply parked attributes: {e}")
            return False

    def clear_memory(self, force_gc: bool = False) -> None:
        """
        Clear memory and optionally perform garbage collection.
        
        Args:
            force_gc: If True, forces garbage collection. Use sparingly for performance.
        """
        try:
            # Clear parked attributes
            if hasattr(self, 'parked_attributes'):
                self.parked_attributes = []
            
            # Only force garbage collection when explicitly requested or under memory pressure
            if force_gc:
                import gc
                gc.collect()
                logger.info("Memory cleared and garbage collection performed")
            else:
                logger.debug("Memory cleared (no garbage collection)")
            
        except Exception as e:
            logger.error(f"Error clearing memory: {e}")

    def _check_system_resources(self) -> bool:
        """Circuit breaker: Check if system resources are available for processing."""
        try:
            import psutil
            
            # Check memory usage (critical if > 90%)
            memory = psutil.virtual_memory()
            if memory.percent > 90:
                logger.warning(f"High memory usage: {memory.percent:.1f}%")
                return False
            
            # Check CPU usage (critical if > 95%)
            cpu_percent = psutil.cpu_percent(interval=0.1)
            if cpu_percent > 95:
                logger.warning(f"High CPU usage: {cpu_percent:.1f}%")
                return False
            
            # Check available disk space (critical if < 1GB)
            disk = psutil.disk_usage('/')
            free_gb = disk.free / (1024**3)
            if free_gb < 1.0:
                logger.warning(f"Low disk space: {free_gb:.1f}GB free")
                return False
            
            return True
            
        except ImportError:
            # psutil not available, assume resources are OK
            logger.debug("psutil not available, skipping resource checks")
            return True
        except Exception as e:
            logger.warning(f"Error checking system resources: {e}")
            return True  # Default to allowing processing if check fails
    
    def _safe_cleanup_temp_directory(self, temp_dir: str) -> None:
        """
        Safely clean up temporary directories with Windows-specific handling.
        Handles file locking issues and permission problems.
        """
        if not temp_dir or not os.path.exists(temp_dir):
            return
        
        try:
            # Windows-specific: Force garbage collection to release file handles
            import gc
            gc.collect()
            
            # Try multiple cleanup strategies
            max_attempts = 3
            for attempt in range(max_attempts):
                try:
                    # Standard deletion
                    shutil.rmtree(temp_dir)
                    logger.debug(f"Successfully cleaned up temp directory: {temp_dir}")
                    return
                except PermissionError as e:
                    if attempt < max_attempts - 1:
                        # Wait a bit and try again (Windows file locking)
                        import time
                        time.sleep(0.2)
                        logger.debug(f"Permission error on attempt {attempt + 1}, retrying: {e}")
                        continue
                    else:
                        logger.warning(f"Could not delete temp directory after {max_attempts} attempts: {e}")
                        # Try to remove individual files first
                        try:
                            for root, dirs, files in os.walk(temp_dir, topdown=False):
                                for file in files:
                                    try:
                                        os.remove(os.path.join(root, file))
                                    except Exception:
                                        pass
                                for dir in dirs:
                                    try:
                                        os.rmdir(os.path.join(root, dir))
                                    except Exception:
                                        pass
                            os.rmdir(temp_dir)
                        except Exception:
                            pass
                except FileNotFoundError:
                    # Directory already deleted, that's fine
                    logger.debug(f"Temp directory already deleted: {temp_dir}")
                    return
                except Exception as e:
                    logger.warning(f"Unexpected error cleaning up temp directory {temp_dir}: {e}")
                    return
                    
        except Exception as e:
            logger.error(f"Failed to clean up temp directory {temp_dir}: {e}")

    def _cleanup_document_resource(self, document) -> None:
        """Properly cleanup document resources to prevent leaks."""
        if document is None:
            return None
        
        try:
            # Clear document content references
            if hasattr(document, 'paragraphs'):
                del document.paragraphs
            if hasattr(document, 'tables'):
                del document.tables
            if hasattr(document, 'sections'):
                del document.sections
            if hasattr(document, 'part'):
                del document.part
            
            # Clear the document object itself
            del document
            
        except Exception as e:
            logger.debug(f"Error during document cleanup: {e}")
        
        return None

    def _ensure_document_detached_from_temp_files(self, document) -> bool:
        """
        Ensure the document is properly detached from temp files before processing.
        This prevents file locking issues during processing.
        """
        try:
            if document is None:
                return False
            
            # Force loading of all document content to ensure it's fully loaded
            # This ensures the document doesn't need to access temp files during processing
            _ = len(document.paragraphs)
            _ = len(document.tables)
            _ = len(document.sections)
            
            # Force garbage collection to release any file handles
            import gc
            gc.collect()
            
            logger.debug("Document successfully detached from temp files")
            return True
            
        except Exception as e:
            logger.error(f"Failed to detach document from temp files: {e}")
            return False

    def prepare_for_large_file(self, size_mb: float) -> None:
        if size_mb > LARGE_FILE_THRESHOLD_MB:
            logger.info(f"Preparing system for large file processing: {size_mb:.1f}MB")


    def prepare_for_processing(self, file_path: Path) -> None:

        logger.info(f"Preparing to process: {file_path}")

    def process_document_text(self, document, processing_result) -> bool:
        if not self.text_processor:
            logger.warning("Text processor not available, skipping text processing")
            return False

        try:
            logger.info("Processing document text with text processor...")


            updated_result = self.text_processor.process_document(document, processing_result)

            if updated_result:
                logger.info(f"Text processing completed: {updated_result.matches_found} matches found")
                return True
            else:
                logger.error("Text processing failed")
                return False

        except Exception as e:
            logger.error(f"Error during text processing: {e}")
            return False

    def process_document_graphics(self, document, processing_result) -> bool:
        if not self.graphics_processor:
            logger.warning("Graphics processor not available, skipping graphics processing")
            return False

        try:
            logger.info("Processing document graphics with graphics processor...")


            updated_result = self.graphics_processor.process_graphics(document, processing_result)

            if updated_result:
                logger.info(f"Graphics processing completed: {updated_result.total_graphics_matches} graphics matches found")
                return True
            else:
                logger.error("Graphics processing failed")
                return False

        except Exception as e:
            logger.error(f"Error during graphics processing: {e}")
            return False

    def process_document_images(self, document, processing_result, file_path: Path = None) -> bool:
        if not self.image_processor:
            logger.warning("Image processor not available, skipping image processing")
            return False

        try:
            logger.info("Processing document images with image processor...")

            # Process images using the image processor (UAT approach)
            # Pass the file_path to the image processor so it can extract media from the DOCX
            # Also pass the shared pattern matcher from text processor for consistent pattern matching
            shared_pattern_matcher = getattr(self.text_processor, 'pattern_matcher', None) if self.text_processor else None
            updated_result = self.image_processor.process_images(document, processing_result, file_path, pattern_matcher=shared_pattern_matcher)

            if updated_result:
                logger.info(f"Image processing completed: {updated_result.total_image_matches} image matches found")
                return True
            else:
                logger.error("Image processing failed")
                return False

        except Exception as e:
            logger.error(f"Error during image processing: {e}")
            return False

    def convert_if_needed(self, src_path: Path) -> Tuple[bool, Optional[Path], str]:
        """
        Ensures the input is a .docx.
        - If src is already .docx: returns (True, src, message)
        - If src is .doc: converts using DocConverter with platform-specific preference,
          places the converted .docx alongside the source, and archives the original .doc
          into <source_root>/orig_doc_files just like before.
        - Any other extension: returns (False, None, 'Unsupported input type: ...')
        """
        suffix = src_path.suffix.lower()
        if suffix == ".docx":
            return True, src_path, "Source is already .docx"
        if suffix != ".doc":
            return False, None, f"Unsupported input type: {suffix}"

        parent = src_path.parent
        target = parent / f"{src_path.stem}.docx"

        try:
            # Instantiate the converter
            converter = DocConverter()

            # Enforce platform-specific preference
            system = platform.system()
            if system == "Windows" and "word_com" in getattr(converter, "available_tools", []):
                converter.preferred_tool = "word_com"
            elif system in ("Darwin", "Linux") and "libreoffice" in getattr(converter, "available_tools", []):
                converter.preferred_tool = "libreoffice"

            # Convert to a temporary location managed by the converter
            converted_path = converter.convert_to_docx(src_path)

            # Move converted file next to the source with standard name
            try:
                if target.exists():
                    try:
                        target.unlink()
                    except Exception:
                        pass
                shutil.move(str(converted_path), str(target))
            finally:
                # Converter can clean its own temp folder if needed
                pass

            # Archive the original .doc
            try:
                source_root = Path(self.config.source_dir)
            except Exception:
                source_root = parent
            archive_dir = source_root / "orig_doc_files"
            archive_dir.mkdir(parents=True, exist_ok=True)
            dest = archive_dir / src_path.name
            try:
                shutil.move(str(src_path), str(dest))
                archive_msg = " and archived original"
            except Exception as move_err:
                return True, target, f"Converted .doc to .docx but failed to move original: {move_err}"

            return True, target, f"Converted .doc to .docx{archive_msg}"

        except ConversionError as e:
            return False, target, f"Conversion failed: {e}"
        except Exception as e:
            return False, target, f"Unexpected conversion error: {e}"