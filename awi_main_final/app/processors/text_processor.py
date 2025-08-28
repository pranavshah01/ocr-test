"""
Text Processor for DOCX document processing.

This module provides text processing capabilities for DOCX documents including:
- Pattern matching using regex patterns
- Text replacement with formatting preservation
- Comprehensive detection reporting
- Integration with ProcessingResult for unified reporting
"""

import re
import logging
import time
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..core.models import ProcessingResult, MatchDetail, ProcessorType, MatchFlag, FallbackFlag
from ..utils.text_utils.text_docx_utils import TextReconstructor, FontManager, PatternMatcher, create_pattern_matcher, load_patterns_and_mappings
from config import DEFAULT_MAPPING, DEFAULT_SEPARATOR, PROCESSING_MODES

logger = logging.getLogger(__name__)


class TextReplacer:
    """Handles text replacement while preserving formatting."""
    
    def __init__(self, mode: str = PROCESSING_MODES['APPEND'], separator: str = DEFAULT_SEPARATOR, default_mapping: str = DEFAULT_MAPPING):
        """
        Initialize text replacer.
        
        Args:
            mode: Replacement mode ('append' or 'replace')
            separator: Separator between original and appended text in append mode
            default_mapping: Default text to append when no mapping is found
        """
        self.mode = mode
        self.separator = separator
        self.default_mapping = default_mapping
    
    def replace_text_in_runs(self, runs: List, original_text: str, replacement_text: str, 
                           start_pos: int, end_pos: int) -> bool:
        """
        Replace text across multiple runs while preserving formatting.
        
        Args:
            runs: List of runs containing the text
            original_text: Original text to replace
            replacement_text: Replacement text
            start_pos: Start position in reconstructed text
            end_pos: End position in reconstructed text
            
        Returns:
            True if replacement was successful
        """
        try:
            # Find which runs are affected
            current_pos = 0
            affected_runs = []
            
            for run in runs:
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                if run_start < end_pos and run_end > start_pos:
                    affected_runs.append((run, run_start, run_end))
                
                current_pos = run_end
            
            if not affected_runs:
                return False
            
            # Determine replacement text based on mode
            if self.mode == PROCESSING_MODES['APPEND']:
                # Check if the replacement text already exists after the original text
                # This prevents duplicate appending when the same text appears in multiple cells
                expected_append = f"{original_text}{self.separator}{replacement_text}"
                
                # Get the current text in the affected runs to check for existing append
                current_text = ""
                for run in runs:
                    current_text += run.text
                
                # Check if the expected append already exists
                if expected_append in current_text:
                    logger.info(f"APPEND MODE: Skipping duplicate append for '{original_text}' -> '{replacement_text}' (already exists)")
                    return True  # Return True to indicate "success" (no action needed)
                
                final_text = expected_append
            else:  # replace mode
                final_text = replacement_text
            
            # Get font info from the first affected run
            first_run = affected_runs[0][0]
            font_info = FontManager.get_font_info(first_run)
            
            # Clear text from all affected runs except the first
            for i, (run, run_start, run_end) in enumerate(affected_runs):
                if i == 0:
                    # Calculate the portion of text to replace in the first run
                    run_match_start = max(0, start_pos - run_start)
                    run_match_end = min(len(run.text), end_pos - run_start)
                    
                    # Replace the text in the first run
                    before_text = run.text[:run_match_start]
                    after_text = run.text[run_match_end:]
                    run.text = before_text + final_text + after_text
                    
                    # Apply original formatting
                    FontManager.apply_font_info(run, font_info)
                else:
                    # Clear subsequent runs that were part of the match
                    run_match_start = max(0, start_pos - run_start)
                    run_match_end = min(len(run.text), end_pos - run_start)
                    
                    before_text = run.text[:run_match_start]
                    after_text = run.text[run_match_end:]
                    run.text = before_text + after_text
            
            return True
            
        except Exception as e:
            logger.error(f"Failed to replace text: {e}")
            return False


class TextProcessor:
    """Text processor for DOCX document processing with comprehensive detection."""
    
    def __init__(self, config):
        """
        Initialize text processor.
        
        Args:
            config: Configuration object containing patterns and mappings
        """
        self.config = config
        self.initialized = False
        self.patterns = {}
        self.mappings = {}
        self.pattern_matcher = None
        self.text_replacer = None
        self.mode = getattr(config, 'mode', PROCESSING_MODES['APPEND'])
        self.separator = getattr(config, 'separator', DEFAULT_SEPARATOR)
        self.default_mapping = getattr(config, 'default_mapping', DEFAULT_MAPPING)
        
        # Load patterns and mappings
        self._load_patterns_and_mappings()
    
    def _load_patterns_and_mappings(self):
        """Load patterns and mappings from configuration files."""
        self.patterns, self.mappings = load_patterns_and_mappings(self.config)
    
    def initialize(self) -> bool:
        """Initialize the text processor."""
        try:
            # Create pattern matcher once during initialization
            self.pattern_matcher = create_pattern_matcher(self.patterns, self.mappings)
            
            # Create text replacer
            self.text_replacer = TextReplacer(self.mode, self.separator, self.default_mapping)
            
            logger.info(f"Pattern matcher created with {len(self.patterns)} patterns, {len(self.mappings)} mappings")
            
            self.initialized = True
            logger.info(f"Text processor initialized with {len(self.patterns)} patterns, {len(self.mappings)} mappings")
            return True
        except Exception as e:
            logger.error(f"Error initializing Text Processor: {e}")
            self.initialized = False
            return False
    
    def is_initialized(self) -> bool:
        """Check if the processor is initialized."""
        return self.initialized
    
    def process_document(self, document: Document, processing_result: ProcessingResult) -> ProcessingResult:
        """
        Process document text and update ProcessingResult with detection details.
        
        Args:
            document: The DOCX document to process
            processing_result: ProcessingResult object to update with detection details
            
        Returns:
            Updated ProcessingResult object
        """
        if not self.initialized:
            logger.error("Text processor not initialized")
            processing_result.error_message = "Text processor not initialized"
            return processing_result
        
        start_time = time.time()
        logger.info("Starting text processing...")
        
        try:
            # Store document for font information access
            self.document = document
            
            # Process all text content
            all_detections = []
            
            # Process main document body paragraphs
            body_detections = self._process_paragraphs(document.paragraphs, "body")
            all_detections.extend(body_detections)
            
            # Process tables
            table_detections = self._process_tables(document.tables)
            all_detections.extend(table_detections)
            
            # Process headers and footers
            header_footer_detections = self._process_headers_footers(document)
            all_detections.extend(header_footer_detections)
            
            # Log detailed detection information
            logger.info(f"Found {len(all_detections)} pattern matches:")
            for i, detection in enumerate(all_detections, 1):
                pattern_name = detection.get('pattern_name', 'Unknown')
                matched_text = detection.get('matched_text', 'Unknown')
                replacement_text = detection.get('replacement_text', 'Unknown')
                location = detection.get('location', 'Unknown')
                logger.info(f"  {i}. Pattern '{pattern_name}' matched '{matched_text}' -> '{replacement_text}' at {location}")
            
            # Update ProcessingResult with detection details
            self._update_processing_result(processing_result, all_detections)
            
            # Reconstruct the document with the detected matches
            reconstruction_results = self.reconstruct_document(document, all_detections)
            
            # Update match details with reconstruction status
            self._update_reconstruction_status(processing_result, all_detections, reconstruction_results)
            
            processing_time = time.time() - start_time
            logger.info(f"Text processing completed in {processing_time:.2f}s: {len(all_detections)} detections")
            
            return processing_result
            
        except Exception as e:
            error_msg = f"Error during text processing: {e}"
            logger.error(error_msg)
            processing_result.error_message = error_msg
            return processing_result
    
    def reconstruct_document(self, document: Document, detections: List[Dict[str, Any]]) -> Dict[str, bool]:
        """
        Reconstruct the document by applying text replacements based on detections.
        
        Args:
            document: The document to reconstruct
            detections: List of detection dictionaries with match details
            
        Returns:
            Dictionary mapping detection keys to reconstruction success status
        """
        logger.info("Starting document reconstruction...")
        
        reconstruction_results = {}
        
        try:
            # Process main document body paragraphs
            body_results = self._reconstruct_paragraphs(document.paragraphs, detections, "body")
            reconstruction_results.update(body_results)
            
            # Process tables
            table_results = self._reconstruct_tables(document.tables, detections)
            reconstruction_results.update(table_results)
            
            # Process headers and footers
            header_footer_results = self._reconstruct_headers_footers(document, detections)
            reconstruction_results.update(header_footer_results)
            
            logger.info(f"Document reconstruction completed: {sum(reconstruction_results.values())}/{len(reconstruction_results)} matches reconstructed successfully")
            
        except Exception as e:
            logger.error(f"Error during document reconstruction: {e}")
        
        return reconstruction_results
    
    def _reconstruct_paragraphs(self, paragraphs: List[Paragraph], detections: List[Dict[str, Any]], location: str) -> Dict[str, bool]:
        """
        Reconstruct paragraphs by applying text replacements.
        
        Args:
            paragraphs: List of paragraphs to reconstruct
            detections: List of detection dictionaries
            location: Location description
            
        Returns:
            Dictionary mapping detection keys to reconstruction success status
        """
        reconstruction_results = {}
        
        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_location = f"{location}_paragraph_{i}"
                paragraph_results = self._reconstruct_paragraph(paragraph, detections, paragraph_location)
                reconstruction_results.update(paragraph_results)
            except Exception as e:
                logger.error(f"Error reconstructing paragraph {i} in {location}: {e}")
        
        return reconstruction_results
    
    def _reconstruct_paragraph(self, paragraph: Paragraph, detections: List[Dict[str, Any]], location: str) -> Dict[str, bool]:
        """
        Reconstruct individual paragraph by applying text replacements.
        
        Args:
            paragraph: Paragraph to reconstruct
            detections: List of detection dictionaries
            location: Location description
            
        Returns:
            Dictionary mapping detection keys to reconstruction success status
        """
        reconstruction_results = {}
        
        if not paragraph.runs:
            return reconstruction_results
        
        # Reconstruct full text from runs
        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)
        
        if not full_text.strip():
            return reconstruction_results
        
        # Find detections for this paragraph location
        paragraph_detections = [
            detection for detection in detections 
            if detection.get('location') == location and detection.get('is_matched', False)
        ]
        
        # Sort detections by position (reverse order to avoid position shifts)
        paragraph_detections.sort(key=lambda x: x.get('start_pos', 0), reverse=True)
        
        # Apply replacements
        for detection in paragraph_detections:
            try:
                matched_text = detection.get('matched_text', '')
                replacement_text = detection.get('replacement_text', '')
                start_pos = detection.get('start_pos', 0)
                end_pos = detection.get('end_pos', 0)
                
                if not matched_text or not replacement_text:
                    continue
                
                # Create a unique key for this detection
                detection_key = f"{location}_{matched_text}_{start_pos}_{end_pos}"
                
                # Find affected runs
                text_span = TextReconstructor.find_text_in_runs(runs, matched_text, start_pos)
                if not text_span:
                    reconstruction_results[detection_key] = False
                    continue
                
                span_start, span_end, affected_runs = text_span
                
                # Perform text replacement
                success = self.text_replacer.replace_text_in_runs(
                    runs, matched_text, replacement_text, span_start, span_end
                )
                
                reconstruction_results[detection_key] = success
                
                if success:
                    logger.info(f"Text replacement successful: '{matched_text}' -> '{replacement_text}' at {location}")
                else:
                    logger.warning(f"Text replacement failed: '{matched_text}' -> '{replacement_text}' at {location}")
                    
            except Exception as e:
                logger.error(f"Error applying replacement for detection at {location}: {e}")
                detection_key = f"{location}_{matched_text}_{start_pos}_{end_pos}"
                reconstruction_results[detection_key] = False
        
        return reconstruction_results
    
    def _reconstruct_tables(self, tables: List[Table], detections: List[Dict[str, Any]], location_prefix: str = "") -> Dict[str, bool]:
        """
        Reconstruct tables by applying text replacements.
        
        Args:
            tables: List of tables to reconstruct
            detections: List of detection dictionaries
            location_prefix: Prefix for location description (e.g., "header_section_0")
            
        Returns:
            Dictionary mapping detection keys to reconstruction success status
        """
        reconstruction_results = {}
        
        # Import the table utilities
        from ..utils.table_utils import get_table_cells_to_process
        
        for table_idx, table in enumerate(tables):
            try:
                # Get cells that should be processed (skipping vMerge continue cells)
                cells_to_process = get_table_cells_to_process(table)
                
                for row_idx, cell_idx, cell in cells_to_process:
                    if location_prefix:
                        location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    else:
                        location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    cell_results = self._reconstruct_paragraphs(cell.paragraphs, detections, location)
                    reconstruction_results.update(cell_results)
                    
            except Exception as e:
                logger.error(f"Error reconstructing table {table_idx}: {e}")
        
        return reconstruction_results
    
    def _reconstruct_headers_footers(self, document: Document, detections: List[Dict[str, Any]]) -> Dict[str, bool]:
        """
        Reconstruct headers and footers by applying text replacements.
        
        Args:
            document: Document to reconstruct
            detections: List of detection dictionaries
            
        Returns:
            Dictionary mapping detection keys to reconstruction success status
        """
        reconstruction_results = {}
        
        try:
            # Process headers and footers (including tables)
            for section_idx, section in enumerate(document.sections):
                # Process header paragraphs and tables
                if section.header:
                    header_results = self._reconstruct_paragraphs(section.header.paragraphs, detections, f"header_section_{section_idx}")
                    reconstruction_results.update(header_results)
                    header_table_results = self._reconstruct_tables(section.header.tables, detections, f"header_section_{section_idx}")
                    reconstruction_results.update(header_table_results)
                
                # Process footer paragraphs and tables
                if section.footer:
                    footer_results = self._reconstruct_paragraphs(section.footer.paragraphs, detections, f"footer_section_{section_idx}")
                    reconstruction_results.update(footer_results)
                    footer_table_results = self._reconstruct_tables(section.footer.tables, detections, f"footer_section_{section_idx}")
                    reconstruction_results.update(footer_table_results)
                    
        except Exception as e:
            logger.error(f"Error reconstructing headers/footers: {e}")
        
        return reconstruction_results
    
    def _process_paragraphs(self, paragraphs: List[Paragraph], location: str) -> List[Dict[str, Any]]:
        """
        Process paragraphs for pattern detection.
        
        Args:
            paragraphs: List of paragraphs to process
            location: Location description
            
        Returns:
            List of detection dictionaries
        """
        all_detections = []
        
        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_detections = self._process_paragraph(
                    paragraph, f"{location}_paragraph_{i}"
                )
                all_detections.extend(paragraph_detections)
            except Exception as e:
                logger.error(f"Error processing paragraph {i} in {location}: {e}")
        
        return all_detections
    
    def _process_paragraph(self, paragraph: Paragraph, location: str) -> List[Dict[str, Any]]:
        """
        Process individual paragraph for pattern detection (same logic as old AWI).
        
        Args:
            paragraph: Paragraph to process
            location: Location description
            
        Returns:
            List of detection dictionaries
        """
        if not paragraph.runs:
            return []
        
        # Reconstruct full text from runs (same as old AWI)
        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)
        
        if not full_text.strip():
            return []
        
        # Find ALL pattern matches (including those without mappings for comprehensive reporting)
        all_pattern_matches = self.pattern_matcher.find_all_pattern_matches(full_text)
        
        detection_results = []
        
        for pattern_name, matched_text, start_pos, end_pos in all_pattern_matches:
            # Get replacement text
            replacement_text = self.pattern_matcher.get_replacement(matched_text)
            
            # Mode-specific behavior:
            # - In "replace" mode: only process if we have a valid mapping (no default)
            # - In "append" mode: use default mapping if no valid mapping found
            if not replacement_text:
                if self.mode == "replace":
                    # In replace mode, skip if no mapping found
                    logger.debug(f"REPLACE MODE: Skipping '{matched_text}' - no mapping found")
                    replacement_text = None
                elif self.mode == PROCESSING_MODES['APPEND']:
                    # In append mode, use default mapping
                    replacement_text = self.default_mapping
                    logger.info(f"APPEND MODE: Using default mapping '{self.default_mapping}' for '{matched_text}'")
                else:
                    # Unknown mode, skip
                    logger.warning(f"Unknown mode '{self.mode}', skipping '{matched_text}'")
                    replacement_text = None
            
            # Determine if this pattern was successfully matched
            is_matched = replacement_text is not None
            
            # Get the actual pattern from patterns
            actual_pattern = self.patterns.get(pattern_name, pattern_name)
            
            # Get font information for this detection using consolidated method
            font_info = FontManager.extract_font_info_for_detection(
                runs, matched_text, start_pos, getattr(self, 'document', None)
            )
            
            # Determine content type and dimensions (same as old AWI)
            content_type = "Paragraph"
            dimension = ""
            
            if "table" in location.lower():
                content_type = "Table"
                # For tables, we'll need to get cell dimensions
                dimension = "Cell size"  # Placeholder - would need table cell info
            elif "header" in location.lower():
                content_type = "Header"
            elif "footer" in location.lower():
                content_type = "Footer"
            
            # Create detection result in ProcessingResult format (same as old AWI)
            detection_result = {
                'pattern_name': pattern_name,
                'actual_pattern': actual_pattern,
                'matched_text': matched_text,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'replacement_text': replacement_text,
                'location': location,
                'content_type': content_type,
                'dimension': dimension,
                'processor': 'Text',
                'font_info': font_info,
                'is_matched': is_matched,
                'confidence': 1.0
            }
            
            detection_results.append(detection_result)
        
        return detection_results
    
    def _process_tables(self, tables: List[Table], location_prefix: str = "") -> List[Dict[str, Any]]:
        """
        Process tables for pattern detection, handling vMerge cells properly.
        
        Args:
            tables: List of tables to process
            location_prefix: Prefix for location description
            
        Returns:
            List of detection dictionaries
        """
        all_detections = []
        
        # Import the table utilities
        from ..utils.table_utils import get_table_cells_to_process
        
        for table_idx, table in enumerate(tables):
            try:
                # Get cells that should be processed (skipping vMerge continue cells)
                cells_to_process = get_table_cells_to_process(table)
                
                for row_idx, cell_idx, cell in cells_to_process:
                    if location_prefix:
                        location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    else:
                        location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    
                    cell_detections = self._process_paragraphs(cell.paragraphs, location)
                    all_detections.extend(cell_detections)
                    
                    logger.debug(f"Processed cell at {location} with {len(cell_detections)} detections")
                    
            except Exception as e:
                logger.error(f"Error processing table {table_idx}: {e}")
        
        return all_detections
    
    def _process_headers_footers(self, document: Document) -> List[Dict[str, Any]]:
        """
        Process headers and footers for pattern detection (including tables).
        
        Args:
            document: Document to process
            
        Returns:
            List of detection dictionaries
        """
        all_detections = []
        
        try:
            # Process headers and footers (including tables)
            for section_idx, section in enumerate(document.sections):
                # Process header paragraphs and tables
                if section.header:
                    header_detections = self._process_paragraphs(section.header.paragraphs, f"header_section_{section_idx}")
                    all_detections.extend(header_detections)
                    
                    # Process tables in header
                    header_table_detections = self._process_tables(section.header.tables, f"header_section_{section_idx}")
                    all_detections.extend(header_table_detections)
                
                # Process footer paragraphs and tables
                if section.footer:
                    footer_detections = self._process_paragraphs(section.footer.paragraphs, f"footer_section_{section_idx}")
                    all_detections.extend(footer_detections)
                    
                    # Process tables in footer
                    footer_table_detections = self._process_tables(section.footer.tables, f"footer_section_{section_idx}")
                    all_detections.extend(footer_table_detections)
        except Exception as e:
            logger.error(f"Error processing headers/footers: {e}")
        
        return all_detections
    
    def _update_processing_result(self, processing_result: ProcessingResult, detections: List[Dict[str, Any]]):
        """
        Update ProcessingResult with detection details.
        
        Args:
            processing_result: ProcessingResult to update
            detections: List of detection dictionaries
        """
        # Count matches by type
        text_matches = 0
        graphics_matches = 0
        image_matches = 0
        graphics_no_match = 0
        image_no_match = 0
        
        # Convert detections to MatchDetail objects
        match_details = []
        sr_no = 1
        
        for detection in detections:
            # Count matches
            if detection.get('processor') == 'Text':
                if detection.get('is_matched', False):
                    text_matches += 1
            elif detection.get('processor') == 'Graphics':
                if detection.get('is_matched', False):
                    graphics_matches += 1
                else:
                    graphics_no_match += 1
            elif detection.get('processor') == 'Image':
                if detection.get('is_matched', False):
                    image_matches += 1
                else:
                    image_no_match += 1
            
            # Convert to MatchDetail
            match_detail_dict = self._convert_detection_to_match_detail(
                detection, sr_no, ProcessorType.TEXT.value
            )
            
            # Create MatchDetail object
            match_detail = MatchDetail(
                sr_no=match_detail_dict['sr_no'],
                type=ProcessorType.TEXT,
                orig_id_name=match_detail_dict['orig_id_name'],
                src_text=match_detail_dict['src_text'],
                src_text_font=match_detail_dict['src_text_font'],
                src_text_color=match_detail_dict['src_text_color'],
                src_text_size=match_detail_dict['src_text_size'],
                src_dimension=match_detail_dict['src_dimension'],
                mapped_text=match_detail_dict['mapped_text'],
                mapped_text_font=match_detail_dict['mapped_text_font'],
                mapped_text_color=match_detail_dict['mapped_text_color'],
                mapped_text_size=match_detail_dict['mapped_text_size'],
                match_flag=MatchFlag.YES if match_detail_dict['match_flag'] == 'Y' else MatchFlag.NO,
                is_fallback=FallbackFlag.NO,
                reasoning=None
            )
            
            match_details.append(match_detail)
            sr_no += 1
        
        # Update ProcessingResult
        processing_result.total_text_matches = text_matches
        processing_result.total_graphics_matches = graphics_matches
        processing_result.total_image_matches = image_matches
        processing_result.total_graphics_no_match = graphics_no_match
        processing_result.total_image_no_match = image_no_match
        processing_result.matches_found = text_matches + graphics_matches + image_matches
        processing_result.total_matches = len(detections)
        processing_result.match_details = match_details
        
        # Update processor type
        processing_result.processor_type = "text_processor"
        
        # Set success flag
        processing_result.success = True
        
        # Update processed file name
        if hasattr(self, 'document') and self.document:
            # Get the base name without extension and add suffix before .docx
            base_name = Path(processing_result.file_name).stem
            processing_result.processed_file_name = f"{base_name}{getattr(self.config, 'suffix', '_12NC_processed')}.docx"
    
    def _convert_detection_to_match_detail(
        self,
        detection: Dict[str, Any],
        sr_no: int,
        processor_type: str
    ) -> Dict[str, Any]:
        """
        Convert detection dictionary to MatchDetail format for ProcessingResult.
        
        Args:
            detection: Detection dictionary
            sr_no: Serial number for the match
            processor_type: Type of processor (TEXT, GRAPHICS, IMAGE)
            
        Returns:
            Dictionary in MatchDetail format
        """
        # Extract font information from the detection
        font_info = detection.get('font_info', {})
        font_family = font_info.get('font_family', 'Arial')
        font_size = font_info.get('font_size', '12.0')
        font_color = font_info.get('color', '000000')
        
        return {
            'sr_no': sr_no,
            'type': processor_type,
            'orig_id_name': detection.get('location', ''),
            'src_text': detection.get('matched_text', ''),
            'src_text_font': font_family,
            'src_text_color': font_color,
            'src_text_size': font_size,
            'src_dimension': detection.get('dimension', ''),
            'mapped_text': detection.get('replacement_text', ''),
            'mapped_text_font': font_family,  # Use same font for mapped text
            'mapped_text_color': font_color,  # Use same color for mapped text
            'mapped_text_size': font_size,    # Use same size for mapped text
            'match_flag': 'Y' if detection.get('is_matched', False) else 'N',
            'is_fallback': 'N',  # Will be updated during processing
            'reasoning': None,  # Will be populated by specific processors
            'reconstructed': False  # Will be updated with reconstruction status
        }
    
    def _update_reconstruction_status(self, processing_result: ProcessingResult, detections: List[Dict[str, Any]], reconstruction_results: Dict[str, bool]):
        """
        Update match details with reconstruction status.
        
        Args:
            processing_result: ProcessingResult to update
            detections: List of detection dictionaries
            reconstruction_results: Dictionary mapping detection keys to reconstruction success status
        """
        for match_detail in processing_result.match_details:
            # Find corresponding detection
            for detection in detections:
                if (detection.get('matched_text') == match_detail.src_text and
                    detection.get('location') == match_detail.orig_id_name):
                    
                    # Create the same key used in reconstruction
                    detection_key = f"{detection.get('location')}_{detection.get('matched_text')}_{detection.get('start_pos', 0)}_{detection.get('end_pos', 0)}"
                    
                    # Update reconstruction status
                    match_detail.reconstructed = reconstruction_results.get(detection_key, False)
                    break
    
    def cleanup(self):
        """Clean up resources used by the text processor."""
        logger.info("Cleaning up text processor")
        self.initialized = False


def create_text_processor(config) -> TextProcessor:
    """Factory function to create a TextProcessor instance."""
    return TextProcessor(config)
