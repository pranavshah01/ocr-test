"""
Core document processor for the document processing pipeline.
Orchestrates the processing of Word documents through text, graphics, and image processors.
"""

import json
import time
from pathlib import Path
from typing import Dict, List, Optional, Any
import logging

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from ..utils.shared_constants import ERROR_CODES, SharedUtilities
from ..utils.platform_utils import PathManager
from ..converters.doc_converter import DocConverter, ConversionError
from .models import ProcessingResult, ProcessingLog

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Main orchestrator for document processing pipeline."""
    
    def __init__(self, config):
        """
        Initialize document processor with configuration.
        
        Args:
            config: ProcessingConfig instance
        """
        self.config = config
        self.converter = DocConverter() if config else None
        
        # Load patterns and mappings
        self.patterns = self._load_json_file(config.patterns_file) if config else {}
        self.mappings = self._load_json_file(config.mappings_file) if config else {}
        
        # Initialize processors (will be set when processors are implemented)
        self.text_processor = None
        self.graphics_processor = None
        self.image_processor = None
        
        logger.info("Document processor initialized")
    
    def _load_json_file(self, file_path: Path) -> Dict[str, Any]:
        """Load JSON file with error handling."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.debug(f"Loaded JSON file: {file_path}")
            return data
        except Exception as e:
            logger.error(f"Failed to load JSON file {file_path}: {e}")
            return {}
    
    def process_document(self, file_path: Path) -> ProcessingResult:
        """
        Process a single document through the complete pipeline.
        
        Args:
            file_path: Path to document file
            
        Returns:
            ProcessingResult with processing details
        """
        start_time = time.time()
        result = ProcessingResult(
            success=False,
            input_path=file_path,
            processing_time=0.0
        )
        
        try:
            logger.info(f"Starting processing of: {file_path}")
            
            # Step 1: Handle format conversion if needed
            docx_path = self._ensure_docx_format(file_path)
            if not docx_path:
                result.error_message = "Failed to convert document to DOCX format"
                return result
            
            # Step 2: Load document
            document = self._load_document(docx_path)
            if not document:
                result.error_message = "Failed to load document"
                return result
            
            # Step 3: Create processing log
            processing_log = ProcessingLog(
                document_path=file_path,
                processing_time=0.0
            )
            
            # Step 4: Process through pipeline with graceful degradation
            self._process_pipeline(document, processing_log)
            
            # Step 5: Save processed document
            output_path = self._generate_output_path(file_path)
            if self._save_document(document, output_path):
                result.output_path = output_path
                result.success = True
                logger.info(f"Successfully processed: {file_path} -> {output_path}")
            else:
                result.error_message = "Failed to save processed document"
            
            # Update result with processing log
            result.processing_log = processing_log
            
        except Exception as e:
            logger.error(f"Unexpected error processing {file_path}: {e}")
            result.error_message = str(e)
        
        finally:
            result.processing_time = time.time() - start_time
            logger.info(f"Processing completed in {result.processing_time:.2f}s")
        
        return result
    
    def _ensure_docx_format(self, file_path: Path) -> Optional[Path]:
        """
        Ensure document is in DOCX format, convert if necessary.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Path to DOCX file or None if conversion failed
        """
        if file_path.suffix.lower() == '.docx':
            return file_path
        
        if file_path.suffix.lower() == '.doc':
            if not self.converter:
                logger.error("No converter available for .doc files")
                return None
            
            try:
                logger.info(f"Converting .doc to .docx: {file_path}")
                return self.converter.convert_to_docx(file_path)
            except ConversionError as e:
                logger.error(f"Conversion failed: {e}")
                return None
        
        logger.error(f"Unsupported file format: {file_path.suffix}")
        return None
    
    def _load_document(self, docx_path: Path) -> Optional[Document]:
        """
        Load DOCX document with error handling and custom parser for large files.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Document instance or None if loading failed
        """
        file_size = docx_path.stat().st_size
        size_mb = file_size / (1024 * 1024)
        
        # Prepare system for large files
        if size_mb > 100:
            self._prepare_for_large_file()
            logger.info(f"Loading large document: {size_mb:.1f}MB - {docx_path.name}")
        
        try:
            # For files under 100MB, try standard loading first
            if size_mb < 100:
                try:
                    document = Document(docx_path)
                    logger.debug(f"Successfully loaded document: {docx_path}")
                    return document
                except Exception:
                    logger.debug(f"Standard loading failed for {size_mb:.1f}MB file, using enhanced parser")
            
            # Use enhanced parser for larger files or fallback
            return self._load_with_enhanced_parser(docx_path)
            
        except PackageNotFoundError:
            logger.error(f"Invalid DOCX file: {docx_path}")
            return None
        except Exception as e:
            logger.error(f"Failed to load {size_mb:.1f}MB document {docx_path}: {e}")
            return None

    def _load_with_enhanced_parser(self, docx_path: Path) -> Optional[Document]:
        """Enhanced parser specifically tuned for 200-300MB files."""
        try:
            from lxml import etree
            
            # Parser optimized for large file size range
            parser = etree.XMLParser(
                huge_tree=True,
                attribute_value_length_limit=200000000,  # 200M - covers max case
                text_value_length_limit=200000000,       # Handle large text blocks
                recover=True,
                strip_cdata=False,
                resolve_entities=False,
                attribute_defaults=False,
                dtd_validation=False,
                load_dtd=False,
                no_network=True,
                collect_ids=False  # Performance boost
            )
            
            from docx.opc.package import OpcPackage
            package = OpcPackage.open(docx_path, parser=parser)
            
            if package.main_document_part is None:
                logger.error(f"Main document part not found in: {docx_path}")
                return None
            
            from docx.document import Document as DocxDocument
            document = DocxDocument(package.main_document_part._element, package.main_document_part)
            
            logger.info(f"Successfully loaded large document with enhanced parser: {docx_path.name}")
            return document
            
        except Exception as e:
            logger.error(f"Enhanced parser failed: {e}")
            return None

    def _prepare_for_large_file(self):
        """Lightweight preparation for files over 100MB."""
        import gc
        import os
        
        # Force garbage collection
        gc.collect()
        
        # Set lxml environment variable
        os.environ['LXML_HUGE_TREE'] = '1'
        
        # Log memory status
        try:
            import psutil
            process = psutil.Process()
            memory_mb = process.memory_info().rss / 1024 / 1024
            logger.debug(f"Memory usage before large file: {memory_mb:.1f}MB")
        except ImportError:
            pass  # psutil not available, continue without logging
    
    def _process_pipeline(self, document: Document, processing_log: ProcessingLog):
        """
        Process document through the complete pipeline with graceful degradation.
        
        Args:
            document: Document instance
            processing_log: Processing log to update
        """
        # Process text content
        if self.text_processor:
            try:
                logger.info("Processing text content...")
                text_results = self.text_processor.process_document_text(document)
                processing_log.text_matches.extend(text_results)
                logger.info(f"Text processing completed: {len(text_results)} matches")
            except Exception as e:
                error_msg = f"Text processing failed: {e}"
                logger.error(error_msg)
                processing_log.errors.append(error_msg)
        else:
            logger.warning("Text processor not available, skipping text processing")
        
        # Process graphics content
        if self.graphics_processor:
            try:
                logger.info("Processing graphics content...")
                graphics_results = self.graphics_processor.process_graphics(document)
                processing_log.graphics_matches.extend(graphics_results)
                logger.info(f"Graphics processing completed: {len(graphics_results)} matches")
            except Exception as e:
                error_msg = f"Graphics processing failed: {e}"
                logger.error(error_msg)
                processing_log.errors.append(error_msg)
        else:
            logger.warning("Graphics processor not available, skipping graphics processing")
        
        # Process image content
        if self.image_processor:
            try:
                logger.info("Processing image content...")
                # Extract media directory from DOCX
                media_dir = self._get_media_directory(document)
                image_results = self.image_processor.process_images(document, media_dir)
                processing_log.image_matches.extend(image_results)
                
                # Collect OCR comparison data for reporting
                ocr_comparison_data = self.image_processor.get_ocr_comparison_data()
                processing_log.ocr_comparison_data.extend(ocr_comparison_data)
                
                logger.info(f"Image processing completed: {len(image_results)} matches")
                logger.info(f"OCR comparison data collected for {len(ocr_comparison_data)} images")
            except Exception as e:
                error_msg = f"Image processing failed: {e}"
                logger.error(error_msg)
                processing_log.errors.append(error_msg)
        else:
            logger.warning("Image processor not available, skipping image processing")
    
    def _get_media_directory(self, document: Document) -> Optional[Path]:
        """
        Get media directory from DOCX document.
        
        Args:
            document: Document instance
            
        Returns:
            Path to media directory or None
        """
        try:
            # Access the document's package to get media directory
            package = document.part.package
            media_parts = [part for part in package.iter_parts() if 'media' in part.partname]
            
            if media_parts:
                # Create temporary media directory
                temp_dir = PathManager.get_temp_directory() / "docx_media"
                PathManager.ensure_directory(temp_dir)
                return temp_dir
            
            return None
        except Exception as e:
            logger.error(f"Failed to get media directory: {e}")
            return None
    
    def _generate_output_path(self, input_path: Path) -> Path:
        """
        Generate output path for processed document.
        
        Args:
            input_path: Original document path
            
        Returns:
            Output path for processed document
        """
        output_dir = self.config.output_dir if self.config else Path("processed")
        PathManager.ensure_directory(output_dir)
        
        # Generate filename with _processed suffix
        stem = input_path.stem
        if stem.endswith('_processed'):
            # Avoid double suffixes
            output_filename = f"{stem}.docx"
        else:
            output_filename = f"{stem}_processed.docx"
        
        return output_dir / output_filename
    
    def _save_document(self, document: Document, output_path: Path) -> bool:
        """
        Save processed document with error handling.
        
        Args:
            document: Document instance
            output_path: Path to save document
            
        Returns:
            True if save was successful
        """
        try:
            document.save(output_path)
            logger.debug(f"Document saved to: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to save document to {output_path}: {e}")
            return False
    
    def validate_document_integrity(self, document_path: Path) -> bool:
        """
        Validate that a document can be opened without corruption.
        
        Args:
            document_path: Path to document
            
        Returns:
            True if document is valid
        """
        try:
            document = Document(document_path)
            # Try to access basic document properties
            _ = len(document.paragraphs)
            _ = len(document.tables)
            logger.debug(f"Document integrity validation passed: {document_path}")
            return True
        except Exception as e:
            logger.error(f"Document integrity validation failed for {document_path}: {e}")
            return False
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics and configuration info."""
        return {
            'config': {
                'text_mode': self.config.text_mode if self.config else 'unknown',
                'ocr_mode': self.config.ocr_mode if self.config else 'unknown',
                'use_gpu': self.config.use_gpu if self.config else False,
                'max_workers': self.config.max_workers if self.config else 1
            },
            'patterns_loaded': len(self.patterns),
            'mappings_loaded': len(self.mappings),
            'processors_available': {
                'text': self.text_processor is not None,
                'graphics': self.graphics_processor is not None,
                'image': self.image_processor is not None
            },
            'converter_available': self.converter is not None
        }
    
    def set_processors(self, text_processor=None, graphics_processor=None, image_processor=None):
        """
        Set processor instances (used when processors are implemented).
        
        Args:
            text_processor: Text processor instance
            graphics_processor: Graphics processor instance
            image_processor: Image processor instance
        """
        self.text_processor = text_processor
        self.graphics_processor = graphics_processor
        self.image_processor = image_processor
        
        logger.info("Processors updated")

def create_document_processor(config) -> DocumentProcessor:
    """
    Factory function to create a DocumentProcessor instance.
    
    Args:
        config: ProcessingConfig instance
        
    Returns:
        Configured DocumentProcessor instance
    """
    return DocumentProcessor(config)