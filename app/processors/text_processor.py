"""
Text processor for main document content processing.
Handles pattern matching, text replacement, and font preservation in document body text.
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
import logging

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..core.models import Match, create_match
from ..utils.text_utils.text_docx_utils import (
    TextReconstructor, FontManager, TextReplacer,
    create_text_replacer
)
from ..utils.pattern_matcher import PatternMatcher, create_pattern_matcher

logger = logging.getLogger(__name__)

class TextProcessor:
    """Processes main document text content with pattern matching and replacement."""
    
    def __init__(self, patterns: Dict[str, str], mappings: Dict[str, str], mode: str = "append"):
        """
        Initialize text processor.
        
        Args:
            patterns: Dictionary of pattern names to regex patterns
            mappings: Dictionary of original text to replacement text
            mode: Processing mode ('append' or 'replace')
        """
        self.patterns = patterns
        self.mappings = mappings
        self.mode = mode
        
        self.pattern_matcher = create_pattern_matcher(patterns, mappings)
        self.text_replacer = create_text_replacer(mode)
        
        logger.info(f"Text processor initialized with {len(patterns)} patterns, {len(mappings)} mappings, mode: {mode}")
    
    def process_document_text(self, document: Document) -> List[Match]:
        """
        Process all text content in the document.
        
        Args:
            document: Document to process
            
        Returns:
            List of Match objects representing successful replacements
        """
        matches = []
        processed_patterns = set()  # Track processed patterns to avoid merged cell duplicates
        
        logger.info("Starting text processing...")
        
        # Process main document body paragraphs
        body_matches = self._process_paragraphs(document.paragraphs, "body", processed_patterns)
        matches.extend(body_matches)
        
        # Process tables
        table_matches = self._process_tables(document.tables, "", processed_patterns)
        matches.extend(table_matches)
        
        # Process headers and footers (WITH TABLE SUPPORT)
        header_footer_matches = self._process_headers_footers(document, processed_patterns)
        matches.extend(header_footer_matches)
        
        logger.info(f"Text processing completed: {len(matches)} matches found")
        logger.info(f"Processed unique patterns: {len(processed_patterns)}")
        if len(processed_patterns) > 0:
            logger.info(f"Unique patterns processed: {', '.join(sorted(processed_patterns))}")
        return matches
    
    def _process_paragraphs(self, paragraphs: List[Paragraph], location: str, processed_patterns: set = None) -> List[Match]:
        """
        Process a list of paragraphs for pattern matches.
        
        Args:
            paragraphs: List of paragraphs to process
            location: Location description for logging
            
        Returns:
            List of Match objects
        """
        matches = []
        
        for i, paragraph in enumerate(paragraphs):
            try:
                paragraph_matches = self._process_paragraph(paragraph, f"{location}_paragraph_{i}", processed_patterns)
                matches.extend(paragraph_matches)
            except Exception as e:
                logger.error(f"Error processing paragraph {i} in {location}: {e}")
        
        return matches
    
    def _process_paragraph(self, paragraph: Paragraph, location: str, processed_patterns: set = None) -> List[Match]:
        """
        Process individual paragraph for pattern matches.
        
        Args:
            paragraph: Paragraph to process
            location: Location description
            
        Returns:
            List of Match objects
        """
        matches = []
        
        if not paragraph.runs:
            return matches
        
        # Reconstruct full text from runs (combines all w:t elements)
        full_text, runs = TextReconstructor.reconstruct_paragraph_text(paragraph)
        
        if not full_text.strip():
            return matches
        
        # Find pattern matches
        pattern_matches = self.pattern_matcher.find_matches(full_text)
        
        # Process each match
        for pattern_name, matched_text, start_pos, end_pos in pattern_matches:
            try:
                replacement_text = self.pattern_matcher.get_replacement(matched_text)
                if not replacement_text:
                    continue
                
                # Smart deduplication: Only prevent merged cell duplicates, allow same text in different locations
                if processed_patterns is not None:
                    # Create a location-specific key for merged cell deduplication
                    # Extract table context from location (e.g., "header_section_0_table_0")
                    location_parts = location.split('_')
                    if 'table' in location_parts:
                        # For table cells, use table-specific deduplication to prevent merged cell duplicates
                        table_context = '_'.join(location_parts[:location_parts.index('table')+2])  # Include table index
                        pattern_key = f"{table_context}:{matched_text}->{replacement_text}"
                    else:
                        # For non-table locations, use the full location to allow duplicates across different sections
                        pattern_key = f"{location}:{matched_text}->{replacement_text}"
                    
                    if pattern_key in processed_patterns:
                        logger.info(f"MERGED CELL DEDUPLICATION: Skipping duplicate pattern '{matched_text}' in same table context at {location}")
                        continue
                    processed_patterns.add(pattern_key)
                
                # Find affected runs
                text_span = TextReconstructor.find_text_in_runs(runs, matched_text, start_pos)
                if not text_span:
                    continue
                
                span_start, span_end, affected_runs = text_span
                
                # Get font information from the first affected run
                font_info = FontManager.get_font_info(affected_runs[0]) if affected_runs else {}
                
                # Perform text replacement
                success = self.text_replacer.replace_text_in_runs(
                    runs, matched_text, replacement_text, span_start, span_end
                )
                
                if success:
                    match = create_match(
                        pattern=pattern_name,
                        original=matched_text,
                        replacement=replacement_text,
                        location=location,
                        font_info={
                            'font_family': font_info.get('font_family', 'Unknown'),
                            'font_size': font_info.get('font_size', 'Unknown'),
                            'is_bold': font_info.get('is_bold', False),
                            'is_italic': font_info.get('is_italic', False)
                        }
                    )
                    matches.append(match)
                    
                    logger.info(f"Text replacement: '{matched_text}' -> '{replacement_text}' at {location}")
                
            except Exception as e:
                logger.error(f"Error processing match '{matched_text}' at {location}: {e}")
        
        return matches
    
    def _process_tables(self, tables: List[Table], location_prefix: str = "", processed_patterns: set = None) -> List[Match]:
        """
        Process tables in the document.
        
        Args:
            tables: List of tables to process
            location_prefix: Prefix for location description
            
        Returns:
            List of Match objects
        """
        matches = []
        
        for table_idx, table in enumerate(tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if location_prefix:
                            location = f"{location_prefix}_table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                        else:
                            location = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                        cell_matches = self._process_paragraphs(cell.paragraphs, location, processed_patterns)
                        matches.extend(cell_matches)
            except Exception as e:
                logger.error(f"Error processing table {table_idx} in {location_prefix}: {e}")
        
        return matches
    
    def _process_headers_footers(self, document: Document, processed_patterns: set = None) -> List[Match]:
        """
        Process headers and footers in the document, INCLUDING TABLES.
        
        Args:
            document: Document to process
            
        Returns:
            List of Match objects
        """
        matches = []
        
        try:
            # Process headers
            for section_idx, section in enumerate(document.sections):
                # Primary header
                if section.header:
                    # Process header paragraphs
                    header_matches = self._process_paragraphs(
                        section.header.paragraphs, 
                        f"header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(header_matches)
                    
                    # Process header tables - THIS IS THE KEY FIX
                    header_table_matches = self._process_tables(
                        section.header.tables,
                        f"header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(header_table_matches)
                
                # First page header
                if hasattr(section, 'first_page_header') and section.first_page_header:
                    first_header_matches = self._process_paragraphs(
                        section.first_page_header.paragraphs,
                        f"first_page_header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(first_header_matches)
                    
                    first_header_table_matches = self._process_tables(
                        section.first_page_header.tables,
                        f"first_page_header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(first_header_table_matches)
                
                # Even page header
                if hasattr(section, 'even_page_header') and section.even_page_header:
                    even_header_matches = self._process_paragraphs(
                        section.even_page_header.paragraphs,
                        f"even_page_header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(even_header_matches)
                    
                    even_header_table_matches = self._process_tables(
                        section.even_page_header.tables,
                        f"even_page_header_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(even_header_table_matches)
                
                # Primary footer
                if section.footer:
                    footer_matches = self._process_paragraphs(
                        section.footer.paragraphs,
                        f"footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(footer_matches)
                    
                    footer_table_matches = self._process_tables(
                        section.footer.tables,
                        f"footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(footer_table_matches)
                
                # First page footer
                if hasattr(section, 'first_page_footer') and section.first_page_footer:
                    first_footer_matches = self._process_paragraphs(
                        section.first_page_footer.paragraphs,
                        f"first_page_footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(first_footer_matches)
                    
                    first_footer_table_matches = self._process_tables(
                        section.first_page_footer.tables,
                        f"first_page_footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(first_footer_table_matches)
                
                # Even page footer
                if hasattr(section, 'even_page_footer') and section.even_page_footer:
                    even_footer_matches = self._process_paragraphs(
                        section.even_page_footer.paragraphs,
                        f"even_page_footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(even_footer_matches)
                    
                    even_footer_table_matches = self._process_tables(
                        section.even_page_footer.tables,
                        f"even_page_footer_section_{section_idx}",
                        processed_patterns
                    )
                    matches.extend(even_footer_table_matches)
                    
        except Exception as e:
            logger.error(f"Error processing headers/footers: {e}")
        
        return matches
    
    def get_processing_info(self) -> Dict[str, Any]:
        """Get information about the text processor configuration."""
        return {
            'mode': self.mode,
            'patterns_count': len(self.patterns),
            'mappings_count': len(self.mappings),
            'compiled_patterns': len(self.pattern_matcher.compiled_patterns)
        }

def create_text_processor(patterns: Dict[str, str], mappings: Dict[str, str], 
                         mode: str = "append") -> TextProcessor:
    """
    Factory function to create a TextProcessor instance.
    
    Args:
        patterns: Dictionary of pattern names to regex patterns
        mappings: Dictionary of original text to replacement text
        mode: Processing mode ('append' or 'replace')
        
    Returns:
        TextProcessor instance
    """
    return TextProcessor(patterns, mappings, mode)