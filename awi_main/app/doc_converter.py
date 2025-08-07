from pathlib import Path
import sys
import shutil
import subprocess
import tempfile
import os
import logging
from typing import Optional, List

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DOCConverter:
    """Enhanced DOC to DOCX converter with multiple fallback methods for cross-platform compatibility."""
    
    def __init__(self):
        self.conversion_methods = [
            self._convert_with_libreoffice,
            self._convert_with_win32com,
            self._convert_with_textract,
            self._convert_with_mammoth,
            self._convert_with_unoconv
        ]
    
    def convert_doc_to_docx(self, input_path: Path, output_path: Optional[Path] = None) -> Path:
        """
        Convert DOC file to DOCX using the first available method.
        
        Args:
            input_path: Path to input DOC file
            output_path: Optional path to output DOCX file (defaults to input_path with .docx extension)
            
        Returns:
            Path: Path to the converted DOCX file
            
        Raises:
            RuntimeError: If all conversion methods fail
        """
        input_path = Path(input_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # If already DOCX, return as-is
        if input_path.suffix.lower() == '.docx':
            return input_path
        
        # Set output path
        if output_path is None:
            output_path = input_path.with_suffix('.docx')
        else:
            output_path = Path(output_path)
        
        # If output already exists, return it
        if output_path.exists():
            logger.info(f"Output file already exists: {output_path}")
            return output_path
        
        if not input_path.suffix.lower() == '.doc':
            raise ValueError(f"Input file is not a DOC file: {input_path}")
        
        # Try each conversion method
        last_error = None
        for method in self.conversion_methods:
            try:
                logger.info(f"Trying conversion method: {method.__name__}")
                if method(str(input_path), str(output_path)):
                    logger.info(f"Successfully converted {input_path} to {output_path}")
                    return output_path
            except Exception as e:
                logger.debug(f"Method {method.__name__} failed: {e}")
                last_error = e
                continue
        
        # All methods failed
        error_msg = f"All conversion methods failed for {input_path}"
        if last_error:
            error_msg += f". Last error: {last_error}"
        raise RuntimeError(error_msg)
    
    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> bool:
        """Convert using LibreOffice command line (cross-platform)."""
        try:
            # Try different LibreOffice executable names
            libreoffice_commands = [
                'libreoffice',
                'soffice',
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
                'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe'  # Windows 32-bit
            ]
            
            for cmd in libreoffice_commands:
                if self._command_exists(cmd) or os.path.exists(cmd):
                    # Create output directory
                    output_dir = Path(output_path).parent
                    output_dir.mkdir(parents=True, exist_ok=True)
                    
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Convert to DOCX
                        result = subprocess.run([
                            cmd,
                            '--headless',
                            '--convert-to', 'docx',
                            '--outdir', temp_dir,
                            input_path
                        ], capture_output=True, text=True, timeout=60)
                        
                        if result.returncode == 0:
                            # LibreOffice creates file with same name but .docx extension
                            generated_file = Path(temp_dir) / (Path(input_path).stem + '.docx')
                            if generated_file.exists():
                                shutil.move(str(generated_file), output_path)
                                return True
                    break
            
            return False
        except Exception as e:
            logger.debug(f"LibreOffice conversion failed: {e}")
            return False
    
    def _convert_with_win32com(self, input_path: str, output_path: str) -> bool:
        """Convert using Windows COM automation (Windows only)."""
        try:
            if sys.platform != 'win32':
                return False
            
            import win32com.client
            
            # Create Word application
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # Open DOC file
                doc = word.Documents.Open(os.path.abspath(input_path))
                
                # Save as DOCX (format 16 is DOCX)
                doc.SaveAs2(os.path.abspath(output_path), FileFormat=16)
                doc.Close()
                return True
            finally:
                word.Quit()
                
        except ImportError:
            logger.debug("win32com not available")
            return False
        except Exception as e:
            logger.debug(f"win32com conversion failed: {e}")
            return False
    
    def _convert_with_textract(self, input_path: str, output_path: str) -> bool:
        """Convert using textract library."""
        try:
            import textract
            from docx import Document
            
            # Extract text from DOC
            text = textract.process(input_path).decode('utf-8')
            
            # Create new DOCX with extracted text
            doc = Document()
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            doc.save(output_path)
            return True
        except ImportError:
            logger.debug("textract not available")
            return False
        except Exception as e:
            logger.debug(f"textract conversion failed: {e}")
            return False
    
    def _convert_with_mammoth(self, input_path: str, output_path: str) -> bool:
        """Convert using mammoth library - NOTE: Mammoth only works with .docx files, not .doc files."""
        try:
            import mammoth
            from docx import Document
            
            # Mammoth only works with .docx files, not .doc files
            # This method is kept for potential future use with .docx processing
            if not input_path.lower().endswith('.docx'):
                logger.debug("mammoth: skipping .doc file (mammoth only supports .docx)")
                return False
            
            with open(input_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
            
            if text:
                doc = Document()
                for paragraph in text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                doc.save(output_path)
                logger.debug(f"mammoth: successfully converted {input_path}")
                return True
            
            logger.debug("mammoth: no text extracted")
            return False
            
        except ImportError:
            logger.debug("mammoth not available")
            return False
        except Exception as e:
            logger.debug(f"mammoth conversion failed: {e}")
            return False
    
    def _convert_with_unoconv(self, input_path: str, output_path: str) -> bool:
        """Convert using unoconv (LibreOffice backend)."""
        try:
            if not self._command_exists('unoconv'):
                return False
            
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            with tempfile.TemporaryDirectory() as temp_dir:
                result = subprocess.run([
                    'unoconv',
                    '-f', 'docx',
                    '-o', temp_dir,
                    input_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    generated_file = Path(temp_dir) / (Path(input_path).stem + '.docx')
                    if generated_file.exists():
                        shutil.move(str(generated_file), output_path)
                        return True
            
            return False
        except Exception as e:
            logger.debug(f"unoconv conversion failed: {e}")
            return False
    
    def _command_exists(self, command: str) -> bool:
        """Check if a command exists in PATH."""
        try:
            subprocess.run([command, '--version'], 
                         capture_output=True, timeout=5)
            return True
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
            return False
    
    def get_available_methods(self) -> List[str]:
        """Get list of available conversion methods on current system."""
        available = []
        
        # Check LibreOffice
        libreoffice_commands = ['libreoffice', 'soffice']
        for cmd in libreoffice_commands:
            if self._command_exists(cmd):
                available.append("LibreOffice")
                break
        
        # Check Windows COM
        if sys.platform == 'win32':
            try:
                import win32com.client
                available.append("win32com (Windows)")
            except ImportError:
                pass
        
        # Check Python libraries
        try:
            import textract
            available.append("textract")
        except ImportError:
            pass
        
        try:
            import mammoth
            available.append("mammoth")
        except ImportError:
            pass
        
        # Check unoconv
        if self._command_exists('unoconv'):
            available.append("unoconv")
        
        return available


class FileDiscovery:
    """Enhanced file discovery with batch processing and nested directory scanning."""
    
    SUPPORTED_EXTENSIONS = {'.doc', '.docx'}
    
    @staticmethod
    def discover_files(source_path: Path, recursive: bool = True, 
                      extensions: Optional[List[str]] = None) -> List[Path]:
        """
        Discover document files with enhanced filtering and nested directory support.
        
        Args:
            source_path: Path to search for files
            recursive: Whether to search subdirectories recursively
            extensions: List of file extensions to include (default: ['.doc', '.docx'])
            
        Returns:
            List[Path]: Sorted list of discovered files
            
        Raises:
            FileNotFoundError: If source_path doesn't exist
            ValueError: If no valid files found
        """
        source_path = Path(source_path)
        
        if not source_path.exists():
            raise FileNotFoundError(f"Source directory not found: {source_path}")
        
        if not source_path.is_dir():
            # If it's a single file, return it if it matches criteria
            if source_path.suffix.lower() in (extensions or FileDiscovery.SUPPORTED_EXTENSIONS):
                return [source_path]
            else:
                raise ValueError(f"File {source_path} is not a supported document type")
        
        # Set default extensions if not provided
        if extensions is None:
            extensions = list(FileDiscovery.SUPPORTED_EXTENSIONS)
        
        # Normalize extensions to lowercase
        extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' for ext in extensions]
        
        discovered_files = []
        
        try:
            if recursive:
                # Recursive search using rglob
                for ext in extensions:
                    pattern = f"**/*{ext}"
                    discovered_files.extend(source_path.rglob(pattern))
            else:
                # Non-recursive search using glob
                for ext in extensions:
                    pattern = f"*{ext}"
                    discovered_files.extend(source_path.glob(pattern))
            
            # Filter out directories and ensure files exist
            valid_files = [f for f in discovered_files if f.is_file() and f.exists()]
            
            # Remove duplicates and sort
            unique_files = list(set(valid_files))
            unique_files.sort()
            
            logger.info(f"Discovered {len(unique_files)} files in {source_path}")
            if recursive and len(unique_files) > 0:
                logger.info(f"Files found in subdirectories: {len([f for f in unique_files if f.parent != source_path])}")
            
            return unique_files
            
        except Exception as e:
            logger.error(f"Error during file discovery: {e}")
            raise
    
    @staticmethod
    def batch_discover_files(source_paths: List[Path], recursive: bool = True,
                           extensions: Optional[List[str]] = None) -> List[Path]:
        """
        Discover files from multiple source directories.
        
        Args:
            source_paths: List of paths to search
            recursive: Whether to search subdirectories recursively
            extensions: List of file extensions to include
            
        Returns:
            List[Path]: Combined sorted list of discovered files from all sources
        """
        all_files = []
        
        for source_path in source_paths:
            try:
                files = FileDiscovery.discover_files(source_path, recursive, extensions)
                all_files.extend(files)
                logger.info(f"Found {len(files)} files in {source_path}")
            except (FileNotFoundError, ValueError) as e:
                logger.warning(f"Skipping {source_path}: {e}")
                continue
        
        # Remove duplicates and sort
        unique_files = list(set(all_files))
        unique_files.sort()
        
        logger.info(f"Total unique files discovered: {len(unique_files)}")
        return unique_files
    
    @staticmethod
    def validate_file_access(file_path: Path) -> bool:
        """
        Validate that a file can be read and processed.
        
        Args:
            file_path: Path to validate
            
        Returns:
            bool: True if file is accessible, False otherwise
        """
        try:
            # Check if file exists and is readable
            if not file_path.exists():
                logger.warning(f"File does not exist: {file_path}")
                return False
            
            if not file_path.is_file():
                logger.warning(f"Path is not a file: {file_path}")
                return False
            
            # Try to read the file to check permissions
            with open(file_path, 'rb') as f:
                f.read(1)  # Read just one byte to test access
            
            # Check file size (warn about very large files)
            file_size = file_path.stat().st_size
            if file_size > 100 * 1024 * 1024:  # 100MB
                logger.warning(f"Large file detected ({file_size // (1024*1024)}MB): {file_path}")
            
            return True
            
        except PermissionError:
            logger.error(f"Permission denied accessing file: {file_path}")
            return False
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {e}")
            return False


# Global converter instance
_converter = DOCConverter()

def convert(doc_path: Path) -> Path:
    """
    Convert .doc to .docx using multiple fallback methods for cross-platform compatibility.
    Returns Path to .docx file (original if already .docx).
    
    This function maintains backward compatibility with the original API.
    """
    return _converter.convert_doc_to_docx(doc_path)


def batch_convert_files(file_paths: List[Path], preserve_structure: bool = True) -> List[Path]:
    """
    Convert multiple DOC files to DOCX with enhanced error handling.
    
    Args:
        file_paths: List of file paths to convert
        preserve_structure: Whether to maintain directory structure in output
        
    Returns:
        List[Path]: List of successfully converted DOCX file paths
    """
    converted_files = []
    failed_conversions = []
    
    logger.info(f"Starting batch conversion of {len(file_paths)} files")
    
    for file_path in file_paths:
        try:
            # Validate file access first
            if not FileDiscovery.validate_file_access(file_path):
                failed_conversions.append((file_path, "File access validation failed"))
                continue
            
            # Convert the file
            converted_path = convert(file_path)
            converted_files.append(converted_path)
            
            if file_path.suffix.lower() == '.doc':
                logger.info(f"Successfully converted: {file_path.name} -> {converted_path.name}")
            else:
                logger.debug(f"File already DOCX: {file_path.name}")
                
        except Exception as e:
            error_msg = f"Conversion failed: {str(e)}"
            failed_conversions.append((file_path, error_msg))
            logger.error(f"Failed to convert {file_path}: {error_msg}")
            continue
    
    # Log summary
    logger.info(f"Batch conversion complete: {len(converted_files)} successful, {len(failed_conversions)} failed")
    
    if failed_conversions:
        logger.warning("Failed conversions:")
        for file_path, error in failed_conversions:
            logger.warning(f"  - {file_path}: {error}")
    
    return converted_files


def discover_and_convert_files(source_path: Path, recursive: bool = True, 
                              extensions: Optional[List[str]] = None) -> List[Path]:
    """
    Discover and convert document files in one operation.
    
    Args:
        source_path: Path to search for files
        recursive: Whether to search subdirectories recursively
        extensions: List of file extensions to include
        
    Returns:
        List[Path]: List of converted DOCX file paths
        
    Raises:
        FileNotFoundError: If source_path doesn't exist
        RuntimeError: If no files could be processed
    """
    logger.info(f"Discovering files in: {source_path}")
    
    # Discover files
    try:
        discovered_files = FileDiscovery.discover_files(source_path, recursive, extensions)
    except (FileNotFoundError, ValueError) as e:
        logger.error(f"File discovery failed: {e}")
        raise
    
    if not discovered_files:
        raise RuntimeError(f"No supported document files found in {source_path}")
    
    logger.info(f"Found {len(discovered_files)} files to process")
    
    # Convert files
    converted_files = batch_convert_files(discovered_files)
    
    if not converted_files:
        raise RuntimeError("No files could be successfully converted")
    
    return converted_files


def get_available_conversion_methods() -> List[str]:
    """Get list of available conversion methods on current system."""
    return _converter.get_available_methods()


if __name__ == "__main__":
    # Command line interface
    if len(sys.argv) < 2:
        print("Usage: python doc_converter.py <input.doc> [output.docx]")
        print("Example: python doc_converter.py document.doc")
        print("Example: python doc_converter.py document.doc converted.docx")
        print()
        available_methods = get_available_conversion_methods()
        if available_methods:
            print(f"Available conversion methods: {', '.join(available_methods)}")
        else:
            print("No conversion methods available. Please install dependencies:")
            print("- LibreOffice: https://www.libreoffice.org/")
            print("- Python libraries: pip install textract python-docx mammoth")
            if sys.platform == 'win32':
                print("- Windows COM: pip install pywin32")
        sys.exit(1)
    
    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2]) if len(sys.argv) > 2 else None
    
    try:
        result = convert(input_file) if output_file is None else _converter.convert_doc_to_docx(input_file, output_file)
        print(f"Successfully converted to: {result}")
    except Exception as e:
        print(f"Conversion failed: {e}")
        sys.exit(1)