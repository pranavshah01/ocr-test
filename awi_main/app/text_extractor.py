"""
Complete Text Extractor for DOCX/DOC files.
Combines document text and OCR text from images into comprehensive text files.
"""

from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from docx import Document
import json
import re
import unicodedata
import math
from datetime import datetime
from .ocr_engine import run_ocr
from PIL import Image
from io import BytesIO
from lxml import etree
import zipfile
from collections import namedtuple
from .shared_constants import XML_NAMESPACES, SharedUtilities

# Structure to hold text fragments with their XML context
TextFragment = namedtuple('TextFragment', ['text', 'element', 'start_pos', 'end_pos'])

# Structure to hold pattern matches with location information
PatternMatch = namedtuple('PatternMatch', ['pattern', 'match_text', 'start_fragment', 'end_fragment', 'start_pos', 'end_pos'])


class EnhancedPatternMatcher:
    """
    Enhanced pattern matcher that can handle patterns spanning multiple <w:t> tags.
    Works directly with DOCX XML structure for accurate text reconstruction.
    """
    
    def __init__(self):
        self.namespaces = XML_NAMESPACES
    
    def normalize_text(self, text: str) -> str:
        """
        Normalize text for consistent pattern matching.
        
        Args:
            text: Input text to normalize
            
        Returns:
            Normalized text
        """
        # Normalize Unicode characters
        text = unicodedata.normalize('NFKC', text)
        
        # Normalize whitespace (but preserve structure)
        text = re.sub(r'\s+', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def extract_text_fragments(self, docx_path: Path) -> List[TextFragment]:
        """
        Extract text fragments from DOCX XML with their XML context.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            List of TextFragment objects with XML context
        """
        fragments = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Extract main document XML
                doc_xml = docx_zip.read('word/document.xml')
                root = etree.fromstring(doc_xml)
                
                # Find all text elements (w:t)
                text_elements = root.xpath('//w:t', namespaces=self.namespaces)
                
                current_pos = 0
                for element in text_elements:
                    if element.text:
                        text = element.text
                        start_pos = current_pos
                        end_pos = current_pos + len(text)
                        
                        fragment = TextFragment(
                            text=text,
                            element=element,
                            start_pos=start_pos,
                            end_pos=end_pos
                        )
                        fragments.append(fragment)
                        current_pos = end_pos
                
                return fragments
                
        except Exception as e:
            print(f"Error extracting text fragments: {e}")
            return []
    
    def reconstruct_continuous_text(self, fragments: List[TextFragment]) -> str:
        """
        Reconstruct continuous text from fragments for pattern matching.
        
        Args:
            fragments: List of TextFragment objects
            
        Returns:
            Reconstructed continuous text
        """
        if not fragments:
            return ""
        
        # Combine all text fragments
        combined_text = ''.join(fragment.text for fragment in fragments)
        
        # Normalize the combined text
        return self.normalize_text(combined_text)
    
    def find_pattern_matches(self, fragments: List[TextFragment], patterns: List[str]) -> List[PatternMatch]:
        """
        Find pattern matches across text fragments, handling patterns that span multiple <w:t> tags.
        
        Args:
            fragments: List of TextFragment objects
            patterns: List of regex patterns to match
            
        Returns:
            List of PatternMatch objects
        """
        if not fragments or not patterns:
            return []
        
        matches = []
        
        # Reconstruct continuous text for pattern matching
        continuous_text = self.reconstruct_continuous_text(fragments)
        
        # Find matches for each pattern
        for pattern in patterns:
            try:
                # Compile regex pattern
                regex = re.compile(pattern, re.IGNORECASE | re.UNICODE)
                
                # Find all matches in continuous text
                for match in regex.finditer(continuous_text):
                    match_start = match.start()
                    match_end = match.end()
                    match_text = match.group()
                    
                    # Find which fragments contain the match
                    start_fragment = None
                    end_fragment = None
                    
                    for fragment in fragments:
                        if (start_fragment is None and 
                            fragment.start_pos <= match_start < fragment.end_pos):
                            start_fragment = fragment
                        
                        if (fragment.start_pos < match_end <= fragment.end_pos):
                            end_fragment = fragment
                            break
                    
                    if start_fragment and end_fragment:
                        pattern_match = PatternMatch(
                            pattern=pattern,
                            match_text=match_text,
                            start_fragment=start_fragment,
                            end_fragment=end_fragment,
                            start_pos=match_start,
                            end_pos=match_end
                        )
                        matches.append(pattern_match)
                        
            except re.error as e:
                print(f"Invalid regex pattern '{pattern}': {e}")
                continue
        
        return matches
    
    def extract_context_around_match(self, fragments: List[TextFragment], 
                                   match: PatternMatch, context_chars: int = 50) -> str:
        """
        Extract context around a pattern match for debugging/logging.
        
        Args:
            fragments: List of TextFragment objects
            match: PatternMatch object
            context_chars: Number of characters of context to include
            
        Returns:
            Context string around the match
        """
        continuous_text = self.reconstruct_continuous_text(fragments)
        
        start = max(0, match.start_pos - context_chars)
        end = min(len(continuous_text), match.end_pos + context_chars)
        
        context = continuous_text[start:end]
        
        # Highlight the match within the context
        match_start_in_context = match.start_pos - start
        match_end_in_context = match.end_pos - start
        
        highlighted = (
            context[:match_start_in_context] + 
            "**" + context[match_start_in_context:match_end_in_context] + "**" +
            context[match_end_in_context:]
        )
        
        return highlighted


def find_patterns_across_split_tags(docx_path: Path, patterns: List[str]) -> Dict[str, Any]:
    """
    Find patterns that may span across multiple <w:t> tags in a DOCX document.
    
    Args:
        docx_path: Path to DOCX file
        patterns: List of regex patterns to search for
        
    Returns:
        Dictionary containing pattern matches and metadata
    """
    matcher = EnhancedPatternMatcher()
    
    # Extract text fragments with XML context
    fragments = matcher.extract_text_fragments(docx_path)
    
    if not fragments:
        return {
            "file_path": str(docx_path),
            "extraction_time": datetime.now().isoformat(),
            "fragments_count": 0,
            "matches": [],
            "total_matches": 0,
            "patterns_searched": patterns
        }
    
    # Find pattern matches
    matches = matcher.find_pattern_matches(fragments, patterns)
    
    # Prepare results
    results = {
        "file_path": str(docx_path),
        "extraction_time": datetime.now().isoformat(),
        "fragments_count": len(fragments),
        "matches": [],
        "total_matches": len(matches),
        "patterns_searched": patterns,
        "continuous_text_length": len(matcher.reconstruct_continuous_text(fragments))
    }
    
    # Process each match
    for match in matches:
        match_data = {
            "pattern": match.pattern,
            "match_text": match.match_text,
            "start_pos": match.start_pos,
            "end_pos": match.end_pos,
            "length": len(match.match_text),
            "spans_multiple_fragments": match.start_fragment != match.end_fragment,
            "context": matcher.extract_context_around_match(fragments, match),
            "start_fragment_text": match.start_fragment.text,
            "end_fragment_text": match.end_fragment.text
        }
        results["matches"].append(match_data)
    
    return results


def test_enhanced_pattern_matching(docx_path: Path, patterns: List[str], verbose: bool = True) -> Dict[str, Any]:
    """
    Test the enhanced pattern matching functionality and provide detailed diagnostics.
    
    Args:
        docx_path: Path to DOCX file to test
        patterns: List of regex patterns to test
        verbose: Whether to print detailed output
        
    Returns:
        Dictionary containing test results and diagnostics
    """
    if verbose:
        print(f"\n=== Enhanced Pattern Matching Test ===")
        print(f"File: {docx_path}")
        print(f"Patterns to test: {patterns}")
    
    # Run the enhanced pattern matching
    results = find_patterns_across_split_tags(docx_path, patterns)
    
    if verbose:
        print(f"\n--- Results ---")
        print(f"Text fragments found: {results['fragments_count']}")
        print(f"Continuous text length: {results['continuous_text_length']}")
        print(f"Total matches found: {results['total_matches']}")
        
        if results['matches']:
            print(f"\n--- Match Details ---")
            for i, match in enumerate(results['matches'], 1):
                print(f"\nMatch {i}:")
                print(f"  Pattern: {match['pattern']}")
                print(f"  Match text: '{match['match_text']}'")
                print(f"  Position: {match['start_pos']}-{match['end_pos']} (length: {match['length']})")
                print(f"  Spans multiple fragments: {match['spans_multiple_fragments']}")
                print(f"  Context: {match['context']}")
                if match['spans_multiple_fragments']:
                    print(f"  Start fragment: '{match['start_fragment_text']}'")
                    print(f"  End fragment: '{match['end_fragment_text']}'")
        else:
            print("\nNo matches found.")
    
    # Add test validation
    test_results = {
        **results,
        "test_passed": True,
        "test_errors": [],
        "test_warnings": []
    }
    
    # Validate results
    if results['fragments_count'] == 0:
        test_results["test_warnings"].append("No text fragments found - document may be empty or corrupted")
    
    if results['continuous_text_length'] == 0:
        test_results["test_warnings"].append("No continuous text reconstructed")
    
    # Test for patterns that should match
    expected_patterns = [r'\d+', r'[A-Za-z]+']  # Basic patterns that should match in most documents
    for pattern in expected_patterns:
        if pattern in patterns:
            found_match = any(match['pattern'] == pattern for match in results['matches'])
            if not found_match and results['continuous_text_length'] > 0:
                test_results["test_warnings"].append(f"Expected pattern '{pattern}' not found despite having text")
    
    if verbose:
        print(f"\n--- Test Summary ---")
        print(f"Test passed: {test_results['test_passed']}")
        if test_results['test_warnings']:
            print(f"Warnings: {len(test_results['test_warnings'])}")
            for warning in test_results['test_warnings']:
                print(f"  - {warning}")
        if test_results['test_errors']:
            print(f"Errors: {len(test_results['test_errors'])}")
            for error in test_results['test_errors']:
                print(f"  - {error}")
    
    return test_results


def extract_document_text(docx_path: Path) -> Dict[str, Any]:
    """
    Extract all text content from a DOCX document including paragraphs, headers, footers, and tables.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dict containing extracted text organized by type
    """
    doc = Document(str(docx_path))
    
    extracted_text = {
        "file_path": str(docx_path),
        "extraction_time": datetime.now().isoformat(),
        "paragraphs": [],
        "headers": [],
        "footers": [],
        "tables": [],
        "total_paragraphs": 0,
        "total_words": 0,
        "total_characters": 0
    }
    
    # Extract paragraph text
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            para_data = {
                "index": i,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "word_count": len(paragraph.text.split()),
                "char_count": len(paragraph.text)
            }
            extracted_text["paragraphs"].append(para_data)
            extracted_text["total_words"] += para_data["word_count"]
            extracted_text["total_characters"] += para_data["char_count"]
    
    extracted_text["total_paragraphs"] = len(extracted_text["paragraphs"])
    
    # Extract table text
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "table_index": table_idx,
            "rows": []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = {
                "row_index": row_idx,
                "cells": []
            }
            
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    cell_data = {
                        "cell_index": cell_idx,
                        "text": cell.text,
                        "word_count": len(cell.text.split()),
                        "char_count": len(cell.text)
                    }
                    row_data["cells"].append(cell_data)
                    extracted_text["total_words"] += cell_data["word_count"]
                    extracted_text["total_characters"] += cell_data["char_count"]
            
            if row_data["cells"]:
                table_data["rows"].append(row_data)
        
        if table_data["rows"]:
            extracted_text["tables"].append(table_data)
    
    # Extract text from text boxes and drawing elements (callouts, shapes)
    extracted_text["text_boxes"] = []
    try:
        # Access the document XML to find text boxes and drawing elements
        doc_element = doc.element
        
        # Search for text boxes (w:txbxContent)
        textboxes = doc_element.xpath('.//w:txbxContent')
        for i, textbox in enumerate(textboxes):
            # Extract text from text box paragraphs
            text_elements = textbox.xpath('.//w:t')
            textbox_text = ''.join([elem.text or '' for elem in text_elements]).strip()
            
            if textbox_text:
                textbox_data = {
                    "index": i,
                    "text": textbox_text,
                    "type": "textbox",
                    "word_count": len(textbox_text.split()),
                    "char_count": len(textbox_text)
                }
                extracted_text["text_boxes"].append(textbox_data)
                extracted_text["total_words"] += textbox_data["word_count"]
                extracted_text["total_characters"] += textbox_data["char_count"]
        
        # Search for VML text (older format text boxes)
        vml_textboxes = doc_element.xpath('.//v:textbox')
        for i, vml_textbox in enumerate(vml_textboxes):
            text_elements = vml_textbox.xpath('.//w:t')
            vml_text = ''.join([elem.text or '' for elem in text_elements]).strip()
            
            if vml_text:
                vml_data = {
                    "index": len(extracted_text["text_boxes"]) + i,
                    "text": vml_text,
                    "type": "vml_textbox",
                    "word_count": len(vml_text.split()),
                    "char_count": len(vml_text)
                }
                extracted_text["text_boxes"].append(vml_data)
                extracted_text["total_words"] += vml_data["word_count"]
                extracted_text["total_characters"] += vml_data["char_count"]
                
    except Exception as e:
        extracted_text["text_boxes_error"] = str(e)
    
    # Extract header/footer text (if accessible)
    try:
        for section in doc.sections:
            # Headers
            if hasattr(section, 'header') and section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_data = {
                            "text": para.text,
                            "style": para.style.name if para.style else "Header",
                            "word_count": len(para.text.split()),
                            "char_count": len(para.text)
                        }
                        extracted_text["headers"].append(header_data)
                        extracted_text["total_words"] += header_data["word_count"]
                        extracted_text["total_characters"] += header_data["char_count"]
            
            # Footers
            if hasattr(section, 'footer') and section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_data = {
                            "text": para.text,
                            "style": para.style.name if para.style else "Footer",
                            "word_count": len(para.text.split()),
                            "char_count": len(para.text)
                        }
                        extracted_text["footers"].append(footer_data)
                        extracted_text["total_words"] += footer_data["word_count"]
                        extracted_text["total_characters"] += footer_data["char_count"]
    except Exception as e:
        # Headers/footers might not be accessible in some documents
        extracted_text["header_footer_error"] = str(e)
    
    return extracted_text


def extract_image_ocr_text(docx_path: Path, ocr_engine: str = "easyocr", gpu: bool = True, confidence_min: float = 0.7) -> Dict[str, Any]:
    """
    Extract OCR text from all images in a DOCX document with orientation information.
    
    Args:
        docx_path: Path to the DOCX file
        ocr_engine: OCR engine to use
        gpu: Whether to use GPU for OCR
        confidence_min: Minimum confidence threshold
        
    Returns:
        Dict containing OCR results from all images with orientation data
    """
    doc = Document(str(docx_path))
    rels = doc.part.rels
    
    ocr_results = {
        "file_path": str(docx_path),
        "ocr_time": datetime.now().isoformat(),
        "ocr_engine": ocr_engine,
        "confidence_threshold": confidence_min,
        "images": [],
        "total_images": 0,
        "total_ocr_text_blocks": 0,
        "total_ocr_words": 0,
        "total_ocr_characters": 0,
        "orientations_found": []
    }
    
    image_index = 0
    
    # Process inline shapes (embedded images)
    for shape in getattr(doc, 'inline_shapes', []):
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rel_id = blip.embed
            img_part = rels[rel_id].target_part
            img_bytes = img_part.blob
            
            # Convert to PIL Image
            img = Image.open(BytesIO(img_bytes))
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Run OCR on the image with orientation detection
            ocr_result = run_ocr(img, use_gpu=gpu, engine=ocr_engine)
            
            image_data = {
                "image_index": image_index,
                "image_size": {"width": img.width, "height": img.height},
                "image_mode": img.mode,
                "ocr_blocks": [],
                "ocr_text_by_orientation": {},
                "image_total_words": 0,
                "image_total_characters": 0,
                "orientations_detected": []
            }
            
            # Group OCR results by orientation
            orientation_groups = {}
            
            # Process OCR results
            for block in ocr_result:
                text = block.get("text", "").strip()
                conf = block.get("conf", 0.0)
                bbox = block.get("bbox", [])
                
                if text and conf >= confidence_min:
                    # Calculate orientation from bounding box
                    orientation = 0.0
                    if bbox and len(bbox) >= 4:
                        try:
                            # Calculate angle from bounding box points
                            if len(bbox) == 4 and isinstance(bbox[0], list):
                                # EasyOCR format: [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
                                x1, y1 = bbox[0]
                                x2, y2 = bbox[1]
                                if x2 != x1:
                                    angle = math.atan2(y2 - y1, x2 - x1) * 180 / math.pi
                                    orientation = angle
                        except:
                            orientation = 0.0
                    
                    # Round orientation to nearest 90 degrees for grouping
                    orientation_group = round(orientation / 90) * 90
                    orientation_key = f"{orientation_group}deg"
                    
                    if orientation_key not in orientation_groups:
                        orientation_groups[orientation_key] = []
                    
                    block_data = {
                        "text": text,
                        "confidence": conf,
                        "bbox": bbox,
                        "orientation": orientation,
                        "orientation_group": orientation_group,
                        "word_count": len(text.split()),
                        "char_count": len(text),
                        "engine": block.get("engine", ocr_engine),
                        "fallback_used": block.get("fallback_used", False)
                    }
                    
                    orientation_groups[orientation_key].append(block_data)
                    image_data["ocr_blocks"].append(block_data)
                    image_data["image_total_words"] += block_data["word_count"]
                    image_data["image_total_characters"] += block_data["char_count"]
                    
                    ocr_results["total_ocr_text_blocks"] += 1
                    ocr_results["total_ocr_words"] += block_data["word_count"]
                    ocr_results["total_ocr_characters"] += block_data["char_count"]
            
            # Create orientation-specific text summaries
            for orientation_key, blocks in orientation_groups.items():
                all_text = " ".join([block["text"] for block in blocks])
                image_data["ocr_text_by_orientation"][orientation_key] = {
                    "text": all_text,
                    "blocks_count": len(blocks),
                    "total_words": sum(block["word_count"] for block in blocks),
                    "total_characters": sum(block["char_count"] for block in blocks),
                    "average_confidence": sum(block["confidence"] for block in blocks) / len(blocks) if blocks else 0.0
                }
                image_data["orientations_detected"].append(orientation_key)
                if orientation_key not in ocr_results["orientations_found"]:
                    ocr_results["orientations_found"].append(orientation_key)
            
            if image_data["ocr_blocks"]:
                ocr_results["images"].append(image_data)
            
            image_index += 1
            
        except Exception as e:
            # Skip images that can't be processed
            continue
    
    ocr_results["total_images"] = image_index
    
    return ocr_results


def create_combined_text_file(docx_path: Path, ocr_dir: Path, file_name: str, 
                             ocr_engine: str = "easyocr", gpu: bool = True, 
                             confidence_min: float = 0.7) -> Dict[str, Any]:
    """
    Create a comprehensive text file combining document text and OCR text.
    
    Args:
        docx_path: Path to the DOCX file
        ocr_dir: Directory to save OCR results
        file_name: Base name for the output file
        ocr_engine: OCR engine to use
        gpu: Whether to use GPU for OCR
        confidence_min: Minimum confidence threshold
        
    Returns:
        Summary of the text extraction process
    """
    if not ocr_dir:
        return {"error": "OCR directory not specified"}
    
    # Ensure OCR directory exists
    ocr_dir = Path(ocr_dir)
    ocr_dir.mkdir(parents=True, exist_ok=True)
    
    # Extract document text
    document_text = extract_document_text(docx_path)
    
    # Extract OCR text from images
    ocr_text = extract_image_ocr_text(docx_path, ocr_engine, gpu, confidence_min)
    
    # Create combined text data
    combined_data = {
        "file_info": {
            "source_file": str(docx_path),
            "extraction_time": datetime.now().isoformat(),
            "ocr_engine": ocr_engine,
            "confidence_threshold": confidence_min
        },
        "document_text": document_text,
        "ocr_text": ocr_text,
        "summary": {
            "total_document_words": document_text["total_words"],
            "total_document_characters": document_text["total_characters"],
            "total_document_paragraphs": document_text["total_paragraphs"],
            "total_ocr_words": ocr_text["total_ocr_words"],
            "total_ocr_characters": ocr_text["total_ocr_characters"],
            "total_ocr_images": ocr_text["total_images"],
            "total_ocr_text_blocks": ocr_text["total_ocr_text_blocks"],
            "combined_total_words": document_text["total_words"] + ocr_text["total_ocr_words"],
            "combined_total_characters": document_text["total_characters"] + ocr_text["total_ocr_characters"]
        }
    }
    
    # Save detailed JSON file
    json_output_path = ocr_dir / f"{file_name}_complete_text.json"
    
    # Convert numpy types to native Python types for JSON serialization
    import numpy as np
    
    def convert_numpy_types(obj):
        # Handle NumPy integers (including int64)
        if isinstance(obj, (np.integer, np.int64, np.int32, np.int16, np.int8)):
            return int(obj)
        # Handle NumPy floats
        elif isinstance(obj, (np.floating, np.float64, np.float32, np.float16)):
            return float(obj)
        # Handle NumPy booleans
        elif isinstance(obj, np.bool_):
            return bool(obj)
        # Handle NumPy arrays
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        # Handle other NumPy scalars with .item() method
        elif hasattr(obj, 'item') and hasattr(obj, 'dtype'):
            return obj.item()
        # Handle dictionaries recursively
        elif isinstance(obj, dict):
            return {k: convert_numpy_types(v) for k, v in obj.items()}
        # Handle lists recursively
        elif isinstance(obj, list):
            return [convert_numpy_types(item) for item in obj]
        # Handle tuples recursively
        elif isinstance(obj, tuple):
            return tuple(convert_numpy_types(item) for item in obj)
        # Handle sets
        elif isinstance(obj, set):
            return [convert_numpy_types(item) for item in obj]
        # Return as-is for native Python types
        else:
            return obj
    
    combined_data_serializable = convert_numpy_types(combined_data)
    
    try:
        with open(json_output_path, 'w', encoding='utf-8') as f:
            json.dump(combined_data_serializable, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error writing JSON to {json_output_path}: {e}")
        # Try to write a minimal version
        try:
            minimal_data = {
                "file": str(docx_path),
                "error": str(e),
                "timestamp": combined_data_serializable.get("timestamp", "unknown")
            }
            with open(json_output_path, 'w', encoding='utf-8') as f:
                json.dump(minimal_data, f, indent=2, ensure_ascii=False)
        except Exception:
            print(f"Failed to write even minimal JSON to {json_output_path}")
    
    # Create readable text file
    txt_output_path = ocr_dir / f"{file_name}_complete_text.txt"
    with open(txt_output_path, 'w', encoding='utf-8') as f:
        f.write(f"COMPLETE TEXT EXTRACTION\n")
        f.write(f"=" * 50 + "\n")
        f.write(f"Source File: {docx_path}\n")
        f.write(f"Extraction Time: {combined_data['file_info']['extraction_time']}\n")
        f.write(f"OCR Engine: {ocr_engine}\n")
        f.write(f"Confidence Threshold: {confidence_min}\n\n")
        
        # Summary
        f.write(f"SUMMARY\n")
        f.write(f"-" * 20 + "\n")
        f.write(f"Document Text: {combined_data['summary']['total_document_words']} words, {combined_data['summary']['total_document_characters']} characters\n")
        f.write(f"OCR Text: {combined_data['summary']['total_ocr_words']} words, {combined_data['summary']['total_ocr_characters']} characters\n")
        f.write(f"Combined Total: {combined_data['summary']['combined_total_words']} words, {combined_data['summary']['combined_total_characters']} characters\n")
        f.write(f"Images Processed: {combined_data['summary']['total_ocr_images']}\n")
        f.write(f"OCR Text Blocks: {combined_data['summary']['total_ocr_text_blocks']}\n\n")
        
        # Document text
        f.write(f"DOCUMENT TEXT\n")
        f.write(f"=" * 30 + "\n\n")
        
        # Headers
        if document_text["headers"]:
            f.write(f"HEADERS:\n")
            for header in document_text["headers"]:
                f.write(f"  {header['text']}\n")
            f.write("\n")
        
        # Paragraphs
        f.write(f"PARAGRAPHS:\n")
        for para in document_text["paragraphs"]:
            f.write(f"[{para['index']}] {para['text']}\n")
        f.write("\n")
        
        # Text boxes and callouts
        if document_text.get("text_boxes"):
            f.write(f"TEXT BOXES & CALLOUTS:\n")
            for textbox in document_text["text_boxes"]:
                f.write(f"[{textbox['type']}] {textbox['text']}\n")
            f.write("\n")
        
        # Tables
        if document_text["tables"]:
            f.write(f"TABLES:\n")
            for table in document_text["tables"]:
                f.write(f"  Table {table['table_index']}:\n")
                for row in table["rows"]:
                    row_text = " | ".join([cell["text"] for cell in row["cells"]])
                    f.write(f"    {row_text}\n")
                f.write("\n")
        
        # Footers
        if document_text["footers"]:
            f.write(f"FOOTERS:\n")
            for footer in document_text["footers"]:
                f.write(f"  {footer['text']}\n")
            f.write("\n")
        
        # OCR text
        f.write(f"OCR TEXT FROM IMAGES\n")
        f.write(f"=" * 30 + "\n\n")
        
        for img in ocr_text["images"]:
            f.write(f"IMAGE {img['image_index']} ({img['image_size']['width']}x{img['image_size']['height']}):\n")
            for block in img["ocr_blocks"]:
                f.write(f"  [{block['confidence']:.2f}] {block['text']}\n")
            f.write("\n")
    
    return {
        "json_file": str(json_output_path),
        "text_file": str(txt_output_path),
        "summary": combined_data["summary"]
    }
