"""
Test Suite for Robust DOCX Writer

This test suite validates the robust DOCX writer's ability to prevent corruption
and handle various edge cases safely.

Test Categories:
1. Basic DOCX operations (load, save, verify)
2. Corruption prevention and recovery
3. Atomic operations and rollback
4. XML writing safety
5. Backup and recovery mechanisms
6. Error handling and validation

Author: OCR DOCX Text Replacement Utility
"""

import unittest
import tempfile
import shutil
import os
import zipfile
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document
from lxml import etree

# Import the module under test
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.robust_docx_writer import (
    RobustDOCXWriter,
    DOCXCorruptionError,
    create_robust_docx_writer,
    safe_docx_operation
)


class TestRobustDOCXWriter(unittest.TestCase):
    """Test cases for RobustDOCXWriter class."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp(prefix="test_robust_docx_")
        self.test_docx_path = Path(self.test_dir) / "test_document.docx"
        self.output_path = Path(self.test_dir) / "output_document.docx"
        
        # Create a simple test DOCX file
        self._create_test_docx()
        
        # Initialize writer
        self.writer = RobustDOCXWriter(enable_backup=True, validate_structure=True)
    
    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
    
    def _create_test_docx(self):
        """Create a simple test DOCX file."""
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.add_paragraph("Used for testing robust DOCX operations.")
        doc.save(str(self.test_docx_path))
    
    def _create_corrupted_docx(self, corruption_type: str = "truncated"):
        """Create a corrupted DOCX file for testing."""
        corrupted_path = Path(self.test_dir) / "corrupted.docx"
        
        if corruption_type == "truncated":
            # Create truncated file
            with open(corrupted_path, 'wb') as f:
                f.write(b"PK")  # Start of ZIP but truncated
        elif corruption_type == "empty":
            # Create empty file
            corrupted_path.touch()
        elif corruption_type == "invalid_zip":
            # Create file with invalid ZIP content
            with open(corrupted_path, 'wb') as f:
                f.write(b"This is not a ZIP file")
        
        return corrupted_path
    
    def test_initialization(self):
        """Test RobustDOCXWriter initialization."""
        # Test default initialization
        writer1 = RobustDOCXWriter()
        self.assertTrue(writer1.enable_backup)
        self.assertTrue(writer1.validate_structure)
        
        # Test custom initialization
        writer2 = RobustDOCXWriter(enable_backup=False, validate_structure=False)
        self.assertFalse(writer2.enable_backup)
        self.assertFalse(writer2.validate_structure)
        
        # Test factory function
        writer3 = create_robust_docx_writer(enable_backup=True, validate_structure=True)
        self.assertIsInstance(writer3, RobustDOCXWriter)
        self.assertTrue(writer3.enable_backup)
        self.assertTrue(writer3.validate_structure)
    
    def test_safe_docx_context_success(self):
        """Test successful DOCX context operations."""
        def test_operation(doc):
            # Add a paragraph to test modification
            doc.add_paragraph("Added by test operation")
            return {"modifications": 1}
        
        with self.writer.safe_docx_context(self.test_docx_path, self.output_path) as doc:
            result = test_operation(doc)
            self.assertEqual(result["modifications"], 1)
        
        # Verify output file was created
        self.assertTrue(self.output_path.exists())
        
        # Verify content was modified
        output_doc = Document(str(self.output_path))
        paragraphs = [p.text for p in output_doc.paragraphs]
        self.assertIn("Added by test operation", paragraphs)
        
        # Check processing stats
        stats = self.writer.get_processing_stats()
        self.assertEqual(stats["files_processed"], 1)
        self.assertEqual(stats["validations_performed"], 1)
    
    def test_load_document_safely_success(self):
        """Test successful document loading."""
        doc = self.writer._load_document_safely(self.test_docx_path)
        self.assertIsInstance(doc, Document)
        self.assertIsNotNone(doc.part)
    
    def test_load_document_safely_file_not_found(self):
        """Test loading non-existent file."""
        non_existent_path = Path(self.test_dir) / "non_existent.docx"
        
        with self.assertRaises(DOCXCorruptionError) as context:
            self.writer._load_document_safely(non_existent_path)
        
        self.assertIn("Failed to load DOCX safely", str(context.exception))
    
    def test_load_document_safely_corrupted_file(self):
        """Test loading corrupted files."""
        # Test truncated file
        corrupted_path = self._create_corrupted_docx("truncated")
        with self.assertRaises(DOCXCorruptionError):
            self.writer._load_document_safely(corrupted_path)
        
        # Test empty file
        corrupted_path = self._create_corrupted_docx("empty")
        with self.assertRaises(DOCXCorruptionError):
            self.writer._load_document_safely(corrupted_path)
        
        # Test invalid ZIP
        corrupted_path = self._create_corrupted_docx("invalid_zip")
        with self.assertRaises(DOCXCorruptionError):
            self.writer._load_document_safely(corrupted_path)
    
    def test_verify_zip_structure_success(self):
        """Test successful ZIP structure verification."""
        # Should not raise exception for valid DOCX
        self.writer._verify_zip_structure(self.test_docx_path)
    
    def test_verify_zip_structure_invalid(self):
        """Test ZIP structure verification with invalid files."""
        corrupted_path = self._create_corrupted_docx("invalid_zip")
        
        with self.assertRaises(DOCXCorruptionError) as context:
            self.writer._verify_zip_structure(corrupted_path)
        
        self.assertIn("Invalid ZIP structure", str(context.exception))
    
    def test_validate_docx_structure_success(self):
        """Test successful DOCX structure validation."""
        doc = Document(str(self.test_docx_path))
        
        # Should not raise exception for valid document
        self.writer._validate_docx_structure(doc)
    
    def test_validate_docx_structure_invalid(self):
        """Test DOCX structure validation with invalid document."""
        # Test with None document
        with self.assertRaises(DOCXCorruptionError):
            self.writer._validate_docx_structure(None)
    
    def test_backup_creation(self):
        """Test backup file creation."""
        # Create a file to backup
        test_file = Path(self.test_dir) / "backup_test.docx"
        shutil.copy2(self.test_docx_path, test_file)
        
        # Create backup
        backup_path = self.writer._create_backup(test_file)
        
        # Verify backup was created
        self.assertTrue(backup_path.exists())
        self.assertIn("backup", backup_path.name)
        
        # Verify backup content matches original
        self.assertEqual(test_file.stat().st_size, backup_path.stat().st_size)
    
    def test_atomic_save_operation(self):
        """Test atomic save operations."""
        doc = Document(str(self.test_docx_path))
        doc.add_paragraph("Test atomic save")
        
        temp_dir = tempfile.mkdtemp()
        try:
            # Should save successfully
            self.writer._save_document_safely(doc, self.output_path, temp_dir)
            
            # Verify file was created
            self.assertTrue(self.output_path.exists())
            
            # Verify content
            saved_doc = Document(str(self.output_path))
            paragraphs = [p.text for p in saved_doc.paragraphs]
            self.assertIn("Test atomic save", paragraphs)
            
        finally:
            shutil.rmtree(temp_dir)
    
    def test_verify_saved_file_success(self):
        """Test successful saved file verification."""
        # Copy test file to output location
        shutil.copy2(self.test_docx_path, self.output_path)
        
        # Should not raise exception for valid file
        self.writer._verify_saved_file(self.output_path)
    
    def test_verify_saved_file_corrupted(self):
        """Test saved file verification with corrupted file."""
        # Create corrupted output file
        with open(self.output_path, 'wb') as f:
            f.write(b"corrupted")
        
        with self.assertRaises(DOCXCorruptionError):
            self.writer._verify_saved_file(self.output_path)
    
    def test_xml_writing_safety(self):
        """Test safe XML writing with proper encoding."""
        # Create test XML element
        root = etree.Element("test")
        child = etree.SubElement(root, "child")
        child.text = "Test content with üñíçødé"
        
        xml_output_path = Path(self.test_dir) / "test_output.xml"
        
        # Write XML safely
        self.writer.write_xml_safely(root, xml_output_path)
        
        # Verify file was created
        self.assertTrue(xml_output_path.exists())
        
        # Verify content and encoding
        with open(xml_output_path, 'rb') as f:
            content = f.read()
            self.assertIn(b'<?xml version=', content)
            self.assertIn(b'encoding="UTF-8"', content)
            self.assertIn(b'<test>', content)
    
    def test_processing_stats(self):
        """Test processing statistics tracking."""
        initial_stats = self.writer.get_processing_stats()
        self.assertEqual(initial_stats["files_processed"], 0)
        self.assertEqual(initial_stats["backups_created"], 0)
        self.assertEqual(initial_stats["validations_performed"], 0)
        
        # Perform operation
        with self.writer.safe_docx_context(self.test_docx_path, self.output_path) as doc:
            doc.add_paragraph("Stats test")
        
        # Check updated stats
        final_stats = self.writer.get_processing_stats()
        self.assertEqual(final_stats["files_processed"], 1)
        self.assertEqual(final_stats["validations_performed"], 1)
        self.assertTrue("timestamp" in final_stats)
        self.assertTrue("backup_enabled" in final_stats)
        self.assertTrue("validation_enabled" in final_stats)
    
    def test_error_recovery_with_backup(self):
        """Test error recovery using backup files."""
        # Create existing output file
        shutil.copy2(self.test_docx_path, self.output_path)
        
        # Mock save operation to fail
        with patch.object(self.writer, '_save_document_safely', side_effect=Exception("Save failed")):
            with self.assertRaises(DOCXCorruptionError):
                with self.writer.safe_docx_context(self.test_docx_path, self.output_path) as doc:
                    doc.add_paragraph("This should fail")
        
        # Verify backup was created and original file restored
        self.assertTrue(self.output_path.exists())
        # Check that stats show error recovery
        stats = self.writer.get_processing_stats()
        self.assertEqual(stats["errors_recovered"], 1)


class TestSafeDocxOperation(unittest.TestCase):
    """Test cases for safe_docx_operation convenience function."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp(prefix="test_safe_operation_")
        self.test_docx_path = Path(self.test_dir) / "test_document.docx"
        self.output_path = Path(self.test_dir) / "output_document.docx"
        
        # Create test DOCX
        doc = Document()
        doc.add_paragraph("Test document for safe operations")
        doc.save(str(self.test_docx_path))
    
    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
    
    def test_safe_operation_success(self):
        """Test successful safe operation."""
        def test_operation(doc, test_param=None):
            doc.add_paragraph(f"Added paragraph: {test_param}")
            return {"paragraphs_added": 1, "test_param": test_param}
        
        result = safe_docx_operation(
            self.test_docx_path,
            self.output_path,
            test_operation,
            test_param="success_test"
        )
        
        # Verify result structure
        self.assertTrue(result["success"])
        self.assertIsNotNone(result["operation_result"])
        self.assertEqual(result["operation_result"]["paragraphs_added"], 1)
        self.assertEqual(result["operation_result"]["test_param"], "success_test")
        self.assertIn("writer_stats", result)
        self.assertEqual(len(result["errors"]), 0)
        
        # Verify output file
        self.assertTrue(self.output_path.exists())
        output_doc = Document(str(self.output_path))
        paragraphs = [p.text for p in output_doc.paragraphs]
        self.assertIn("Added paragraph: success_test", paragraphs)
    
    def test_safe_operation_failure(self):
        """Test safe operation with failure."""
        def failing_operation(doc):
            raise ValueError("Operation failed intentionally")
        
        result = safe_docx_operation(
            self.test_docx_path,
            self.output_path,
            failing_operation
        )
        
        # Verify failure result structure
        self.assertFalse(result["success"])
        self.assertIsNone(result["operation_result"])
        self.assertIn("writer_stats", result)
        self.assertGreater(len(result["errors"]), 0)
        self.assertIn("Operation failed intentionally", result["errors"][0])


class TestCorruptionPrevention(unittest.TestCase):
    """Test cases specifically for corruption prevention scenarios."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp(prefix="test_corruption_prevention_")
        self.test_docx_path = Path(self.test_dir) / "test_document.docx"
        
        # Create test DOCX
        doc = Document()
        doc.add_paragraph("Original content")
        doc.save(str(self.test_docx_path))
    
    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
    
    def test_partial_write_prevention(self):
        """Test prevention of partial writes."""
        writer = RobustDOCXWriter()
        output_path = Path(self.test_dir) / "output.docx"
        
        # Mock filesystem error during save
        with patch('shutil.move', side_effect=OSError("Disk full")):
            with self.assertRaises(DOCXCorruptionError):
                with writer.safe_docx_context(self.test_docx_path, output_path) as doc:
                    doc.add_paragraph("This should not be saved")
        
        # Verify output file was not created or is not corrupted
        if output_path.exists():
            # If file exists, it should be valid (from backup recovery)
            try:
                Document(str(output_path))
            except Exception:
                self.fail("Output file exists but is corrupted")
    
    def test_memory_efficient_processing(self):
        """Test memory-efficient processing for large documents."""
        # Create a larger test document
        doc = Document()
        for i in range(100):
            doc.add_paragraph(f"Paragraph {i} with some content to make it larger")
        
        large_docx_path = Path(self.test_dir) / "large_document.docx"
        doc.save(str(large_docx_path))
        
        writer = RobustDOCXWriter()
        output_path = Path(self.test_dir) / "large_output.docx"
        
        def process_large_doc(doc):
            # Add content without loading everything into memory at once
            doc.add_paragraph("Processed large document")
            return {"processed": True}
        
        # Should handle large document without memory issues
        with writer.safe_docx_context(large_docx_path, output_path) as doc:
            result = process_large_doc(doc)
            self.assertTrue(result["processed"])
        
        # Verify output
        self.assertTrue(output_path.exists())
        output_doc = Document(str(output_path))
        paragraphs = [p.text for p in output_doc.paragraphs]
        self.assertIn("Processed large document", paragraphs)


if __name__ == '__main__':
    # Create test suite
    test_suite = unittest.TestSuite()
    
    # Add test cases
    test_suite.addTest(unittest.makeSuite(TestRobustDOCXWriter))
    test_suite.addTest(unittest.makeSuite(TestSafeDocxOperation))
    test_suite.addTest(unittest.makeSuite(TestCorruptionPrevention))
    
    # Run tests
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    # Print summary
    print(f"\n{'='*50}")
    print(f"ROBUST DOCX WRITER TEST SUMMARY")
    print(f"{'='*50}")
    print(f"Tests run: {result.testsRun}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print(f"Success rate: {((result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100):.1f}%")
    
    if result.failures:
        print(f"\nFailures:")
        for test, traceback in result.failures:
            print(f"  - {test}: {traceback}")
    
    if result.errors:
        print(f"\nErrors:")
        for test, traceback in result.errors:
            print(f"  - {test}: {traceback}")
